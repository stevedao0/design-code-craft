"""
Công văn (Dispatch/Official Letter) API.

Migrated from F:\\VCPMC_BK\\APPS\\contract\\api\\apps\\bg_congvan\\routes.py
All logic adapted to F:\\APPs backend patterns (get_db, UserRow, etc.)

Endpoints:
  GET  /api/dispatches                    — list Công văn
  GET  /api/dispatches/{id}              — detail
  DELETE /api/dispatches/{id}            — delete
  GET  /api/dispatches/expired-contracts — expired Karaoke contracts
  GET  /api/dispatches/envelope-layout-config  — get envelope print settings
  PUT  /api/dispatches/envelope-layout-config  — save envelope print settings
  POST /api/dispatches/create-renewal     — create batch renewal letters
  POST /api/dispatches/batches/{id}/envelope — generate envelope DOCX
  POST /api/dispatches/batches/{id}/envelope-calibration — calibration sheet
  GET  /api/dispatches/{id}/logs          — get process logs
  POST /api/dispatches/{id}/logs         — add log entry
  GET  /api/dispatches/download/{year}/{filename} — download generated DOCX
"""
from __future__ import annotations

import json
import logging
import re
import zipfile
from calendar import monthrange
from datetime import date, datetime
from pathlib import Path
from typing import Any

# [DEBUG] Helper for NDJSON logging
def _debug_log(location: str, message: str, data: dict, hypothesis_id: str = "A"):
    """Append debug log entry to NDJSON file for this session."""
    try:
        log_path = Path(r"f:\APPs\debug-420958.log")
        entry = {
            "sessionId": "420958",
            "id": f"log_{int(datetime.now().timestamp()*1000)}",
            "timestamp": int(datetime.now().timestamp() * 1000),
            "location": location,
            "message": message,
            "data": data,
            "hypothesisId": hypothesis_id
        }
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from fastapi import APIRouter, Body, Depends, HTTPException, Query, status
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import HTTPAuthorizationCredentials
from lxml import etree
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.database import get_db
from app.core.security import decode_access_token, get_bearer_token, security_scheme
from app.models import (
    BgCongVanBatchRow,
    BgCongVanProcessRow,
    BgCongVanRow,
    ContractRecordRow,
    SystemSettingRow,
    UserRow,
)

logger = logging.getLogger(__name__)

from app.core.config import settings

_CONG_VAN_TEMPLATE_PATH = (
    Path(settings.export_template_root)
    / "Background"
    / "cong van_tai ky_karaoke.docx"
).resolve()

_NEW_KARAOKE_TEMPLATE_PATH = (
    Path(settings.export_template_root)
    / "Background"
    / "cong van_ky moi_karaoke.docx"
).resolve()

NEW_KARAOKE_PLACEHOLDERS = [
    "TEN_DON_VI",
    "DIA_CHI",
    "SO_CONG_VAN",
    "NGAY_KY_CONG_VAN",
    "THANG_KY_CONG_VAN",
    "NAM_KY_CONG_VAN",
]

NEW_KARAOKE_LEGACY_PLACEHOLDERS = [
    "so_cong_van",
    "ngay_ky_cong_van",
    "thang_ky_cong_van ",
    "thang_ky_cong_van",
    "nam_ky_cong_van",
]


def _repair_template_placeholders(template_xml: str) -> str:
    """Repair malformed Jinja2 tags split across Word XML runs.

    Mirrors the original VCPMC app renderer:
    F:\\VCPMC_BK\\APPS\\contract\\api\\domains\\common\\renderers\\docx\\text_renderer.py
    """
    def _fix_segment(seg: str) -> str:
        seg = re.sub(r"<w:proofErr[^>]*/>", "", seg)
        seg = re.sub(r"<w:(?:br|cr|tab)\b[^>]*/>", "", seg)
        seg = re.sub(
            r"</w:t>\s*</w:r>\s*(?:</w:proofErr>)?\s*<w:r[^>]*>\s*(?:<w:rPr>.*?</w:rPr>)?\s*<w:t[^>]*>",
            "",
            seg,
            flags=re.DOTALL,
        )
        seg = re.sub(
            r"</w:t>\s*</w:r>\s*<w:r[^>]*>\s*(?:<w:rPr>.*?</w:rPr>)?\s*<w:t[^>]*>",
            "",
            seg,
            flags=re.DOTALL,
        )
        return seg

    for pat in [
        re.compile(r"\{%.*?%\}", re.DOTALL),
        re.compile(r"\{\{.*?\}\}", re.DOTALL),
    ]:
        while True:
            m = pat.search(template_xml)
            if not m:
                break
            raw = m.group(0)
            fixed = _fix_segment(raw)
            if fixed == raw:
                break
            template_xml = template_xml[: m.start()] + fixed + template_xml[m.end():]

    return template_xml


def _normalize_docx_formatting(docx_path: Path) -> None:
    """Strip yellow highlight shading and force black font color after rendering.

    Mirrors the original VCPMC app post-processing:
    F:\\VCPMC_BK\\APPS\\contract\\api\\domains\\common\\renderers\\docx\\formatting.py
    """
    import shutil
    import tempfile

    tmp_dir = Path(tempfile.mkdtemp(prefix="docx_norm_"))
    tmp_zip = tmp_dir / "out.docx"
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_zip, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    try:
                        parser = etree.XMLParser(recover=True, huge_tree=True)
                        root = etree.fromstring(data, parser=parser)

                        for el in root.xpath(".//w:highlight", namespaces=ns):
                            parent = el.getparent()
                            if parent is not None:
                                parent.remove(el)

                        for el in root.xpath(".//w:shd", namespaces=ns):
                            fill = str(el.get("{" + ns["w"] + "}fill") or "").upper().strip()
                            if fill in {"FFFF00", "FFFF99", "FFF2CC", "FFEB9C"}:
                                parent = el.getparent()
                                if parent is not None:
                                    parent.remove(el)

                        for el in root.xpath(".//w:color", namespaces=ns):
                            el.set("{" + ns["w"] + "}val", "000000")

                        data = etree.tostring(
                            root,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone="yes",
                        )
                    except Exception:
                        pass
                zout.writestr(item, data)

        shutil.copyfile(tmp_zip, docx_path)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)





def _render_cong_van_from_template(
    *,
    template_path: Path,
    output_path: Path,
    context: dict,
) -> None:
    """Render a công văn DOCX from the Word template using docxtpl."""
    import tempfile

    with zipfile.ZipFile(template_path, "r") as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")

    repaired_xml = _repair_template_placeholders(doc_xml)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(template_path, "r") as zf_in, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zf_out:
            for item in zf_in.infolist():
                if item.filename == "word/document.xml":
                    zf_out.writestr(item, repaired_xml.encode("utf-8"))
                else:
                    zf_out.writestr(item, zf_in.read(item.filename))

        tpl = DocxTemplate(str(tmp_path))
        tpl.render(context)
        tpl.save(str(output_path))
        _normalize_docx_formatting(output_path)
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

router = APIRouter(prefix="/api/dispatches", tags=["dispatches"])
download_router = APIRouter(prefix="/api/dispatches/download", tags=["dispatches"])

# --- Envelope layout constants ---
_ENVELOPE_LAYOUT_SETTING_KEY = "BG_CONGVAN_ENVELOPE_LAYOUT_CONFIG"

# VCPMC Envelope preset - correct size: 230 x 170 mm (NOT C5 229x162)
# This is the actual VCPMC envelope size based on their template BIA THU.pdf
#
# Positioning uses recipient_box_top_mm (distance from page TOP), NOT bottom.
# This is more stable: when page_height changes, top stays fixed while bottom auto-calculates.
#
# Formula: resolved_top = recipient_box_top_mm - printer_offset_y_mm
#          resolved_bottom = page_height_mm - resolved_top
#
# For 230x170 envelope with top=110mm:
#   block at 110mm from top, block height=42mm, bottom = 170-110-42 = 18mm
VCPMC_ENVELOPE_230X170 = {
    "preset_name": "VCPMC Bia Thu 230x170mm",
    "page_width_mm": 230.0,           # Width: 230mm
    "page_height_mm": 170.0,          # Height: 170mm (NOT 162mm C5)
    "recipient_box_left_mm": 130.0,    # Start position for "Kinh gui" area (from left)
    "recipient_box_top_mm": 110.0,     # Top of recipient block from page top (110mm from top)
    "recipient_box_width_mm": 95.0,    # Width of recipient area
    "recipient_box_height_mm": 42.0,   # Height: name(13pt~6.5mm) + addr(8mm) + phone(8mm) + gaps(17mm) = ~40mm
    "line_spacing_mm": 8.0,           # Line spacing 8mm
    "printer_offset_x_mm": 0.0,
    "printer_offset_y_mm": 0.0,       # No offset by default - adjust per printer
    "non_printable_left_mm": 0.0,
    "non_printable_right_mm": 0.0,
    "non_printable_top_mm": 0.0,
    "non_printable_bottom_mm": 0.0,
}

# Canon LBP325x Printer Profile - Portrait 17x23, no offset
CANON_LBP325X_PRINTER_PROFILE = {
    "printer_profile": "canon_lbp325x",
    "safe_margin_mm": 0.0,
    "recipient_left_mm": 50.0,
    "recipient_top_mm": 120.0,
    "line_spacing_mm": 8.0,
}

# =============================================================================
# PRINTER PROFILE SYSTEM
# Profiles override layout + transform for specific printers.
# Default (legacy_default) uses legacy app logic unchanged.
# =============================================================================

PRINTER_PROFILES = {
    "legacy_default": {
        "id": "legacy_default",
        "name": "Mac dinh / App cu",
        "description": "Su dung layout app cu, khong doi gi",
        "transform_mode": "none",         # No transform
        "page_width_mm": 230.0,
        "page_height_mm": 162.0,
        "offset_x_mm": 0.0,
        "offset_y_mm": 0.0,
        "phone_on_envelope": False,
    },
    "canon_lbp325x": {
        "id": "canon_lbp325x",
        "name": "Canon LBP325x",
        "description": "Canon LBP325x / Portrait / Custom 17x23 cm / Khay giua",
        "transform_mode": "portrait",
        "paper_width_mm": 170.0,
        "paper_height_mm": 230.0,
        "safe_margin_mm": 0.0,
        "offset_x_mm": 0.0,
        "offset_y_mm": 0.0,
        "phone_on_envelope": False,
        # Portrait 17x23 - vi tri "Kinh gui" o giua-duoi
        "recipient_left_mm": 50.0,
        "recipient_top_mm": 120.0,
        "recipient_width_mm": 110.0,
        "recipient_height_mm": 35.0,
        "line_spacing_mm": 8.0,
    },
    "brother_hl_l2360d": {
        "id": "brother_hl_l2360d",
        "name": "Brother HL-L2360D",
        "description": "Brother HL-L2360D / Portrait / Custom 17x23 cm / Khay chinh",
        "transform_mode": "portrait",
        "paper_width_mm": 170.0,
        "paper_height_mm": 230.0,
        "safe_margin_mm": 0.0,
        "offset_x_mm": 0.0,
        "offset_y_mm": 0.0,
        "phone_on_envelope": False,
        # Portrait 17x23 - vi tri "Kinh gui" o giua-duoi (khoang 60% tu tren)
        "recipient_left_mm": 50.0,
        "recipient_top_mm": 120.0,
        "recipient_width_mm": 110.0,
        "recipient_height_mm": 35.0,
        "line_spacing_mm": 8.0,
    },
}

# Brother HL-L2360D transform modes (for testing)
BROTHER_TRANSFORM_MODES = {
    "portrait": {
        "id": "portrait",
        "name": "Portrait 170x230",
        "description": "Brother dung Portrait 170x230, khong dung Landscape",
    },
}

_ENVELOPE_PROFILE_SETTING_KEY = "BG_CONGVAN_ENVELOPE_PRINTER_PROFILE"


def _get_printer_profile(profile_id: str | None) -> dict:
    """Get printer profile by ID, default to legacy_default."""
    if not profile_id or profile_id not in PRINTER_PROFILES:
        return PRINTER_PROFILES["legacy_default"].copy()
    return PRINTER_PROFILES[profile_id].copy()


def _load_envelope_profile_config(db: Session) -> dict:
    """Load saved printer profile setting from DB."""
    setting = (
        db.query(SystemSettingRow)
        .filter(SystemSettingRow.key == _ENVELOPE_PROFILE_SETTING_KEY)
        .first()
    )
    if setting:
        try:
            val = json.loads(str(setting.value or "{}"))
            return val
        except Exception:
            pass
    return {"profile_id": "legacy_default", "brother_transform_mode": "portrait"}


def _save_envelope_profile_config(db: Session, config: dict, actor: str = "") -> dict:
    """Save printer profile setting to DB."""
    row = (
        db.query(SystemSettingRow)
        .filter(SystemSettingRow.key == _ENVELOPE_PROFILE_SETTING_KEY)
        .first()
    )
    now = datetime.now()
    json_val = json.dumps(config, ensure_ascii=False)
    if row:
        row.value = json_val
        row.updated_at = now
        row.updated_by = actor
    else:
        row = SystemSettingRow(
            key=_ENVELOPE_PROFILE_SETTING_KEY,
            value=json_val,
            updated_at=now,
            updated_by=actor,
        )
        db.add(row)
    db.commit()
    return config


def _apply_printer_profile_to_layout(
    base_layout: dict, profile_id: str, brother_transform_mode: str | None = None
) -> dict:
    """
    Apply printer profile overrides to base envelope layout.
    Returns a new layout dict with profile-specific values.
    """
    profile = _get_printer_profile(profile_id)
    layout = base_layout.copy()

    transform_mode = profile.get("transform_mode", "none")

    # Apply page size from profile
    if transform_mode == "portrait":
        # Brother uses Portrait 170x230 (NOT Landscape 230x170)
        layout["page_width_mm"] = profile.get("paper_width_mm", 170.0)
        layout["page_height_mm"] = profile.get("paper_height_mm", 230.0)
        # Use Brother Portrait coordinates (from top, not bottom)
        layout["recipient_box_left_mm"] = profile.get("recipient_left_mm", 55.0)
        layout["recipient_box_top_mm"] = profile.get("recipient_top_mm", 150.0)
        layout["recipient_box_width_mm"] = profile.get("recipient_width_mm", 105.0)
        layout["recipient_box_height_mm"] = profile.get("recipient_height_mm", 35.0)
        layout["line_spacing_mm"] = profile.get("line_spacing_mm", 8.0)
        # Mark as portrait mode
        layout["_brother_portrait_mode"] = True
    else:
        layout["page_width_mm"] = profile.get("page_width_mm", layout.get("page_width_mm", 230.0))
        layout["page_height_mm"] = profile.get("page_height_mm", layout.get("page_height_mm", 162.0))

    # Apply offset from profile (additive to existing)
    offset_x = profile.get("offset_x_mm", 0.0)
    offset_y = profile.get("offset_y_mm", 0.0)
    existing_x = layout.get("printer_offset_x_mm", 0.0)
    existing_y = layout.get("printer_offset_y_mm", 0.0)
    layout["printer_offset_x_mm"] = existing_x + offset_x
    layout["printer_offset_y_mm"] = existing_y + offset_y

    # Apply phone setting from profile
    layout["phone_on_envelope"] = profile.get("phone_on_envelope", False)

    # Store profile info in layout for reference
    layout["_printer_profile_id"] = profile_id
    layout["_brother_transform_mode"] = brother_transform_mode

    return layout


# =============================================================================
# ENVELOPE LAYOUT — IDENTICAL TO OLD APP (F:\VCPMC_BK)
# Page: 230mm x 162mm (landscape)
# Layout: bottom-anchored, NO safe margin
# phone_on_envelope: False by default (don't print phone on bia thu)
# =============================================================================
_ENVELOPE_LAYOUT_SETTING_KEY = "BG_CONGVAN_ENVELOPE_LAYOUT_CONFIG"
_DEFAULT_ENVELOPE_LAYOUT_CONFIG = {
    "page_width_mm": 230.0,
    "page_height_mm": 162.0,
    "recipient_box_left_mm": 130.0,
    "recipient_box_bottom_mm": 41.0,
    "recipient_box_width_mm": 95.0,
    "recipient_box_height_mm": 26.0,
    "line_spacing_mm": 8.0,
    "printer_offset_x_mm": 0.0,
    "printer_offset_y_mm": 0.0,
    "non_printable_left_mm": 0.0,
    "non_printable_right_mm": 0.0,
    "non_printable_top_mm": 0.0,
    "non_printable_bottom_mm": 0.0,
}

# =============================================================================
# BACKWARD-COMPATIBLE ALIASES — kept for existing endpoints that reference them
# =============================================================================

# VCPMC Bia Thu with PRE-PRINTED form lines (default for real-world envelopes)
VCPMC_BIA_THU_PRINTED = {
    "preset_name": "VCPMC Bia Thu 230x170 (physical)",
    "printer_profile": "generic",
    "page_width_mm": 230.0,
    "page_height_mm": 170.0,
    "safe_margin_mm": 0.0,   # 0 = legacy (no safe margin), 5mm = recommended
    "recipient_x_mm": 130.0,
    "recipient_width_mm": 95.0,
    "recipient_block_width_mm": 95.0,
    "recipient_block_height_mm": 28.0,
    "first_line_baseline_from_bottom_mm": 32.0,
    "font_baseline_offset_mm": 4.0,
    "line_gap_mm": 8.0,
    "font_name": "Times New Roman",
    "font_size_pt": 13.0,
    "phone_on_envelope": True,
    "phone_render_mode": "separate_line",
    "rotate_180": False,
    "printer_offset_x_mm": 0.0,
    "printer_offset_y_mm": 0.0,
}

_TRACKING_PROGRESS_LABELS = {
    "draft": "Nháp",
    "processing": "Đang theo dõi",
    "sent": "Đã gửi công văn",
    "closed": "Hoàn tất",
}
_CONTACT_ACTIONS = {
    "contacted", "follow_up", "sent", "delivery_success",
    "delivery_failed", "wrong_address", "renewed", "closed"
}

# File output root
_STORAGE_ROOT = Path(settings.export_output_root)
_CONGVAN_DOCX_ROOT = _STORAGE_ROOT / "docx"


# =============================================================================
# Auth helper
# =============================================================================

def _get_current_user(*, credentials: HTTPAuthorizationCredentials | None, db: Session) -> UserRow:
    token = get_bearer_token(credentials)
    username = decode_access_token(token)
    user = db.query(UserRow).filter(func.lower(UserRow.username) == username.lower()).first()
    if user is None:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="User not found")
    return user


# =============================================================================
# Helper utilities
# =============================================================================

def _format_dd_mm_yyyy(value: date | None) -> str:
    if value is None:
        return ""
    return value.strftime("%d/%m/%Y")

def _to_float_mm(value: Any, default: float) -> float:
    try:
        return float(value)
    except Exception:
        return default

def _normalize_envelope_layout_config(raw: dict | None) -> dict:
    src = raw or {}
    # Preserve preset_name if provided, otherwise use default
    preset_name = str(src.get("preset_name", "")) or _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("preset_name", "")

    def _pick(key: str, default: float) -> float:
        return _to_float_mm(src.get(key), default)

    # Recipient render mode: "printed_form_lines" (default) or "free_block"
    render_mode = str(src.get("recipient_render_mode", "") or _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_render_mode", "printed_form_lines"))
    # Phone render mode: "separate_line" (default) or "inline_address"
    phone_mode = str(src.get("phone_render_mode", "") or _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("phone_render_mode", "separate_line"))
    # phone_on_envelope: default True (print phone on bia thu)
    if "phone_on_envelope" in src:
        phone_on = bool(src.get("phone_on_envelope"))
    else:
        phone_on = bool(_DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("phone_on_envelope", True))

    page_h = _pick("page_height_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG["page_height_mm"])
    baseline_from_bottom = _pick("first_line_baseline_from_bottom_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("first_line_baseline_from_bottom_mm", 32.0))
    font_baseline_offset = _pick("font_baseline_offset_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("font_baseline_offset_mm", 4.0))
    # Auto-compute start_y from baseline if not explicitly provided
    baseline_y_from_top = page_h - baseline_from_bottom
    default_start_y = baseline_y_from_top - font_baseline_offset  # 170 - 32 - 4 = 134

    cfg = {
        "preset_name": preset_name,
        # Phone on envelope
        "phone_on_envelope": phone_on,
        "phone_render_mode": phone_mode,
        # Page
        "page_width_mm": _pick("page_width_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG["page_width_mm"]),
        "page_height_mm": page_h,
        # Physical positioning (for _generate_vcpmc_envelope_docx)
        "recipient_x_mm": _pick("recipient_x_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_x_mm", 130.0)),
        "recipient_width_mm": _pick("recipient_width_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_width_mm", 95.0)),
        # Baseline anchor
        "first_line_baseline_from_bottom_mm": baseline_from_bottom,
        "font_baseline_offset_mm": font_baseline_offset,
        "baseline_y_from_top_mm": baseline_y_from_top,
        # Line gap + font
        "line_gap_mm": _pick("line_gap_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("line_gap_mm", 8.0)),
        "font_name": str(src.get("font_name", "") or _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("font_name", "Times New Roman")),
        "font_size_pt": _pick("font_size_pt", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("font_size_pt", 13.0)),
        # Printer offset
        "printer_offset_x_mm": _pick("printer_offset_x_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG["printer_offset_x_mm"]),
        "printer_offset_y_mm": _pick("printer_offset_y_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG["printer_offset_y_mm"]),
        # Legacy box fields (kept for _migrate_envelope_layout compat + old generator)
        "recipient_box_left_mm": _pick("recipient_box_left_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_box_left_mm", 57.0)),
        "recipient_box_top_mm": _pick("recipient_box_top_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_box_top_mm", 110.0)),
        "recipient_box_bottom_mm": _pick("recipient_box_bottom_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_box_bottom_mm", 49.0)),
        "recipient_box_width_mm": max(40.0, _pick("recipient_box_width_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_box_width_mm", 95.0))),
        "recipient_box_height_mm": max(8.0, _pick("recipient_box_height_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_box_height_mm", 42.0))),
        "line_spacing_mm": _pick("line_spacing_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("line_spacing_mm", 8.0)),
        "non_printable_left_mm": _pick("non_printable_left_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("non_printable_left_mm", 0.0)),
        "non_printable_right_mm": _pick("non_printable_right_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("non_printable_right_mm", 0.0)),
        "non_printable_top_mm": _pick("non_printable_top_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("non_printable_top_mm", 0.0)),
        "non_printable_bottom_mm": _pick("non_printable_bottom_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("non_printable_bottom_mm", 0.0)),
        # Derived (for API response)
        "recipient_start_x_mm": _pick("recipient_x_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_start_x_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_x_mm", 57.0))),
        "recipient_start_y_mm": _pick("recipient_start_y_mm", default_start_y),
        "recipient_line_gap_mm": _pick("line_gap_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("line_gap_mm", 8.0)),
        "recipient_max_width_mm": _pick("recipient_width_mm", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("recipient_max_width_mm", 95.0)),
        "recipient_font_name": str(src.get("font_name", "") or _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("font_name", "Times New Roman")),
        "recipient_font_size_pt": _pick("font_size_pt", _DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("font_size_pt", 13.0)),
        # Rotation — critical for portrait vs landscape orientation
        "rotate_180": bool(src.get("rotate_180")) if "rotate_180" in src else bool(_DEFAULT_ENVELOPE_LAYOUT_CONFIG.get("rotate_180", False)),
    }
    print(f"[ENVELOPE_BASELINE] page_height={page_h} "
          f"baseline_from_bottom={baseline_from_bottom} "
          f"baseline_y_from_top={baseline_y_from_top} "
          f"font_baseline_offset={font_baseline_offset} "
          f"recipient_start_y={cfg['recipient_start_y_mm']} "
          f"phone_on_envelope={phone_on}")
    return cfg


def _migrate_envelope_layout(cfg: dict) -> dict:
    """Legacy stub — no migration needed. Identical to old app."""
    return dict(cfg)

def _resolve_envelope_layout(config: dict) -> dict:
    """
    Resolve envelope layout. Identical to F:\\VCPMC_BK\\APPS\\contract\\api\\apps\\bg_congvan\\routes.py

    Uses BOTTOM-based positioning:
      bottom = recipient_box_bottom_mm + printer_offset_y_mm
      top = page_h - bottom
      left = recipient_box_left_mm + printer_offset_x_mm
      right_indent = page_w - left - width
    """
    cfg = _normalize_envelope_layout_config(config)
    page_w = cfg["page_width_mm"]
    page_h = cfg["page_height_mm"]
    safe_left = min(page_w - 1.0, max(0.0, cfg["non_printable_left_mm"]))
    safe_right = min(page_w - 1.0, max(0.0, cfg["non_printable_right_mm"]))
    safe_top = min(page_h - 1.0, max(0.0, cfg["non_printable_top_mm"]))
    safe_bottom = min(page_h - 1.0, max(0.0, cfg["non_printable_bottom_mm"]))

    max_width = max(20.0, page_w - safe_left - safe_right)
    width = min(cfg["recipient_box_width_mm"], max_width)
    height = min(cfg["recipient_box_height_mm"], max(page_h - safe_top - safe_bottom, page_h - safe_top - cfg["line_spacing_mm"]))

    left_raw = cfg["recipient_box_left_mm"] + cfg["printer_offset_x_mm"]
    bottom_raw = cfg["recipient_box_bottom_mm"] + cfg["printer_offset_y_mm"]

    left = min(max(left_raw, safe_left), page_w - safe_right - width)
    min_bottom_anchor = safe_bottom + cfg["line_spacing_mm"]
    max_bottom_anchor = page_h - safe_top
    bottom = min(max(bottom_raw, min_bottom_anchor), max_bottom_anchor)
    top = page_h - bottom
    right_indent = max(0.0, page_w - left - width)

    return {
        **cfg,
        "resolved_left_mm": left,
        "resolved_bottom_mm": bottom,
        "resolved_top_mm": top,
        "resolved_width_mm": width,
        "resolved_height_mm": height,
        "resolved_right_indent_mm": right_indent,
    }

def _load_envelope_layout_config(*, db: Session) -> dict:
    """Load from DB and resolve. Identical to old app."""
    row = db.query(SystemSettingRow).filter(SystemSettingRow.key == _ENVELOPE_LAYOUT_SETTING_KEY).first()
    parsed = {}
    if row and getattr(row, "value", None):
        try:
            parsed = json.loads(str(getattr(row, "value", "") or "{}"))
        except Exception:
            parsed = {}
    print(f"[ENVELOPE_CONFIG_RESET_TO_LEGACY] before={parsed}")
    resolved = _resolve_envelope_layout(parsed)
    print(f"[ENVELOPE_CONFIG_RESET_TO_LEGACY] after=page_w={resolved.get('page_width_mm')} page_h={resolved.get('page_height_mm')} left={resolved.get('resolved_left_mm')} top={resolved.get('resolved_top_mm')} bottom={resolved.get('resolved_bottom_mm')}")
    return resolved

def _save_envelope_layout_config(*, db: Session, config: dict, actor: str | None) -> dict:
    """Save envelope layout config. Identical to old app."""
    normalized = _normalize_envelope_layout_config(config)
    row = db.query(SystemSettingRow).filter(SystemSettingRow.key == _ENVELOPE_LAYOUT_SETTING_KEY).first()
    serialized = json.dumps(normalized, ensure_ascii=False)
    if row is None:
        row = SystemSettingRow(key=_ENVELOPE_LAYOUT_SETTING_KEY, value=serialized)
        db.add(row)
    else:
        row.value = serialized
    try:
        db.commit()
    except Exception as e:
        db.rollback()
        raise
    print(f"[ENVELOPE_CONFIG_SAVED] page_w={normalized.get('page_width_mm')} page_h={normalized.get('page_height_mm')} left={normalized.get('recipient_box_left_mm')} bottom={normalized.get('recipient_box_bottom_mm')}")
    return _resolve_envelope_layout(normalized)

def _download_url_for_path(path_value: str | Path | None, *, year: int | None = None) -> str:
    raw = str(path_value or "").strip()
    if not raw:
        return ""
    p = Path(raw)
    resolved_year = int(year or 0)
    if resolved_year <= 0:
        m = re.search(r"[/\\](\d{4})[/\\]", raw)
        if m:
            try:
                resolved_year = int(m.group(1))
            except Exception:
                resolved_year = 0
    if resolved_year <= 0:
        resolved_year = date.today().year
    return f"/api/dispatches/download/{resolved_year}/{p.name}"

def _slug_name(value: str | None) -> str:
    text = str(value or "").strip().lower()
    if not text:
        return "don-vi"
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"\s+", "-", text)
    text = re.sub(r"-{2,}", "-", text)
    return text.strip("-") or "don-vi"


def _coerce_new_karaoke_rows(rows_raw: Any) -> tuple[list[dict[str, str]], list[dict[str, Any]]]:
    rows: list[dict[str, str]] = []
    issues: list[dict[str, Any]] = []
    if not isinstance(rows_raw, list):
        return rows, [{"row_index": 0, "missing_fields": ["rows"], "message": "Thiếu danh sách đơn vị ký mới"}]

    for idx, raw in enumerate(rows_raw, start=1):
        if not isinstance(raw, dict):
            issues.append({"row_index": idx, "missing_fields": ["row"], "message": "Dòng dữ liệu không hợp lệ"})
            continue
        row = {
            "stt": str(raw.get("stt") or "").strip(),
            "ten_don_vi": _normalize_unit_name(str(raw.get("ten_don_vi") or raw.get("TEN_DON_VI") or "")),
            "dia_chi": _normalize_address_spacing(str(raw.get("dia_chi") or raw.get("DIA_CHI") or "")),
            "so_cong_van": str(raw.get("so_cong_van") or raw.get("SO_CONG_VAN") or "").strip(),
            "ngay_ky_cong_van": str(raw.get("ngay_ky_cong_van") or raw.get("NGAY_KY_CONG_VAN") or "").strip(),
            "thang_ky_cong_van": str(raw.get("thang_ky_cong_van") or raw.get("THANG_KY_CONG_VAN") or "").strip(),
            "nam_ky_cong_van": str(raw.get("nam_ky_cong_van") or raw.get("NAM_KY_CONG_VAN") or "").strip(),
            # optional envelope-only fields
            "so_dien_thoai": str(raw.get("so_dien_thoai") or raw.get("SO_DIEN_THOAI")
                                 or raw.get("sdt") or raw.get("SDT")
                                 or raw.get("dien_thoai") or raw.get("DIEN_THOAI")
                                 or raw.get("phone") or raw.get("PHONE")
                                 or raw.get("mobile") or raw.get("MOBILE")
                                 or raw.get("tel") or raw.get("TEL")
                                 or "").strip(),
            "nguoi_nhan_bia_thu": str(raw.get("nguoi_nhan_bia_thu") or raw.get("NGUOI_NHAN_BIA_THU")
                                        or raw.get("nguoi_nhan") or raw.get("NGUOI_NHAN")
                                        or raw.get("nguoi_nhan_bia_thu_") or ""
                                        ).strip(),
            # raw_line: original pasted line, used as final fallback for phone extraction
            "raw_line": str(raw.get("raw_line") or raw.get("raw_text") or raw.get("raw") or "").strip(),
        }
        missing_fields: list[str] = []
        if not row["ten_don_vi"]:
            missing_fields.append("TEN_DON_VI")
        if not row["dia_chi"]:
            missing_fields.append("DIA_CHI")
        if missing_fields:
            issues.append({
                "row_index": idx,
                "missing_fields": missing_fields,
                "message": "Thiếu dữ liệu bắt buộc theo template ký mới Karaoke",
            })
        rows.append(row)
    return rows, issues


def _build_new_karaoke_context(*, row: dict[str, str], cong_van_no: str, issue_date: date) -> dict[str, str]:
    so_cong_van = str(cong_van_no or row.get("so_cong_van") or "").strip()
    ngay = f"{issue_date.day:02d}" if issue_date else str(row.get("ngay_ky_cong_van") or "").strip()
    thang = f"{issue_date.month:02d}" if issue_date else str(row.get("thang_ky_cong_van") or "").strip()
    nam = str(issue_date.year) if issue_date else str(row.get("nam_ky_cong_van") or "").strip()
    return {
        "TEN_DON_VI": row.get("ten_don_vi", ""),
        "DIA_CHI": row.get("dia_chi", ""),
        "SO_CONG_VAN": so_cong_van,
        "NGAY_KY_CONG_VAN": ngay,
        "THANG_KY_CONG_VAN": thang,
        "NAM_KY_CONG_VAN": nam,
        # legacy fallback for old template variants still present in split XML runs
        "so_cong_van": so_cong_van,
        "ngay_ky_cong_van": ngay,
        "thang_ky_cong_van": thang,
        "thang_ky_cong_van ": thang,
        "nam_ky_cong_van": nam,
    }

def _normalize_unit_name(value: str | None) -> str:
    """Normalize unit name: strip, collapse whitespace. Do NOT change case."""
    return re.sub(r"\s+", " ", str(value or "").strip())

def _normalize_address_spacing(value: str | None) -> str:
    """Normalize address: strip, collapse whitespace. Do NOT change case."""
    return re.sub(r"\s+", " ", str(value or "").strip())

def _is_karaoke_contract_row(row: ContractRecordRow) -> bool:
    field_code = str(getattr(row, "field_code", "") or "").strip().lower()
    linh_vuc = str(getattr(row, "linh_vuc", "") or "").strip().upper()
    template_code = str(getattr(row, "template_code", "") or "").strip().upper()
    return (
        field_code == "karaoke"
        or linh_vuc in {"SCKS", "KARAOKE"}
        or template_code == "SCKS"
    )

def _to_single_line(value: str | None) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())

def _to_payload_dict(raw: str | None) -> dict:
    text = str(raw or "").strip()
    if not text:
        return {}
    try:
        data = json.loads(text)
    except Exception:
        return {}
    return data if isinstance(data, dict) else {}

def _pick_first_non_empty(*values: object) -> str:
    for value in values:
        normalized = _to_single_line(str(value or ""))
        if normalized:
            return normalized
    return ""

def _parse_int_range(value: int | None, default: int, min_value: int, max_value: int) -> int:
    try:
        parsed = int(value or default)
    except Exception:
        parsed = int(default)
    return max(min_value, min(parsed, max_value))

def _resolve_expiry_date_range(
    *, filter_mode: str, filter_year: int,
    filter_quarter: int, filter_month: int, filter_week: int,
) -> tuple[date | None, date | None]:
    mode = str(filter_mode or "year").strip().lower()
    year = _parse_int_range(filter_year, date.today().year, 2000, 2100)
    quarter = _parse_int_range(filter_quarter, 1, 1, 4)
    month = _parse_int_range(filter_month, 1, 1, 12)
    week = _parse_int_range(filter_week, 1, 1, 4)
    if mode == "all":
        return None, None
    if mode == "year":
        return date(year, 1, 1), date(year, 12, 31)
    if mode == "quarter":
        start_month = (quarter - 1) * 3 + 1
        end_month = start_month + 2
        end_day = monthrange(year, end_month)[1]
        return date(year, start_month, 1), date(year, end_month, end_day)
    if mode == "month":
        end_day = monthrange(year, month)[1]
        return date(year, month, 1), date(year, month, end_day)
    if mode == "week":
        end_day_month = monthrange(year, month)[1]
        start_day = (week - 1) * 7 + 1
        if week == 4:
            end_day = end_day_month
        else:
            end_day = min(week * 7, end_day_month)
        start_day = min(start_day, end_day_month)
        return date(year, month, start_day), date(year, month, end_day)
    return date(year, 1, 1), date(year, 12, 31)

def _derive_progress_fields(*, status: str, action_types: list[str]) -> dict[str, Any]:
    actions = [str(a or "").strip().lower() for a in action_types if str(a or "").strip()]
    status_norm = str(status or "").strip().lower()
    latest_action = actions[0] if actions else ""
    contacted = any(a in _CONTACT_ACTIONS for a in actions)
    contact_status = "Đã liên hệ" if contacted else "Chưa liên hệ"
    if "wrong_address" in actions:
        delivery_status = "Sai địa chỉ"
    elif "delivery_success" in actions:
        delivery_status = "Đã đến"
    elif "delivery_failed" in actions:
        delivery_status = "Không đến"
    elif "sent" in actions:
        delivery_status = "Đã gửi - chờ phản hồi"
    else:
        delivery_status = "Chưa gửi"
    if "renewed" in actions:
        tracking_progress = "Đã tái ký"
    else:
        tracking_progress = _TRACKING_PROGRESS_LABELS.get(status_norm, "Đang theo dõi")
    return {
        "contacted": contacted,
        "contact_status": contact_status,
        "delivery_status": delivery_status,
        "tracking_progress": tracking_progress,
        "latest_action": latest_action or "-",
    }


# =============================================================================
# Envelope generation helpers (pure python-docx, no external template needed)
# =============================================================================

def _set_envelope_run_font(run, *, size_pt: int = 13) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")

def _prepare_envelope_doc(doc: Document, *, layout: dict) -> None:
    section = doc.sections[0]
    page_w = float(layout["page_width_mm"])
    page_h = float(layout["page_height_mm"])
    if page_w >= page_h:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(page_w)
        section.page_height = Mm(page_h)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(page_w)
        section.page_height = Mm(page_h)
    # Old app: no margins — coordinates are absolute from page edges
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

def _apply_envelope_block(paragraph, *, lines: list[str], layout: dict, page_break_before: bool = False) -> None:
    """
    Render recipient text on a paragraph.

    Two modes:
    - "printed_form_lines" (default): the bia thu already has the printed dotted
      lines. We render text at top-left anchor (resolved_start_x_mm, resolved_start_y_mm)
      with NO border, NO big block. Use for real envelopes.
    - "free_block" (legacy): a free-floating block at resolved_left/top with given
      width/height. Use for testing only.
    """
    paragraph.text = ""
    p_format = paragraph.paragraph_format
    p_format.space_after = Mm(0)
    p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    render_mode = str(layout.get("recipient_render_mode", "printed_form_lines") or "printed_form_lines")

    # Check for Brother Portrait mode
    is_brother_portrait = layout.get("_brother_portrait_mode", False)

    if render_mode == "printed_form_lines":
        # Coordinate system: top-left of the PRINTABLE AREA (after margins).
        # safe_margin_mm is set in _prepare_envelope_doc and propagated via layout.
        safe_margin_mm = float(layout.get("safe_margin_mm", 0.0))

        # Absolute position from top-left of page
        start_x_abs = float(layout.get("resolved_start_x_mm", 57.0))
        start_y_abs = float(layout.get("resolved_start_y_mm", 118.0))
        line_gap = float(layout.get("recipient_line_gap_mm", 8.0))
        max_w = float(layout.get("recipient_max_width_mm", 120.0))
        font_size = float(layout.get("recipient_font_size_pt", 13.0))
        page_w = float(layout.get("page_width_mm", 230.0))

        # Offset to be relative to printable area
        # If start_x < safe_margin, clamp to 0 (won't happen with real values)
        word_left_indent = max(0.0, start_x_abs - safe_margin_mm)
        # If start_y < safe_margin, clamp to 0
        word_space_before = max(0.0, start_y_abs - safe_margin_mm)
        # right_indent: page right edge (absolute) minus safe margin minus max_w
        right_indent = max(0.0, page_w - safe_margin_mm - word_left_indent - max_w)

        p_format.left_indent = Mm(word_left_indent)
        p_format.right_indent = Mm(right_indent)
        p_format.space_before = Mm(word_space_before)
        p_format.line_spacing = Mm(line_gap)

        print(f"[ENVELOPE_PRINTABLE_MARGIN] "
              f"page={page_w:.0f}x170 "
              f"safe_margin={safe_margin_mm} "
              f"desired_x={start_x_abs} desired_y={start_y_abs} "
              f"word_left_indent={word_left_indent} "
              f"word_space_before={word_space_before} "
              f"margins={safe_margin_mm}/{safe_margin_mm}/{safe_margin_mm}/{safe_margin_mm}")
    elif is_brother_portrait:
        # Brother Portrait mode: use top-based coordinates (NOT bottom-based)
        # page_size = 170x230 Portrait
        brother_left = float(layout.get("recipient_box_left_mm", 55.0))
        brother_top = float(layout.get("recipient_box_top_mm", 150.0))
        brother_line_spacing = float(layout.get("line_spacing_mm", 8.0))
        font_size = 13.0

        p_format.left_indent = Mm(brother_left)
        p_format.right_indent = Mm(0)
        p_format.space_before = Mm(brother_top)
        p_format.line_spacing = Mm(brother_line_spacing)

        print(f"[BROTHER_PORTRAIT] left={brother_left} top={brother_top} spacing={brother_line_spacing}")
    else:
        # Legacy free block
        p_format.left_indent = Mm(layout["resolved_left_mm"])
        p_format.right_indent = Mm(layout["resolved_right_indent_mm"])
        p_format.space_before = Mm(layout["resolved_top_mm"])
        p_format.line_spacing = Mm(layout["line_spacing_mm"])
        font_size = 13.0

    if page_break_before:
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        _set_envelope_run_font(run, size_pt=int(font_size))
        if idx < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)

def _trim_document_trailing_blanks(doc: Document) -> None:
    """Remove trailing empty paragraphs from a document (envelope only)."""
    body = doc.element.body
    children = list(body)
    for element in reversed(children):
        tag = str(getattr(element, "tag", ""))
        if tag.endswith("}sectPr"):
            continue
        if tag.endswith("}tbl"):
            break
        if tag.endswith("}p"):
            texts = element.xpath(".//*[local-name()='t']")
            has_text = any(str(t.text or "").strip() for t in texts)
            has_draw = bool(element.xpath(".//*[local-name()='drawing' or local-name()='pict']"))
            if has_text or has_draw:
                break
        else:
            break
        body.remove(element)

def _merge_docx_files(*, inputs: list[Path], output: Path) -> None:
    """Merge rendered công văn DOCX files using docxcompose.

    Each rendered công văn is a complete document with its own layout (A4, headers, tables).
    We add a page-break + explicit section-break paragraph before each appended doc so that
    each công văn starts on a new page with its own section properties.
    """
    if not inputs:
        return
    if len(inputs) == 1:
        import shutil
        output.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(inputs[0]), str(output))
        return

    # Pre-load all source docs so we can extract their sectPr for the section break
    src_docs = [Document(str(p)) for p in inputs]

    master = src_docs[0]
    composer = Composer(master)

    for i, path in enumerate(inputs[1:], start=1):
        src = src_docs[i]

        # --- Build sectPr paragraph: ends the previous section, starts a new page ---
        # Strategy: add a paragraph with a sectPr that has pageBreakBefore and nextPage type.
        # This ensures the appended doc starts in a fresh section on a new page.
        pg_break_para = master.add_paragraph()
        run = pg_break_para.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Add a sectPr paragraph to close the previous section (nextPage)
        sect_para = master.add_paragraph()
        pPr = sect_para._p.get_or_add_pPr()
        sectPr = _make_sectPr_element(src)
        pPr.append(sectPr)

        # Now append the source document body (docxcompose skips sectPr elements)
        composer.append(src)

    output.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output))


def _make_sectPr_element(src_doc: Document) -> Any:
    """Build a sectPr element with nextPage type, copying page layout from src_doc."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    sectPr = OxmlElement("w:sectPr")

    # Section type: nextPage (explicit page break between sections)
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sectPr.append(type_el)

    # Copy page size and margins from source doc's final sectPr
    src_sectPr_list = src_doc.element.body.findall(qn("w:sectPr"))
    if src_sectPr_list:
        src_sectPr = src_sectPr_list[-1]
        for child_tag in ["w:pgSz", "w:pgMar", "w:cols", "w:docGrid"]:
            for child in src_sectPr.findall(qn(child_tag)):
                sectPr.append(deepcopy(child))

    return sectPr

def _normalize_envelope_recipient(row: dict | Any) -> dict[str, str]:
    """
    Normalize one recipient dict for the envelope block, with FALLBACK phone extraction.

    Accepts a row from various sources (frontend payload, ORM object, dict).
    Returns a clean dict: {"name": ..., "address": ..., "phone": ...}

    Fallback chain for phone:
      1) explicit phone field (so_dien_thoai / phone / contact_phone / phone_number / recipient_phone)
      2) extract from raw_line (if provided)
      3) extract from address (and remove from address)

    Fallback chain for name:
      - ten_don_vi / TEN_DON_VI / name / recipient_unit
    Fallback chain for address:
      - dia_chi / DIA_CHI / address / recipient_address
    """
    def _g(*keys: str) -> str:
        for k in keys:
            try:
                v = row.get(k) if hasattr(row, "get") else getattr(row, k, None)
            except Exception:
                v = None
            if v is not None and str(v).strip():
                return str(v).strip()
        return ""

    name = _g("ten_don_vi", "TEN_DON_VI", "name", "recipient_unit", "don_vi_nguoi_nhan")
    address = _g("dia_chi", "DIA_CHI", "address", "recipient_address")
    phone = _g("so_dien_thoai", "phone", "contact_phone", "phone_number",
               "recipient_phone", "dien_thoai", "sdt")
    raw_line = _g("raw_line", "raw_text", "raw")

    # Fallback: extract phone from raw_line or address if empty
    if not phone:
        embedded = re.compile(r'(?:\+?84[\s.\-]?|0)(?:\d[\s.\-]?){9,11}(?=\D|$)')
        search_text = raw_line or address or ""
        if search_text:
            m = embedded.findall(search_text)
            if m:
                phone = " / ".join(s.strip() for s in m)
                # If extracted from address, strip it from address
            if m:
                phone = " / ".join(s.strip() for s in m)
                # If extracted from address, strip it from address
                if not raw_line and address:
                    for s in m:
                        address = address.replace(s, " ")
                    address = re.sub(r'[\s,;.\-]+$', '', address).strip()
                elif raw_line:
                    # Try to find a clean name/address split using the phone boundary
                    cleaned = raw_line
                    for s in m:
                        cleaned = cleaned.replace(s, " ")
                    cleaned = re.sub(r'[\s,;.\-]+$', '', cleaned).strip()
                    if not name and not address:
                        # split cleaned at first numeric segment
                        mm = re.match(r'^(.+?)\s+(\d.*)$', cleaned)
                        if mm:
                            name = mm.group(1).strip()
                            address = mm.group(2).strip()
                        else:
                            name = cleaned
                    elif not address:
                        # Use cleaned as address (strip from end)
                        if cleaned:
                            address = cleaned

    return {
        "name": name,
        "address": address,
        "phone": phone,
    }


def _build_envelope_recipient_lines(
    *, recipient_unit, recipient_address, recipient_contact=None, recipient_phone=None,
    phone_render_mode: str = "inline_address",
    phone_on_envelope: bool = False,
) -> list[str]:
    """
    Build envelope block lines.

    phone_on_envelope (default False):
      - When False: phone is NOT printed on the bia thu. Only "name" + "address".
      - When True: phone may be printed per phone_render_mode.

    phone_render_mode (only used when phone_on_envelope=True):
      - "inline_address": phone is appended to address line as " - DT: xxx"
      - "separate_line": phone is rendered as a 3rd line "DT: xxx"
    """
    lines: list[str] = []
    unit = _to_single_line(recipient_unit)
    address = _to_single_line(recipient_address)
    phone = _to_single_line(recipient_phone) if phone_on_envelope else ""

    if phone_on_envelope and phone_render_mode == "inline_address":
        # Render: name / address - DT: phone
        if unit:
            lines.append(unit)
        if address and phone:
            lines.append(f"{address} - ĐT: {phone}")
        elif address:
            lines.append(address)
        elif phone:
            lines.append(f"ĐT: {phone}")
    elif phone_on_envelope and phone_render_mode == "separate_line":
        # 3 lines: name / address / DT: phone
        if unit:
            lines.append(unit)
        if address:
            lines.append(address)
        if phone:
            lines.append(f"ĐT: {phone}")
    else:
        # Default: name + address only, no phone
        if unit:
            lines.append(unit)
        if address:
            lines.append(address)
    return lines or ["-"]

def _build_envelope_docx(*, recipients: list[dict], output: Path, layout: dict, calibration: bool = False) -> None:
    doc = Document()
    _prepare_envelope_doc(doc, layout=layout)
    # DEBUG: envelope recipient count
    print(f"[ENVELOPE_RECIPIENT_COUNT] {len(recipients)}")
    print(f"[ENVELOPE_LAYOUT_DEBUG] page_width={layout.get('page_width_mm')} page_height={layout.get('page_height_mm')} "
          f"resolved_left={layout.get('resolved_left_mm')} resolved_bottom={layout.get('resolved_bottom_mm')} "
          f"resolved_top={layout.get('resolved_top_mm')} line_spacing={layout.get('line_spacing_mm')}")
    if recipients:
        first_norm = _normalize_envelope_recipient(recipients[0])
        phone_mode = str(layout.get("phone_render_mode", "inline_address") or "inline_address")
        phone_on = bool(layout.get("phone_on_envelope", False))
        first_lines = _build_envelope_recipient_lines(
            recipient_unit=first_norm["name"],
            recipient_address=first_norm["address"],
            recipient_phone=first_norm["phone"],
            phone_render_mode=phone_mode,
            phone_on_envelope=phone_on,
        )
        try:
            print(f"[ENVELOPE_FIRST_RECIPIENT_LINES] {first_lines}")
        except UnicodeEncodeError:
            print(f"[ENVELOPE_FIRST_RECIPIENT_LINES] {len(first_lines)} lines (cannot print due to encoding)")
    if not recipients:
        recipients = [{"recipient_unit": "", "recipient_address": "", "recipient_contact": "", "recipient_phone": ""}]
    for idx, recipient in enumerate(recipients):
        # SAFETY NET: normalize + fallback-extract phone from raw_line or address.
        # This guarantees the bìa thư always renders phone when one is reachable
        # in any of the input fields.
        normalized = _normalize_envelope_recipient(recipient)
        unit = normalized["name"]
        address = normalized["address"]
        phone = normalized["phone"]

        # Also pass through any extra fields (contact) the caller may have provided
        contact = recipient.get("recipient_contact", "") or recipient.get("contact", "") or ""
        if not phone:
            # last-ditch: pull from raw-style fields
            phone = str(recipient.get("recipient_phone", "") or "").strip()

        phone_mode = str(layout.get("phone_render_mode", "inline_address") or "inline_address")
        phone_on = bool(layout.get("phone_on_envelope", False))
        lines = _build_envelope_recipient_lines(
            recipient_unit=unit,
            recipient_address=address,
            recipient_contact=contact,
            recipient_phone=phone,
            phone_render_mode=phone_mode,
            phone_on_envelope=phone_on,
        )
        if calibration:
            lines = [
                "[CALIBRATION - GOC TRAI TREN]",
                f"LEFT={layout['resolved_left_mm']:.1f}mm | BOTTOM={layout['resolved_bottom_mm']:.1f}mm",
                f"WIDTH={layout['resolved_width_mm']:.1f}mm | LINE={layout['line_spacing_mm']:.1f}mm",
                "[CALIBRATION - GOC PHAI DUOI]",
            ]
        target_para = doc.paragraphs[0] if idx == 0 and doc.paragraphs else doc.add_paragraph()
        _apply_envelope_block(target_para, lines=lines, layout=layout, page_break_before=(idx > 0))
    _trim_document_trailing_blanks(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

def _build_envelope_recipients_from_rows(rows: list[BgCongVanRow]) -> list[dict]:
    """
    Build per-row recipient dict for the envelope block.

    Layout:
      Line 1: recipient_unit (Tên đơn vị)
      Line 2: recipient_address (Địa chỉ)
      Line 3: ĐT: recipient_phone  [only if non-empty]

    For new-karaoke batches, recipient_unit, recipient_address, and recipient_phone are
    stored in BgCongVanRow at creation time (see create_new_karaoke_batch).

    The normalizer applies a final fallback: if recipient_phone is empty but the
    address still contains a phone, it is extracted and removed from the address.
    """
    out: list[dict] = []
    for r in rows:
        raw = {
            "recipient_unit": str(getattr(r, "recipient_unit", "") or ""),
            "recipient_address": str(getattr(r, "recipient_address", "") or ""),
            "recipient_contact": str(getattr(r, "recipient_contact", "") or ""),
            "recipient_phone": str(getattr(r, "recipient_phone", "") or ""),
        }
        normalized = _normalize_envelope_recipient(raw)
        raw["recipient_unit"] = normalized["name"]
        raw["recipient_address"] = normalized["address"]
        raw["recipient_phone"] = normalized["phone"]
        out.append(raw)
    return out


def _build_envelope_docx_from_row_data(
    *, row_data_list: list[dict], output: Path, layout: dict, calibration: bool = False
) -> None:
    """
    Build envelope DOCX directly from raw row_data dicts (from API payload).

    This bypasses the DB so it retains raw_line / so_dien_thoai fields for phone
    extraction even when the caller only passed partial data.

    row_data_list: list of dicts with keys:
      - ten_don_vi / name / recipient_unit
      - dia_chi / address / recipient_address
      - so_dien_thoai / phone / raw_line (optional, used for phone fallback)
      - nguoi_nhan_bia_thu / recipient_contact
    """
    recipients: list[dict] = []
    for row_data in row_data_list:
        # _normalize_envelope_recipient already handles all fallbacks including raw_line
        normalized = _normalize_envelope_recipient(row_data)
        recipients.append({
            "recipient_unit": normalized["name"],
            "recipient_address": normalized["address"],
            "recipient_contact": str(row_data.get("recipient_contact") or row_data.get("contact") or row_data.get("nguoi_nhan_bia_thu") or ""),
            "recipient_phone": normalized["phone"],
        })

    doc = Document()
    _prepare_envelope_doc(doc, layout=layout)
    # DEBUG: envelope recipient count
    print(f"[ENVELOPE_RECIPIENT_COUNT] {len(row_data_list)}")
    if row_data_list:
        first_norm = _normalize_envelope_recipient(row_data_list[0])
        phone_mode = str(layout.get("phone_render_mode", "inline_address") or "inline_address")
        phone_on = bool(layout.get("phone_on_envelope", False))
        first_lines = _build_envelope_recipient_lines(
            recipient_unit=first_norm["name"],
            recipient_address=first_norm["address"],
            recipient_phone=first_norm["phone"],
            phone_render_mode=phone_mode,
            phone_on_envelope=phone_on,
        )
        try:
            print(f"[ENVELOPE_FIRST_RECIPIENT_LINES] {first_lines}")
        except UnicodeEncodeError:
            print(f"[ENVELOPE_FIRST_RECIPIENT_LINES] {len(first_lines)} lines (cannot print due to encoding)")
    if not recipients:
        recipients = [{"recipient_unit": "", "recipient_address": "", "recipient_contact": "", "recipient_phone": ""}]
    for idx, recipient in enumerate(recipients):
        phone_mode = str(layout.get("phone_render_mode", "inline_address") or "inline_address")
        phone_on = bool(layout.get("phone_on_envelope", False))
        lines = _build_envelope_recipient_lines(
            recipient_unit=recipient.get("recipient_unit"),
            recipient_address=recipient.get("recipient_address"),
            recipient_contact=recipient.get("recipient_contact"),
            recipient_phone=recipient.get("recipient_phone"),
            phone_render_mode=phone_mode,
            phone_on_envelope=phone_on,
        )
        if calibration:
            lines = [
                "[CALIBRATION - GOC TRAI TREN]",
                f"LEFT={layout['resolved_left_mm']:.1f}mm | BOTTOM={layout['resolved_bottom_mm']:.1f}mm",
                f"WIDTH={layout['resolved_width_mm']:.1f}mm | LINE={layout['line_spacing_mm']:.1f}mm",
                "[CALIBRATION - GOC PHAI DUOI]",
            ]
        target_para = doc.paragraphs[0] if idx == 0 and doc.paragraphs else doc.add_paragraph()
        _apply_envelope_block(target_para, lines=lines, layout=layout, page_break_before=(idx > 0))
    _trim_document_trailing_blanks(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

def _generate_vcpmc_envelope_docx(
    *,
    recipients: list[dict],
    output: Path,
    layout: dict,
    calibration: bool = False,
) -> None:
    """
    Generator bia thu VCPMC theo mau vat ly.

    Layout chuan:
      page_size: 230mm x 170mm (landscape)
      recipient_x_mm: 130mm (from left)
      first_line_baseline_from_bottom_mm: 32mm
      line_gap_mm: 8mm
      font_name: Times New Roman
      font_size_pt: 13
      font_baseline_offset_mm: 4mm

    Paragraph approach (RELIABLE):
      - Each recipient = 1 paragraph with LINE breaks between lines
      - Paragraph has: left_indent=130mm, space_before=134mm, exact line spacing 8mm
      - Page break BEFORE each new recipient (not after last) using doc.add_page_break()
      - NO table, NO extra paragraphs, NO page break paragraphs

    Lines:
      [0] ten_don_vi
      [1] dia_chi
      [2] "DT: xxx"  (if phone_on and phone)
    """
    doc = Document()
    section = doc.sections[0]

    # Parse layout values early (before margin setup)
    page_w = float(layout.get("page_width_mm", 230.0))
    page_h = float(layout.get("page_height_mm", 170.0))
    safe_margin = float(layout.get("safe_margin_mm", 5.0))

    section.page_width = Mm(page_w)
    section.page_height = Mm(page_h)
    section.top_margin = Mm(safe_margin)
    section.bottom_margin = Mm(safe_margin)
    section.left_margin = Mm(safe_margin)
    section.right_margin = Mm(safe_margin)
    section.header_distance = Mm(0)

    # Parse layout
    rec_x        = float(layout.get("recipient_x_mm", layout.get("recipient_left_mm", 130.0)))
    baseline_bot = float(layout.get("first_line_baseline_from_bottom_mm", 32.0))
    
    # Support new parameter names: recipient_left_mm, recipient_top_mm
    # Convert recipient_top_mm (from top) to baseline_from_bottom
    recipient_top_mm = float(layout.get("recipient_top_mm", 0.0))
    if recipient_top_mm > 0:
        page_h_for_calc = float(layout.get("page_height_mm", 170.0))
        baseline_bot = page_h_for_calc - recipient_top_mm
    
    line_gap     = float(layout.get("line_gap_mm", 8.0))
    font_name    = str(layout.get("font_name", "Times New Roman") or "Times New Roman")
    phone_on     = bool(layout.get("phone_on_envelope", True))
    phone_mode   = str(layout.get("phone_render_mode", "separate_line") or "separate_line")

    baseline_y = page_h - baseline_bot    # 170 - 32 = 138mm
    font_size_pt = float(layout.get("font_size_pt", 13.0))
    font_size_mm = font_size_pt * 25.4 / 72.0   # 13pt = 4.6mm
    # Para top = baseline_y - font_center_offset
    # font_center = font_size_mm / 2
    para_top = baseline_y - font_size_mm / 2.0   # 138 - 2.3 = 135.7mm

    # [DEBUG] Printer profile & rotate 180 support
    printer_profile = str(layout.get("printer_profile", "generic") or "generic")
    rotate_180 = bool(layout.get("rotate_180", False))
    printer_offset_x = float(layout.get("printer_offset_x_mm", 0.0))
    printer_offset_y = float(layout.get("printer_offset_y_mm", 0.0))
    block_w = float(layout.get("recipient_block_width_mm", 95.0))
    block_h = float(layout.get("recipient_block_height_mm", 28.0))

    print(f"[ENVELOPE_PARA_LAYOUT] page={page_w}x{page_h} "
          f"left_indent={rec_x} space_before={para_top:.2f} "
          f"baseline_from_bottom={baseline_bot} line_gap={line_gap} "
          f"font_size={font_size_pt}pt({font_size_mm:.2f}mm) "
          f"phone_on_envelope={phone_on} phone_mode={phone_mode}")
    print(f"[ENVELOPE_PRINTER_PROFILE] profile={printer_profile} rotate_180={rotate_180} "
          f"offset_x={printer_offset_x} offset_y={printer_offset_y}")

    if not recipients:
        recipients = [{"recipient_unit": "", "recipient_address": "", "recipient_contact": "", "recipient_phone": ""}]

    for idx, recipient in enumerate(recipients):
        normalized = _normalize_envelope_recipient(recipient)
        unit  = _to_single_line(normalized["name"])
        addr  = _to_single_line(normalized["address"])
        phone = _to_single_line(normalized["phone"])

        lines = _build_vcpmc_lines(
            unit=unit, addr=addr, phone=phone,
            phone_on=phone_on, phone_mode=phone_mode,
        )

        if calibration:
            lines = [
                f"[CALIB] X={rec_x} TOP={para_top:.0f} GAP={line_gap}",
                f"[CALIB] BASELINE_FROM_BOTTOM={baseline_bot} PAGE_H={page_h}",
                f"[CALIB] NAME={unit or '-'}",
                f"[CALIB] ADDR={addr or '-'}",
                f"[CALIB] ROTATE={rotate_180} OFFSET_X={printer_offset_x} OFFSET_Y={printer_offset_y}",
            ]

        # Calculate final coordinates based on rotate_180 and printer offset
        if rotate_180:
            # Rotated position: mirror coordinates around page center
            # (x, y) -> (page_w - x, page_h - y) then apply margin/offset
            rotated_x = page_w - rec_x - safe_margin + printer_offset_x
            rotated_y = page_h - para_top - safe_margin + printer_offset_y
            # Clamp to prevent negative
            final_left_indent = max(0.0, rotated_x)
            final_space_before = max(0.0, rotated_y)
        else:
            # Normal: apply safe margin and offset directly
            final_left_indent = rec_x - safe_margin + printer_offset_x
            final_space_before = para_top - safe_margin + printer_offset_y
            # Clamp to prevent negative values
            final_left_indent = max(0.0, final_left_indent)
            final_space_before = max(0.0, final_space_before)

        # Page break BEFORE paragraph (not inside runs) — so page 2 starts fresh
        # Use pageBreakBefore in paragraph properties
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent          = Mm(final_left_indent)
        pf.space_before         = Mm(final_space_before)
        pf.space_after          = Pt(0)
        pf.line_spacing_rule    = WD_LINE_SPACING.EXACTLY
        pf.line_spacing         = Mm(line_gap)

        # Page break BEFORE paragraph (idx > 0)
        if idx > 0:
            pPr = p._p.get_or_add_pPr()
            pgBr = pPr.find(qn("w:pageBreakBefore"))
            if pgBr is None:
                pgBr = p._p.makeelement(qn("w:pageBreakBefore"), {})
                pPr.append(pgBr)
        # Build content: one run per line, LINE break at end of each except last
        if lines:
            for i, line_text in enumerate(lines):
                r = p.add_run(line_text)
                r.font.name = font_name
                r.font.size = Pt(font_size_pt)
                rpr = r._element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                rfonts.set(qn("w:ascii"),     font_name)
                rfonts.set(qn("w:hAnsi"),     font_name)
                rfonts.set(qn("w:eastAsia"),  font_name)
                rfonts.set(qn("w:cs"),        font_name)
                if i < len(lines) - 1:
                    r.add_break(WD_BREAK.LINE)
        else:
            r = p.add_run("-")
            r.font.name = font_name
            r.font.size = Pt(font_size_pt)
            rpr = r._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"),     font_name)
            rfonts.set(qn("w:hAnsi"),     font_name)
            rfonts.set(qn("w:eastAsia"),  font_name)
            rfonts.set(qn("w:cs"),        font_name)

    _trim_document_trailing_blanks(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    # Comprehensive final logging
    orientation = "portrait" if page_w < page_h else "landscape"
    print(f"[ENVELOPE_FEED_ORIENTATION_TEST]")
    print(f"  profile={layout.get('feed_profile', 'unknown')}")
    print(f"  page_width={page_w}")
    print(f"  page_height={page_h}")
    print(f"  orientation={orientation}")
    print(f"  safe_margin={safe_margin}")
    print(f"  physical_x={rec_x} physical_y={para_top:.2f}")
    print(f"  word_left_indent={final_left_indent:.2f} word_space_before={final_space_before:.2f}")
    print(f"  offset_x={printer_offset_x} offset_y={printer_offset_y}")
    print(f"  rotate_180={rotate_180}")
    print(f"  file={output}")
    print(f"[ENVELOPE_PARA] saved {output} ({len(recipients)} pages)")


def _build_vcpmc_lines(
    *, unit: str, addr: str, phone: str,
    phone_on: bool = True,
    phone_mode: str = "separate_line",
) -> list[str]:
    """Build recipient lines for VCPMC bia thu.

    Default: name + address + phone (separate line) if phone is available.
    """
    lines: list[str] = []
    if unit:
        lines.append(unit)
    if addr:
        lines.append(addr)
    if phone_on and phone:
        lines.append(f"DT: {phone}")
    return lines or ["-"]


# =============================================================================
# Brother HL-L2360D Test File Generation
# Creates Portrait DOCX files matching Brother tray orientation.
# =============================================================================

def _build_brother_test_docx(
    *, output: Path,
) -> Path:
    """
    Generate a test DOCX for Brother HL-L2360D printer.
    Uses Portrait 170x230 - khay nam ngang, giay dat doc.
    
    Content:
    - BROTHER PORTRAIT 170x230
    - Nha Hang Phuong Nam
    - Ap My An 1, Xa Thap Muoi, Tinh Dong Thap
    - NO PHONE (phone_on_envelope=False)
    """
    # Test content - NO phone
    test_unit = "Nha Hang Phuong Nam"
    test_addr = "Ap My An 1, Xa Thap Muoi, Tinh Dong Thap"
    marker = "BROTHER_PORTRAIT_170x230"

    # Recipients list
    recipients = [{
        "recipient_unit": test_unit,
        "recipient_address": test_addr,
        "recipient_contact": "",
        "recipient_phone": "",
    }]

    # Portrait 170x230 layout - vi tri "Kinh gui" o giua-duoi
    layout = {
        "page_width_mm": 170.0,
        "page_height_mm": 230.0,
        "recipient_box_left_mm": 50.0,
        "recipient_box_top_mm": 120.0,
        "recipient_box_width_mm": 110.0,
        "recipient_box_height_mm": 35.0,
        "line_spacing_mm": 8.0,
        "phone_on_envelope": False,
        "_brother_portrait_mode": True,
        "_profile_marker": marker,
    }

    # Build DOCX using shared builder
    _build_envelope_docx(recipients=recipients, output=output, layout=layout, calibration=False)

    # Inject profile marker as visible header
    _inject_brother_test_marker(output, marker)

    print(f"[BROTHER_TEST] Created Portrait 170x230: {output}")
    return output


def _build_brother_test_a4_docx(
    *, output: Path,
) -> Path:
    """
    Generate a test DOCX with A4 size to test Brother rotation issue.
    If text is NOT rotated on A4, then custom paper size triggers auto-rotation.
    
    Content:
    - BROTHER A4 TEST
    - Nha Hang Phuong Nam
    - Ap My An 1, Xa Thap Muoi, Tinh Dong Thap
    """
    test_unit = "Nha Hang Phuong Nam"
    test_addr = "Ap My An 1, Xa Thap Muoi, Tinh Dong Thap"
    marker = "BROTHER_A4_TEST"

    recipients = [{
        "recipient_unit": test_unit,
        "recipient_address": test_addr,
        "recipient_contact": "",
        "recipient_phone": "",
    }]

    # A4 size layout - vi tri "Kinh gui" o giua-duoi
    layout = {
        "page_width_mm": 210.0,
        "page_height_mm": 297.0,
        "recipient_box_left_mm": 50.0,
        "recipient_box_top_mm": 120.0,
        "recipient_box_width_mm": 110.0,
        "recipient_box_height_mm": 35.0,
        "line_spacing_mm": 8.0,
        "phone_on_envelope": False,
        "_brother_portrait_mode": True,
        "_profile_marker": marker,
    }

    _build_envelope_docx(recipients=recipients, output=output, layout=layout, calibration=False)
    _inject_brother_test_marker(output, marker)

    print(f"[BROTHER_TEST] Created A4 test: {output}")
    return output


def _build_brother_test_files(
    *, output_dir: Path,
) -> dict[str, Path]:
    """
    Generate test DOCX files for Brother HL-L2360D.
    - Portrait 170x230 (for printing on envelope)
    - A4 (for testing rotation issue)
    Returns dict: {filename: Path}
    """
    results = {}

    # File 1: Portrait 170x230 (for actual printing)
    filename1 = "TEST_BROTHER_170x230.docx"
    output1 = output_dir / filename1
    _build_brother_test_docx(output=output1)
    results[filename1] = output1

    # File 2: A4 (for testing rotation - if NOT rotated on A4, custom size triggers rotation)
    filename2 = "TEST_BROTHER_A4.docx"
    output2 = output_dir / filename2
    _build_brother_test_a4_docx(output=output2)
    results[filename2] = output2

    return results


def _inject_brother_test_marker(docx_path: Path, marker: str) -> None:
    """Add a visible test marker to the first paragraph of the DOCX."""
    try:
        doc = Document(str(docx_path))
        # Add header line to first paragraph
        if doc.paragraphs:
            first = doc.paragraphs[0]
            original_text = first.text
            # Prepend marker
            run = first.insert_run(0, f"[{marker}] ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue
            # Re-add original content
            if original_text:
                first.add_run(original_text)
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"[{marker}]")
            run.bold = True
        doc.save(str(docx_path))
    except Exception as e:
        print(f"[BROTHER_TEST] Failed to inject marker: {e}")


# =============================================================================
# Brother HL-L2360D printable test content (for UI preview)
# =============================================================================


def _generate_envelope_for_batch(
    *, db: Session, batch: BgCongVanBatchRow,
    force_regenerate: bool = False, calibration: bool = False,
) -> tuple[Path, int]:
    batch_id = int(getattr(batch, "id", 0) or 0)
    rows = (
        db.query(BgCongVanRow)
        .filter(BgCongVanRow.batch_id == batch_id)
        .order_by(BgCongVanRow.id.asc())
        .all()
    )
    if not rows:
        raise ValueError("Batch chưa có công văn để tạo file bìa thư.")
    issue_date = getattr(batch, "issue_date", None) or getattr(rows[0], "issue_date", None) or date.today()
    existing_attr = "envelope_calibration_docx_path" if calibration else "envelope_docx_path"
    existing = Path(str(getattr(batch, existing_attr, "") or "").strip()) if getattr(batch, existing_attr, None) else None
    if existing and existing.exists() and not force_regenerate:
        return existing, len(rows)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    suffix = "CALIB" if calibration else "GOP"
    out_name = f"BITHU_TAIKY_{suffix}_{issue_date.year}_{ts}.docx"
    out_path = (_CONGVAN_DOCX_ROOT / str(issue_date.year) / out_name).resolve()
    recipients = _build_envelope_recipients_from_rows(rows)
    # Load base envelope layout
    layout = _load_envelope_layout_config(db=db)
    # Apply printer profile if set
    profile_config = _load_envelope_profile_config(db=db)
    profile_id = profile_config.get("profile_id", "legacy_default")
    brother_mode = profile_config.get("brother_transform_mode")
    if profile_id != "legacy_default":
        layout = _apply_printer_profile_to_layout(
            base_layout=layout,
            profile_id=profile_id,
            brother_transform_mode=brother_mode,
        )
        print(f"[ENVELOPE_PROFILE] Using profile={profile_id} brother_mode={brother_mode}")
    # Use the OLD APP logic (bottom-based, no safe margin, no baseline approach)
    _build_envelope_docx(recipients=recipients, output=out_path, layout=layout, calibration=calibration)
    return out_path, len(rows)

def _resolve_recipient_contact_phone(contract: ContractRecordRow) -> tuple[str, str]:
    payload = _to_payload_dict(getattr(contract, "bg_payload_json", None))
    partner_info = payload.get("partner_info") if isinstance(payload.get("partner_info"), dict) else {}
    contact = _pick_first_non_empty(
        partner_info.get("contact_name"),
        partner_info.get("contact_person"),
        payload.get("contact_name"),
        payload.get("nguoi_lien_he"),
        payload.get("nguoi_dai_dien"),
        getattr(contract, "don_vi_nguoi_dai_dien", None),
    )
    phone = _pick_first_non_empty(
        partner_info.get("phone"),
        partner_info.get("contact_phone"),
        payload.get("phone"),
        payload.get("so_dien_thoai"),
        payload.get("dien_thoai"),
        getattr(contract, "don_vi_dien_thoai", None),
    )
    return contact, phone

def _next_cong_van_no(*, db: Session, issue_date: date) -> str:
    year = int(issue_date.year)
    prefix = f"CV-TK/{year}/"
    rows = (
        db.query(BgCongVanRow.cong_van_no)
        .filter(BgCongVanRow.issue_date.is_not(None))
        .filter(BgCongVanRow.issue_date >= date(year, 1, 1))
        .filter(BgCongVanRow.issue_date <= date(year, 12, 31))
        .all()
    )
    max_no = 0
    for row in rows:
        no = str(row[0] or "")
        if not no.startswith(prefix):
            continue
        tail = no.replace(prefix, "", 1).strip()
        if tail.isdigit():
            max_no = max(max_no, int(tail))
    return f"{prefix}{max_no + 1:04d}"

def _split_number_suffix(value: str) -> tuple[str, int, int] | None:
    m = re.match(r"^(.*?)(\d+)$", str(value or "").strip())
    if not m:
        return None
    return m.group(1), int(m.group(2)), len(m.group(2))

def _generate_cong_van_numbers(*, db: Session, issue_date: date, count: int, start_number: str | None) -> list[str]:
    n = max(0, int(count or 0))
    if n <= 0:
        return []
    seed = str(start_number or "").strip()
    if not seed:
        seed = _next_cong_van_no(db=db, issue_date=issue_date)
    parsed = _split_number_suffix(seed)
    if parsed is None:
        values = [seed]
        for _ in range(n - 1):
            values.append(_next_cong_van_no(db=db, issue_date=issue_date))
        return values
    prefix, start_no, width = parsed
    values = [f"{prefix}{(start_no + i):0{width}d}" for i in range(n)]
    existing = {
        str(row[0] or "").strip()
        for row in (
            db.query(BgCongVanRow.cong_van_no)
            .filter(BgCongVanRow.issue_date.is_not(None))
            .filter(BgCongVanRow.issue_date >= date(issue_date.year, 1, 1))
            .filter(BgCongVanRow.issue_date <= date(issue_date.year, 12, 31))
            .filter(BgCongVanRow.deleted_at.is_(None))
            .all()
        )
    }
    conflict = next((no for no in values if no in existing), "")
    if conflict:
        raise ValueError(f"Số công văn bị trùng: {conflict}")
    return values

def _ensure_cong_van_no_available(*, db: Session, issue_date: date, cong_van_no: str) -> None:
    no = str(cong_van_no or "").strip()
    if not no:
        raise ValueError("Thiếu số công văn.")
    exists = (
        db.query(BgCongVanRow.id)
        .filter(BgCongVanRow.issue_date.is_not(None))
        .filter(BgCongVanRow.issue_date >= date(issue_date.year, 1, 1))
        .filter(BgCongVanRow.issue_date <= date(issue_date.year, 12, 31))
        .filter(BgCongVanRow.cong_van_no == no)
        .filter(BgCongVanRow.deleted_at.is_(None))
        .first()
    )
    if exists:
        raise ValueError(f"Số công văn đã tồn tại trong năm {issue_date.year}: {no}")


# =============================================================================
# API Endpoints
# =============================================================================

@router.get("/envelope-layout-config")
def get_envelope_layout_config(
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    layout = _load_envelope_layout_config(db=db)
    return JSONResponse({"ok": True, "layout": layout})


@router.put("/envelope-layout-config")
def save_envelope_layout_config(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    user = _get_current_user(credentials=credentials, db=db)
    body = payload if isinstance(payload, dict) else {}
    raw_layout = body.get("layout") if isinstance(body.get("layout"), dict) else body
    layout = _save_envelope_layout_config(
        db=db,
        config=(raw_layout if isinstance(raw_layout, dict) else {}),
        actor=str(getattr(user, "username", "") or ""),
    )
    return JSONResponse({"ok": True, "layout": layout})


# =============================================================================
# Printer Profile API
# =============================================================================

@router.get("/envelope-profiles")
def list_envelope_profiles(
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """List all available printer profiles."""
    _get_current_user(credentials=credentials, db=db)
    profiles = [
        {
            "id": pid,
            "name": pdata.get("name", pid),
            "description": pdata.get("description", ""),
            "transform_mode": pdata.get("transform_mode", "none"),
        }
        for pid, pdata in PRINTER_PROFILES.items()
    ]
    # Add brother transform modes info
    brother_modes = [
        {"id": mid, "name": mdata.get("name", mid), "description": mdata.get("description", "")}
        for mid, mdata in BROTHER_TRANSFORM_MODES.items()
    ]
    return JSONResponse({
        "ok": True,
        "profiles": profiles,
        "brother_transform_modes": brother_modes,
    })


@router.get("/envelope-profile-config")
def get_envelope_profile_config(
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Get current printer profile configuration."""
    _get_current_user(credentials=credentials, db=db)
    config = _load_envelope_profile_config(db=db)
    profile_id = config.get("profile_id", "legacy_default")
    profile = _get_printer_profile(profile_id)
    return JSONResponse({
        "ok": True,
        "config": config,
        "profile": profile,
    })


@router.put("/envelope-profile-config")
def save_envelope_profile_config(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Save printer profile configuration."""
    user = _get_current_user(credentials=credentials, db=db)
    body = payload if isinstance(payload, dict) else {}
    profile_id = str(body.get("profile_id", "legacy_default"))
    brother_transform_mode = str(body.get("brother_transform_mode", "portrait"))
    config = {
        "profile_id": profile_id,
        "brother_transform_mode": brother_transform_mode,
    }
    saved = _save_envelope_profile_config(
        db=db,
        config=config,
        actor=str(getattr(user, "username", "") or ""),
    )
    profile = _get_printer_profile(profile_id)
    return JSONResponse({"ok": True, "config": saved, "profile": profile})


@router.post("/envelope-test-brother")
def create_brother_test_files(
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Generate test DOCX files for Brother HL-L2360D printer testing.
    - TEST_BROTHER_170x230.docx: For actual printing on envelopes
    - TEST_BROTHER_A4.docx: For testing rotation issue (print on A4)
    
    If text is NOT rotated on A4, then custom paper size triggers auto-rotation.
    """
    _get_current_user(credentials=credentials, db=db)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    test_dir = (_CONGVAN_DOCX_ROOT / "test_brother" / ts).resolve()
    test_dir.mkdir(parents=True, exist_ok=True)

    results = {}
    for filename, output_path in _build_brother_test_files(output_dir=test_dir).items():
        rel_path = output_path.relative_to(_STORAGE_ROOT) if output_path.is_relative_to(_STORAGE_ROOT) else output_path
        results[filename] = {
            "filename": filename,
            "download_url": f"/api/dispatches/download/{str(rel_path).replace(chr(92), '/')}",
        }

    return JSONResponse({
        "ok": True,
        "profile_id": "brother_hl_l2360d",
        "test_files": results,
        "message": "Da tao 2 file test. 170x230 de in bia, A4 de test xem driver co tu xoay khong.",
    })


@router.get("/expired-contracts")
def get_expired_contracts(
    filter_mode: str = Query(default="year"),
    filter_year: int | None = Query(default=None),
    filter_quarter: int | None = Query(default=None),
    filter_month: int | None = Query(default=None),
    filter_week: int | None = Query(default=None),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=50, ge=1, le=200),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    today = date.today()
    year_value = _parse_int_range(filter_year, today.year, 2000, 2100)
    quarter_value = _parse_int_range(filter_quarter, 1, 1, 4)
    month_value = _parse_int_range(filter_month, today.month, 1, 12)
    week_value = _parse_int_range(filter_week, 1, 1, 4)
    range_start, range_end = _resolve_expiry_date_range(
        filter_mode=filter_mode, filter_year=year_value,
        filter_quarter=quarter_value, filter_month=month_value, filter_week=week_value,
    )
    base_q = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.ngay_ket_thuc.is_not(None))
        .filter(ContractRecordRow.ngay_ket_thuc < today)
    )
    if range_start is not None and range_end is not None:
        base_q = base_q.filter(ContractRecordRow.ngay_ket_thuc >= range_start)
        base_q = base_q.filter(ContractRecordRow.ngay_ket_thuc <= range_end)

    # Count total before pagination
    total_count = base_q.count()

    # Filter karaoke only for total
    total_karaoke = 0
    for row in base_q.all():
        if _is_karaoke_contract_row(row):
            total_karaoke += 1

    # Apply pagination
    offset = (page - 1) * page_size
    rows = base_q.order_by(
        ContractRecordRow.ngay_ket_thuc.asc(), ContractRecordRow.contract_year.desc()
    ).offset(offset).limit(page_size).all()

    out = []
    for row in rows:
        if not _is_karaoke_contract_row(row):
            continue
        cong_van_count = (
            db.query(BgCongVanRow.id)
            .filter(BgCongVanRow.contract_id == int(getattr(row, "id", 0) or 0))
            .count()
        )
        # Get latest dispatch info for this contract
        latest_cv = (
            db.query(BgCongVanRow)
            .filter(BgCongVanRow.contract_id == int(getattr(row, "id", 0) or 0))
            .order_by(BgCongVanRow.attempt_no.desc())
            .first()
        )
        latest_cv_no = str(getattr(latest_cv, "cong_van_no", "") or "") if latest_cv else ""
        latest_cv_date = (
            getattr(latest_cv, "issue_date", None).strftime("%d/%m/%Y")
            if getattr(latest_cv, "issue_date", None) else ""
        ) if latest_cv else ""
        latest_cv_status = str(getattr(latest_cv, "status", "") or "") if latest_cv else ""

        contract_no = str(getattr(row, "contract_no", "") or "")
        out.append({
            "contract_id": int(getattr(row, "id", 0) or 0),
            "so_hop_dong": contract_no,
            "contract_no": contract_no,
            "contract_year": int(getattr(row, "contract_year", 0) or 0),
            "don_vi_ten": str(getattr(row, "don_vi_ten", "") or ""),
            "ngay_ky_hop_dong": _format_dd_mm_yyyy(getattr(row, "ngay_lap_hop_dong", None)),
            "ngay_het_hieu_luc_hd": _format_dd_mm_yyyy(getattr(row, "ngay_ket_thuc", None)),
            "days_expired": int((today - row.ngay_ket_thuc).days) if getattr(row, "ngay_ket_thuc", None) else 0,
            "cong_van_count": int(cong_van_count),
            "recipient_address": str(getattr(row, "don_vi_dia_chi", "") or ""),
            "ten_bang_hieu": str(getattr(row, "ten_bang_hieu", "") or ""),
            "latest_dispatch_no": latest_cv_no,
            "latest_dispatch_date": latest_cv_date,
            "latest_dispatch_status": latest_cv_status,
        })
    total_pages = (total_karaoke + page_size - 1) // page_size if total_karaoke > 0 else 0
    return JSONResponse({
        "ok": True,
        "rows": out,
        "total": total_karaoke,
        "page": page,
        "page_size": page_size,
        "total_pages": total_pages,
        "filters": {
            "filter_mode": str(filter_mode or "year").strip().lower(),
            "filter_year": year_value,
            "filter_quarter": quarter_value,
            "filter_month": month_value,
            "filter_week": week_value,
        },
    })


@router.get("")
def list_dispatches(
    status: str | None = None,
    year: int | None = None,
    limit: int = Query(default=100, ge=1, le=500),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    status_filter = str(status or "").strip().lower()
    selected_year = int(year) if year else None
    selected_limit = max(1, min(int(limit or 100), 500))
    q = db.query(BgCongVanRow).filter(BgCongVanRow.field_code == "karaoke")
    if status_filter and status_filter != "all":
        q = q.filter(BgCongVanRow.status == status_filter)
    q = q.filter(BgCongVanRow.issue_date.is_not(None))
    q = q.filter(BgCongVanRow.deleted_at.is_(None))
    if selected_year:
        q = q.filter(BgCongVanRow.issue_date >= date(selected_year, 1, 1))
        q = q.filter(BgCongVanRow.issue_date <= date(selected_year, 12, 31))
    rows = q.order_by(BgCongVanRow.issue_date.desc(), BgCongVanRow.id.desc()).limit(selected_limit).all()
    batch_ids = sorted({int(getattr(r, "batch_id", 0) or 0) for r in rows if int(getattr(r, "batch_id", 0) or 0) > 0})
    batches_map: dict[int, BgCongVanBatchRow] = {}
    if batch_ids:
        batch_rows = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id.in_(batch_ids)).all()
        batches_map = {int(getattr(b, "id", 0) or 0): b for b in batch_rows}
    row_ids = [int(getattr(r, "id", 0) or 0) for r in rows if int(getattr(r, "id", 0) or 0) > 0]
    logs_map: dict[int, list[BgCongVanProcessRow]] = {}
    if row_ids:
        logs = (
            db.query(BgCongVanProcessRow)
            .filter(BgCongVanProcessRow.cong_van_id.in_(row_ids))
            .order_by(BgCongVanProcessRow.created_at.desc(), BgCongVanProcessRow.id.desc())
            .all()
        )
        for log in logs:
            key = int(getattr(log, "cong_van_id", 0) or 0)
            logs_map.setdefault(key, []).append(log)
    out = []
    for r in rows:
        batch_id_val = int(getattr(r, "batch_id", 0) or 0)
        batch = batches_map.get(batch_id_val)
        batch_doc_url = _download_url_for_path(getattr(batch, "merged_docx_path", None), year=None) if batch_id_val > 0 and batch else ""
        batch_env_url = _download_url_for_path(getattr(batch, "envelope_docx_path", None), year=None) if batch_id_val > 0 and batch else ""
        batch_calib_url = _download_url_for_path(getattr(batch, "envelope_calibration_docx_path", None), year=None) if batch_id_val > 0 and batch else ""
        progress = _derive_progress_fields(
            status=str(getattr(r, "status", "") or "draft"),
            action_types=[str(getattr(log, "action_type", "") or "") for log in logs_map.get(int(getattr(r, "id", 0) or 0), [])],
        )
        out.append({
            "id": int(getattr(r, "id", 0) or 0),
            "batch_id": batch_id_val,
            "cong_van_no": str(getattr(r, "cong_van_no", "") or ""),
            "issue_date": _format_dd_mm_yyyy(getattr(r, "issue_date", None)),
            "contract_no": str(getattr(r, "contract_no", "") or ""),
            "recipient_unit": str(getattr(r, "recipient_unit", "") or ""),
            "recipient_address": str(getattr(r, "recipient_address", "") or ""),
            "recipient_contact": str(getattr(r, "recipient_contact", "") or ""),
            "recipient_phone": str(getattr(r, "recipient_phone", "") or ""),
            "status": str(getattr(r, "status", "") or "draft"),
            "docx_path": str(getattr(r, "docx_path", "") or ""),
            "download_url": _download_url_for_path(getattr(r, "docx_path", None), year=None) if getattr(r, "docx_path", None) else "",
            "batch_merged_download_url": batch_doc_url,
            "batch_envelope_download_url": batch_env_url,
            "batch_envelope_calibration_download_url": batch_calib_url,
            "batch_total_items": int(getattr(batch, "total_items", 0) or 0) if batch else 0,
            "batch_envelope_total_items": int(getattr(batch, "envelope_total_items", 0) or 0) if batch else 0,
            "batch_envelope_generated_at": getattr(batch, "envelope_generated_at", None).isoformat() if batch and getattr(batch, "envelope_generated_at", None) else "",
            "note": str(getattr(r, "note", "") or ""),
            **progress,
        })
    return JSONResponse({"ok": True, "rows": out})


# =============================================================================
# Batch-level endpoints
# =============================================================================

@router.get("/batches")
def list_batches(
    year: int | None = Query(default=None),
    dispatch_type: str | None = Query(default=None),
    cong_van_no: str | None = Query(default=None),
    include_deleted: bool = Query(default=False),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    selected_year = int(year) if year else None
    q = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.cong_van_no.is_not(None))
    if not include_deleted:
        q = q.filter(BgCongVanBatchRow.deleted_at.is_(None))
    if selected_year:
        q = q.filter(BgCongVanBatchRow.issue_date >= date(selected_year, 1, 1))
        q = q.filter(BgCongVanBatchRow.issue_date <= date(selected_year, 12, 31))
    if dispatch_type:
        q = q.filter(BgCongVanBatchRow.dispatch_type == dispatch_type)
    if cong_van_no:
        q = q.filter(BgCongVanBatchRow.cong_van_no.ilike(f"%{cong_van_no}%"))
    total_count = q.count()
    offset = (page - 1) * page_size
    rows = q.order_by(BgCongVanBatchRow.issue_date.desc(), BgCongVanBatchRow.id.desc()).offset(offset).limit(page_size).all()
    out = []
    for r in rows:
        batch_id_val = int(getattr(r, "id", 0) or 0)
        merged_path = getattr(r, "merged_docx_path", None)
        envelope_path = getattr(r, "envelope_docx_path", None)
        out.append({
            "id": batch_id_val,
            "cong_van_no": str(getattr(r, "cong_van_no", "") or ""),
            "issue_date": _format_dd_mm_yyyy(getattr(r, "issue_date", None)),
            "dispatch_type": str(getattr(r, "dispatch_type", "") or ""),
            "template_name": str(getattr(r, "template_name", "") or ""),
            "total_items": int(getattr(r, "total_items", 0) or 0) or int(getattr(r, "ready_items", 0) or 0),
            "ready_items": int(getattr(r, "ready_items", 0) or 0),
            "missing_items": int(getattr(r, "missing_items", 0) or 0),
            "create_envelope": bool(getattr(r, "create_envelope", 0)),
            "merge_output": bool(getattr(r, "merge_output", 1)),
            "note": str(getattr(r, "note", "") or ""),
            "created_by": str(getattr(r, "created_by", "") or ""),
            "created_at": getattr(r, "created_at", None).isoformat() if getattr(r, "created_at", None) else "",
            "merged_download_url": _download_url_for_path(merged_path, year=None) if merged_path else "",
            "envelope_download_url": _download_url_for_path(envelope_path, year=None) if envelope_path else "",
            "envelope_generated_at": getattr(r, "envelope_generated_at", None).isoformat() if getattr(r, "envelope_generated_at", None) else "",
            "envelope_total_items": int(getattr(r, "envelope_total_items", 0) or 0),
        })
    total_pages = (total_count + page_size - 1) // page_size if total_count > 0 else 0
    return JSONResponse({
        "ok": True,
        "rows": out,
        "total": total_count,
        "page": page,
        "page_size": page_size,
        "total_pages": total_pages,
    })


@router.get("/batches/{batch_id}")
def get_batch_detail(
    batch_id: int,
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == int(batch_id)).first()
    if batch is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy đợt công văn"}, status_code=404)
    items = (
        db.query(BgCongVanRow)
        .filter(BgCongVanRow.batch_id == int(batch_id))
        .filter(BgCongVanRow.deleted_at.is_(None))
        .order_by(BgCongVanRow.id.asc())
        .all()
    )
    item_rows = []
    for r in items:
        item_id = int(getattr(r, "id", 0) or 0)
        logs = (
            db.query(BgCongVanProcessRow)
            .filter(BgCongVanProcessRow.cong_van_id == item_id)
            .order_by(BgCongVanProcessRow.created_at.desc())
            .limit(3)
            .all()
        )
        progress = _derive_progress_fields(
            status=str(getattr(r, "status", "") or "draft"),
            action_types=[str(getattr(log, "action_type", "") or "") for log in logs],
        )
        item_rows.append({
            "id": item_id,
            "cong_van_no": str(getattr(r, "cong_van_no", "") or ""),
            "issue_date": _format_dd_mm_yyyy(getattr(r, "issue_date", None)),
            "contract_no": str(getattr(r, "contract_no", "") or ""),
            "recipient_unit": str(getattr(r, "recipient_unit", "") or ""),
            "recipient_address": str(getattr(r, "recipient_address", "") or ""),
            "so_dien_thoai": str(getattr(r, "recipient_phone", "") or ""),
            "nguoi_nhan_bia_thu": str(getattr(r, "dong_nguoi_nhan_bia_thu", "") or ""),
            "dong_nguoi_nhan_bia_thu": str(getattr(r, "dong_nguoi_nhan_bia_thu", "") or ""),
            "attempt_no": int(getattr(r, "attempt_no", 0) or 0),
            "lan_gui": int(getattr(r, "lan_gui", 1) or 1),
            "dispatch_type": str(getattr(r, "dispatch_type", "") or ""),
            "status": str(getattr(r, "status", "") or "draft"),
            "trang_thai_lien_he": str(getattr(r, "trang_thai_lien_he", "CHUA_LIEN_HE") or "CHUA_LIEN_HE"),
            "trang_thai_hop_dong": str(getattr(r, "trang_thai_hop_dong", "CHUA_KY_HOP_DONG") or "CHUA_KY_HOP_DONG"),
            "ngay_lien_he_gan_nhat": (
                getattr(r, "ngay_lien_he_gan_nhat", None).isoformat()
                if getattr(r, "ngay_lien_he_gan_nhat", None) else None
            ),
            "ghi_chu_lien_he": str(getattr(r, "ghi_chu_lien_he", "") or ""),
            "ngay_ky_hop_dong": (
                getattr(r, "ngay_ky_hop_dong", None).isoformat()
                if getattr(r, "ngay_ky_hop_dong", None) else None
            ),
            "contract_id": int(getattr(r, "contract_id") or 0) or None,
            "docx_path": str(getattr(r, "docx_path", "") or ""),
            "download_url": _download_url_for_path(getattr(r, "docx_path", None), year=None) if getattr(r, "docx_path", None) else "",
            "envelope_download_url": "",  # envelope is batch-level in karaoke; per-row envelope not stored
            "created_at": getattr(r, "created_at", None).isoformat() if getattr(r, "created_at", None) else "",
            **progress,
        })
    return JSONResponse({
        "ok": True,
        "batch": {
            "id": int(getattr(batch, "id", 0) or 0),
            "cong_van_no": str(getattr(batch, "cong_van_no", "") or ""),
            "issue_date": _format_dd_mm_yyyy(getattr(batch, "issue_date", None)),
            "dispatch_type": str(getattr(batch, "dispatch_type", "") or ""),
            "template_name": str(getattr(batch, "template_name", "") or ""),
            "total_items": len(item_rows),
            "ready_items": int(getattr(batch, "ready_items", 0) or 0),
            "missing_items": int(getattr(batch, "missing_items", 0) or 0),
            "create_envelope": bool(getattr(batch, "create_envelope", False)),
            "merge_output": bool(getattr(batch, "merge_output", True)),
            "envelope_recipient_mode": str(getattr(batch, "envelope_recipient_mode", "") or ""),
            "envelope_custom_prefix": str(getattr(batch, "envelope_custom_prefix", "") or ""),
            "note": str(getattr(batch, "note", "") or ""),
            "created_by": str(getattr(batch, "created_by", "") or ""),
            "created_at": getattr(batch, "created_at", None).isoformat() if getattr(batch, "created_at", None) else "",
            "merged_download_url": _download_url_for_path(getattr(batch, "merged_docx_path", None), year=None),
            "envelope_download_url": _download_url_for_path(getattr(batch, "envelope_docx_path", None), year=None),
        },
        "items": item_rows,
    })


@router.patch("/items/{item_id}/tracking")
def update_item_tracking(
    item_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Update tracking fields for a dispatch item.
    Does NOT modify generated document contents.
    Handles: trang_thai_lien_he, trang_thai_hop_dong, contract linking, notes.
    """
    user = _get_current_user(credentials=credentials, db=db)
    item = db.query(BgCongVanRow).filter(BgCongVanRow.id == int(item_id)).first()
    if item is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy công văn"}, status_code=404)

    logs: list[str] = []
    actor = str(getattr(user, "username", "") or "")

    new_contact_status = str(payload.get("trang_thai_lien_he") or "").strip()
    if new_contact_status:
        valid_contact = {
            "CHUA_LIEN_HE", "DA_LIEN_HE", "DA_GUI_CONG_VAN",
            "DA_PHAN_HOI", "DANG_THUONG_LUONG",
            "NGUNG_HOAT_DONG", "KHONG_HOP_TAC", "SAI_THONG_TIN",
        }
        if new_contact_status not in valid_contact:
            return JSONResponse(
                {"ok": False, "error": f"Trạng thái liên hệ không hợp lệ: {new_contact_status}"},
                status_code=400,
            )
        old_contact = str(getattr(item, "trang_thai_lien_he", "") or "")
        if old_contact != new_contact_status:
            item.trang_thai_lien_he = new_contact_status
            logs.append(f"Liên hệ: {old_contact} → {new_contact_status}")

    new_contract_status = str(payload.get("trang_thai_hop_dong") or "").strip()
    if new_contract_status:
        valid_contract = {
            "CHUA_KY_HOP_DONG", "DANG_XU_LY_HOP_DONG",
            "DA_KY_HOP_DONG", "TU_CHOI_KY", "KHONG_DU_DIEU_KIEN",
        }
        if new_contract_status not in valid_contract:
            return JSONResponse(
                {"ok": False, "error": f"Trạng thái hợp đồng không hợp lệ: {new_contract_status}"},
                status_code=400,
            )
        old_contract = str(getattr(item, "trang_thai_hop_dong", "") or "")
        if old_contract != new_contract_status:
            item.trang_thai_hop_dong = new_contract_status
            logs.append(f"Hợp đồng: {old_contract} → {new_contract_status}")

    new_ngay_lien_he = payload.get("ngay_lien_he_gan_nhat")
    if new_ngay_lien_he:
        try:
            parsed_date = datetime.strptime(str(new_ngay_lien_he), "%Y-%m-%d")
            item.ngay_lien_he_gan_nhat = parsed_date
            logs.append(f"Cập nhật ngày liên hệ gần nhất: {new_ngay_lien_he}")
        except ValueError:
            try:
                parsed_date = datetime.strptime(str(new_ngay_lien_he), "%d/%m/%Y")
                item.ngay_lien_he_gan_nhat = parsed_date
                logs.append(f"Cập nhật ngày liên hệ gần nhất: {new_ngay_lien_he}")
            except ValueError:
                return JSONResponse({"ok": False, "error": f"Ngày liên hệ không hợp lệ: {new_ngay_lien_he}"}, status_code=400)

    new_note = str(payload.get("ghi_chu_lien_he") or "").strip()
    if "ghi_chu_lien_he" in payload:
        old_note = str(getattr(item, "ghi_chu_lien_he", "") or "")
        if old_note != new_note:
            item.ghi_chu_lien_he = new_note or None
            if new_note:
                logs.append(f"Ghi chú: {new_note[:80]}")

    contract_id = payload.get("contract_id")
    if contract_id is not None:
        item.contract_id = int(contract_id) if contract_id else None
        logs.append(f"Gắn contract_id: {contract_id}")

    so_hop_dong = payload.get("so_hop_dong")
    if so_hop_dong is not None:
        item.contract_no = str(so_hop_dong).strip() or None
        logs.append(f"Số HĐ: {so_hop_dong}")

    ngay_ky_hd = payload.get("ngay_ky_hop_dong")
    if ngay_ky_hd is not None:
        raw = str(ngay_ky_hd).strip()
        if not raw:
            item.ngay_ky_hop_dong = None
            logs.append("Xóa ngày ký HĐ")
        else:
            try:
                item.ngay_ky_hop_dong = datetime.strptime(raw, "%Y-%m-%d").date()
                logs.append(f"Ngày ký HĐ: {raw}")
            except ValueError:
                try:
                    item.ngay_ky_hop_dong = datetime.strptime(raw, "%d/%m/%Y").date()
                    logs.append(f"Ngày ký HĐ: {raw}")
                except ValueError:
                    return JSONResponse({"ok": False, "error": f"Ngày ký HĐ không hợp lệ: {raw}"}, status_code=400)

    item.updated_by = actor

    for note in logs:
        db.add(BgCongVanProcessRow(
            cong_van_id=int(item_id),
            action_type="tracking_updated",
            note=note,
            actor=actor,
        ))

    db.commit()
    return JSONResponse({
        "ok": True,
        "id": int(item_id),
        "trang_thai_lien_he": str(getattr(item, "trang_thai_lien_he", "") or ""),
        "trang_thai_hop_dong": str(getattr(item, "trang_thai_hop_dong", "") or ""),
        "ngay_lien_he_gan_nhat": (
            getattr(item, "ngay_lien_he_gan_nhat", None).isoformat()
            if getattr(item, "ngay_lien_he_gan_nhat", None) else None
        ),
        "ghi_chu_lien_he": str(getattr(item, "ghi_chu_lien_he", "") or ""),
        "contract_id": int(getattr(item, "contract_id") or 0) or None,
        "contract_no": str(getattr(item, "contract_no", "") or ""),
        "ngay_ky_hop_dong": (
            getattr(item, "ngay_ky_hop_dong", None).isoformat()
            if getattr(item, "ngay_ky_hop_dong", None) else None
        ),
    })


# =============================================================================
# Batch edit / delete / restore
# =============================================================================

@router.patch("/batches/{batch_id}")
def update_batch(
    batch_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Update batch metadata.
    Does NOT regenerate any DOCX files.
    Returns a warning message that existing files are unchanged.
    """
    user = _get_current_user(credentials=credentials, db=db)
    batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == int(batch_id)).first()
    if batch is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy đợt công văn"}, status_code=404)
    if getattr(batch, "deleted_at", None) is not None:
        return JSONResponse({"ok": False, "error": "Đợt công văn đã bị xóa"}, status_code=410)

    actor = str(getattr(user, "username", "") or "")

    # Fields allowed to update
    new_cong_van_no = str(payload.get("cong_van_no") or "").strip() or None
    new_issue_date = payload.get("issue_date")
    new_note = payload.get("note")
    new_envelope_recipient_mode = str(payload.get("envelope_recipient_mode") or "").strip() or None
    new_envelope_custom_prefix = str(payload.get("envelope_custom_prefix") or "").strip() or None

    changes: list[str] = []

    if new_cong_van_no is not None:
        old = str(getattr(batch, "cong_van_no", "") or "")
        if old != new_cong_van_no:
            batch.cong_van_no = new_cong_van_no
            changes.append(f"Số công văn: {old} → {new_cong_van_no}")

    if new_issue_date is not None:
        raw = str(new_issue_date).strip()
        if raw:
            try:
                parsed = date.fromisoformat(raw)
                old_date = getattr(batch, "issue_date", None)
                batch.issue_date = parsed
                changes.append(f"Ngày ký: {old_date} → {parsed}")
            except ValueError:
                return JSONResponse({"ok": False, "error": f"Ngày ký không hợp lệ: {raw}"}, status_code=400)

    if "note" in payload:
        old_note = str(getattr(batch, "note", "") or "")
        new_note_val = str(new_note or "").strip() or None
        if old_note != (new_note_val or ""):
            batch.note = new_note_val
            changes.append("Ghi chú")

    if new_envelope_recipient_mode is not None:
        old_mode = str(getattr(batch, "envelope_recipient_mode", "") or "")
        if old_mode != new_envelope_recipient_mode:
            batch.envelope_recipient_mode = new_envelope_recipient_mode
            changes.append(f"Cách ghi người nhận bìa thư: {old_mode} → {new_envelope_recipient_mode}")

    if "envelope_custom_prefix" in payload:
        old_prefix = str(getattr(batch, "envelope_custom_prefix", "") or "")
        new_prefix_val = str(new_envelope_custom_prefix or "").strip() or None
        if old_prefix != (new_prefix_val or ""):
            batch.envelope_custom_prefix = new_prefix_val
            changes.append("Tiền tố người nhận bìa thư")

    batch.updated_by = actor
    db.commit()

    return JSONResponse({
        "ok": True,
        "id": int(batch_id),
        "warning": "Thông tin theo dõi đã được cập nhật. File đã tạo trước đó không tự thay đổi.",
        "changes": changes,
        "cong_van_no": str(getattr(batch, "cong_van_no", "") or ""),
        "issue_date": _format_dd_mm_yyyy(getattr(batch, "issue_date", None)),
        "note": str(getattr(batch, "note", "") or ""),
        "envelope_recipient_mode": str(getattr(batch, "envelope_recipient_mode", "") or ""),
        "envelope_custom_prefix": str(getattr(batch, "envelope_custom_prefix", "") or ""),
    })


@router.delete("/batches/{batch_id}")
def delete_batch(
    batch_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Hard delete a batch and all its items (Công văn uses physical delete only).
    """
    user = _get_current_user(credentials=credentials, db=db)
    batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == int(batch_id)).first()
    if batch is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy đợt công văn"}, status_code=404)

    reason = str(payload.get("delete_reason") or "").strip() or None

    # Hard delete all process logs tied to items in this batch first
    item_ids = [
        int(r.id)
        for r in db.query(BgCongVanRow.id).filter(BgCongVanRow.batch_id == int(batch_id)).all()
    ]
    if item_ids:
        db.query(BgCongVanProcessRow).filter(
            BgCongVanProcessRow.cong_van_id.in_(item_ids)
        ).delete(synchronize_session=False)

    # Hard delete all items in this batch
    item_count = db.query(BgCongVanRow).filter(
        BgCongVanRow.batch_id == int(batch_id)
    ).delete(synchronize_session=False)

    # Capture batch file paths before delete
    batch_doc_paths = [
        str(getattr(batch, "merged_docx_path", "") or "").strip(),
        str(getattr(batch, "envelope_docx_path", "") or "").strip(),
        str(getattr(batch, "envelope_calibration_docx_path", "") or "").strip(),
    ]

    # Hard delete the batch
    db.delete(batch)
    db.commit()

    if batch_doc_paths:
        for raw in batch_doc_paths:
            if not raw:
                continue
            try:
                p = Path(raw)
                if p.exists() and p.is_file():
                    p.unlink(missing_ok=True)
            except Exception:
                pass

    return JSONResponse({
        "ok": True,
        "id": int(batch_id),
        "deleted_item_count": int(item_count or 0),
        "message": f"Đợt công văn và {int(item_count or 0)} công văn đã được xóa vĩnh viễn.",
    })


@router.post("/batches/{batch_id}/restore")
def restore_batch(
    batch_id: int,
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Restore a soft-deleted batch and all its items."""
    user = _get_current_user(credentials=credentials, db=db)
    batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == int(batch_id)).first()
    if batch is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy đợt công văn"}, status_code=404)
    if getattr(batch, "deleted_at", None) is None:
        return JSONResponse({"ok": False, "error": "Đợt công văn chưa bị xóa"}, status_code=400)

    actor = str(getattr(user, "username", "") or "")

    # Restore all items belonging to this batch
    item_count = db.query(BgCongVanRow).filter(
        BgCongVanRow.batch_id == int(batch_id),
        BgCongVanRow.deleted_at.isnot(None),
    ).update({
        "deleted_at": None,
        "deleted_by": None,
        "delete_reason": None,
    }, synchronize_session=False)

    batch.deleted_at = None
    batch.deleted_by = None
    batch.delete_reason = None
    batch.updated_by = actor
    db.commit()

    return JSONResponse({
        "ok": True,
        "id": int(batch_id),
        "restored_item_count": item_count,
        "message": f"Đợt công văn và {item_count} công văn đã được khôi phục.",
    })


# =============================================================================
# Item edit / delete / restore
# =============================================================================

@router.patch("/items/{item_id}")
def update_item(
    item_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Full edit for a dispatch item.
    Does NOT modify generated DOCX files.
    Returns a warning that existing files are unchanged.
    """
    user = _get_current_user(credentials=credentials, db=db)
    item = db.query(BgCongVanRow).filter(BgCongVanRow.id == int(item_id)).first()
    if item is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy công văn"}, status_code=404)
    if getattr(item, "deleted_at", None) is not None:
        return JSONResponse({"ok": False, "error": "Công văn đã bị xóa"}, status_code=410)

    actor = str(getattr(user, "username", "") or "")
    changes: list[str] = []

    # recipient_unit
    if "recipient_unit" in payload:
        old = str(getattr(item, "recipient_unit", "") or "")
        new = str(payload["recipient_unit"] or "").strip()
        if old != new:
            item.recipient_unit = new
            changes.append("Tên đơn vị")

    # recipient_address
    if "recipient_address" in payload:
        old = str(getattr(item, "recipient_address", "") or "")
        new = str(payload["recipient_address"] or "").strip()
        if old != new:
            item.recipient_address = new
            changes.append("Địa chỉ")

    # recipient_phone / so_dien_thoai
    if "recipient_phone" in payload or "so_dien_thoai" in payload:
        old = str(getattr(item, "recipient_phone", "") or "")
        new = str(payload.get("recipient_phone") or payload.get("so_dien_thoai") or "").strip()
        if old != new:
            item.recipient_phone = new
            changes.append("Số điện thoại")

    # recipient_contact / nguoi_nhan_bia_thu
    if "recipient_contact" in payload or "nguoi_nhan_bia_thu" in payload:
        old = str(getattr(item, "recipient_contact", "") or "")
        new = str(payload.get("recipient_contact") or payload.get("nguoi_nhan_bia_thu") or "").strip()
        if old != new:
            item.recipient_contact = new
            changes.append("Người nhận bìa thư")

    # trang_thai_lien_he
    if "trang_thai_lien_he" in payload:
        new_status = str(payload["trang_thai_lien_he"] or "").strip()
        valid_contact = {
            "CHUA_LIEN_HE", "DA_LIEN_HE", "DA_GUI_CONG_VAN",
            "DA_PHAN_HOI", "DANG_THUONG_LUONG",
            "NGUNG_HOAT_DONG", "KHONG_HOP_TAC", "SAI_THONG_TIN",
        }
        if new_status and new_status not in valid_contact:
            return JSONResponse(
                {"ok": False, "error": f"Trạng thái liên hệ không hợp lệ: {new_status}"},
                status_code=400,
            )
        old = str(getattr(item, "trang_thai_lien_he", "") or "")
        if old != new_status:
            item.trang_thai_lien_he = new_status or "CHUA_LIEN_HE"
            changes.append(f"Trạng thái liên hệ: {old} → {new_status}")

    # ngay_lien_he_gan_nhat
    if "ngay_lien_he_gan_nhat" in payload:
        raw = str(payload["ngay_lien_he_gan_nhat"] or "").strip()
        if raw:
            try:
                item.ngay_lien_he_gan_nhat = datetime.strptime(raw, "%Y-%m-%d")
                changes.append("Ngày liên hệ gần nhất")
            except ValueError:
                try:
                    item.ngay_lien_he_gan_nhat = datetime.strptime(raw, "%d/%m/%Y")
                    changes.append("Ngày liên hệ gần nhất")
                except ValueError:
                    return JSONResponse({"ok": False, "error": f"Ngày liên hệ không hợp lệ: {raw}"}, status_code=400)
        else:
            item.ngay_lien_he_gan_nhat = None
            changes.append("Ngày liên hệ gần nhất (xóa)")

    # ghi_chu_lien_he
    if "ghi_chu_lien_he" in payload:
        old = str(getattr(item, "ghi_chu_lien_he", "") or "")
        new = str(payload["ghi_chu_lien_he"] or "").strip()
        if old != new:
            item.ghi_chu_lien_he = new or None
            changes.append("Ghi chú liên hệ")

    # trang_thai_hop_dong
    if "trang_thai_hop_dong" in payload:
        new_status = str(payload["trang_thai_hop_dong"] or "").strip()
        valid_contract = {
            "CHUA_KY_HOP_DONG", "DANG_XU_LY_HOP_DONG",
            "DA_KY_HOP_DONG", "TU_CHOI_KY", "KHONG_DU_DIEU_KIEN",
        }
        if new_status and new_status not in valid_contract:
            return JSONResponse(
                {"ok": False, "error": f"Trạng thái hợp đồng không hợp lệ: {new_status}"},
                status_code=400,
            )
        old = str(getattr(item, "trang_thai_hop_dong", "") or "")
        if old != new_status:
            item.trang_thai_hop_dong = new_status or "CHUA_KY_HOP_DONG"
            changes.append(f"Trạng thái hợp đồng: {old} → {new_status}")

    # contract_no / so_hop_dong
    if "contract_no" in payload or "so_hop_dong" in payload:
        old = str(getattr(item, "contract_no", "") or "")
        new = str(payload.get("contract_no") or payload.get("so_hop_dong") or "").strip()
        if old != new:
            item.contract_no = new or None
            changes.append("Số hợp đồng")

    # ngay_ky_hop_dong
    if "ngay_ky_hop_dong" in payload:
        raw = str(payload["ngay_ky_hop_dong"] or "").strip()
        if raw:
            try:
                item.ngay_ky_hop_dong = datetime.strptime(raw, "%Y-%m-%d").date()
                changes.append("Ngày ký hợp đồng")
            except ValueError:
                try:
                    item.ngay_ky_hop_dong = datetime.strptime(raw, "%d/%m/%Y").date()
                    changes.append("Ngày ký hợp đồng")
                except ValueError:
                    return JSONResponse({"ok": False, "error": f"Ngày ký hợp đồng không hợp lệ: {raw}"}, status_code=400)
        else:
            item.ngay_ky_hop_dong = None
            changes.append("Ngày ký hợp đồng (xóa)")

    item.updated_by = actor

    if changes:
        db.add(BgCongVanProcessRow(
            cong_van_id=int(item_id),
            action_type="item_edited",
            note="; ".join(changes),
            actor=actor,
        ))

    db.commit()

    return JSONResponse({
        "ok": True,
        "id": int(item_id),
        "warning": "Thông tin theo dõi đã được cập nhật. File đã tạo trước đó không tự thay đổi.",
        "changes": changes,
        "recipient_unit": str(getattr(item, "recipient_unit", "") or ""),
        "recipient_address": str(getattr(item, "recipient_address", "") or ""),
        "recipient_phone": str(getattr(item, "recipient_phone", "") or ""),
        "recipient_contact": str(getattr(item, "recipient_contact", "") or ""),
        "trang_thai_lien_he": str(getattr(item, "trang_thai_lien_he", "") or ""),
        "ngay_lien_he_gan_nhat": (
            getattr(item, "ngay_lien_he_gan_nhat", None).isoformat()
            if getattr(item, "ngay_lien_he_gan_nhat", None) else None
        ),
        "ghi_chu_lien_he": str(getattr(item, "ghi_chu_lien_he", "") or ""),
        "trang_thai_hop_dong": str(getattr(item, "trang_thai_hop_dong", "") or ""),
        "contract_no": str(getattr(item, "contract_no", "") or ""),
        "ngay_ky_hop_dong": (
            getattr(item, "ngay_ky_hop_dong", None).isoformat()
            if getattr(item, "ngay_ky_hop_dong", None) else None
        ),
    })


@router.delete("/items/{item_id}")
def delete_item(
    item_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Hard delete a single dispatch item (Công văn uses physical delete only).
    Also deletes its process logs and the DOCX file when present.
    """
    user = _get_current_user(credentials=credentials, db=db)
    item = db.query(BgCongVanRow).filter(BgCongVanRow.id == int(item_id)).first()
    if item is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy công văn"}, status_code=404)

    docx_path = str(getattr(item, "docx_path", "") or "").strip()
    row_batch_id = int(getattr(item, "batch_id", 0) or 0)

    db.query(BgCongVanProcessRow).filter(
        BgCongVanProcessRow.cong_van_id == int(item_id)
    ).delete(synchronize_session=False)
    db.delete(item)
    db.flush()

    if row_batch_id > 0:
        remains = db.query(BgCongVanRow.id).filter(BgCongVanRow.batch_id == row_batch_id).first()
        if remains is None:
            batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == row_batch_id).first()
            if batch is not None:
                db.delete(batch)

    db.commit()

    if docx_path:
        try:
            p = Path(docx_path)
            if p.exists() and p.is_file():
                p.unlink(missing_ok=True)
        except Exception:
            pass

    return JSONResponse({
        "ok": True,
        "id": int(item_id),
        "message": "Công văn đã được xóa vĩnh viễn.",
    })


@router.post("/items/{item_id}/restore")
def restore_item(
    item_id: int,
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Restore a soft-deleted dispatch item."""
    user = _get_current_user(credentials=credentials, db=db)
    item = db.query(BgCongVanRow).filter(BgCongVanRow.id == int(item_id)).first()
    if item is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy công văn"}, status_code=404)
    if getattr(item, "deleted_at", None) is None:
        return JSONResponse({"ok": False, "error": "Công văn chưa bị xóa"}, status_code=400)

    actor = str(getattr(user, "username", "") or "")

    item.deleted_at = None
    item.deleted_by = None
    item.delete_reason = None
    item.updated_by = actor
    db.commit()

    return JSONResponse({
        "ok": True,
        "id": int(item_id),
        "message": "Công văn đã được khôi phục.",
    })


@router.post("/batches/bulk-delete")
def bulk_delete_batches(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Bulk soft delete batches.
    Two modes:
      A) explicit IDs:  { "ids": [78, 79] }
      B) all_filtered:  { "scope": "all_filtered", "filters": {...}, "confirm_text": "XOA TOAN BO" }

    Soft deletes matching non-deleted batches and all their non-deleted items.
    Does NOT delete any files.
    """
    user = _get_current_user(credentials=credentials, db=db)
    actor = str(getattr(user, "username", "") or "")
    now = datetime.utcnow()
    reason = str(payload.get("delete_reason") or "").strip() or None
    scope = str(payload.get("scope") or "").strip().lower()

    batch_ids_to_delete: list[int] = []

    # --- Mode A: explicit IDs ---
    if scope == "ids" or ("ids" in payload and not scope):
        ids_raw = payload.get("ids", [])
        if not isinstance(ids_raw, list) or len(ids_raw) == 0:
            return JSONResponse(
                {"ok": False, "error": "ids must be a non-empty array"},
                status_code=400,
            )
        try:
            batch_ids_to_delete = [int(x) for x in ids_raw]
        except (ValueError, TypeError):
            return JSONResponse({"ok": False, "error": "ids must contain integers"}, status_code=400)

    # --- Mode B: all_filtered ---
    elif scope == "all_filtered":
        confirm_text = str(payload.get("confirm_text") or "").strip().upper()
        if confirm_text != "XOA TOAN BO":
            return JSONResponse(
                {"ok": False, "error": "confirm_text must be exactly 'XOA TOAN BO'"},
                status_code=400,
            )
        filters = payload.get("filters", {})
        if not isinstance(filters, dict):
            filters = {}

        q = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.deleted_at.is_(None))

        f_year = filters.get("year")
        if f_year:
            try:
                yr = int(f_year)
                q = q.filter(BgCongVanBatchRow.issue_date >= date(yr, 1, 1))
                q = q.filter(BgCongVanBatchRow.issue_date <= date(yr, 12, 31))
            except (ValueError, TypeError):
                pass

        f_type = filters.get("dispatch_type")
        if f_type and str(f_type).strip():
            q = q.filter(BgCongVanBatchRow.dispatch_type == str(f_type).strip())

        f_congvan_no = filters.get("cong_van_no")
        if f_congvan_no and str(f_congvan_no).strip():
            q = q.filter(BgCongVanBatchRow.cong_van_no.ilike(f"%{str(f_congvan_no).strip()}%"))

        # Never allow all_filtered with no constraints — too dangerous
        if not (f_year or (f_type and str(f_type).strip())):
            return JSONResponse(
                {"ok": False, "error": "all_filtered requires at least year or dispatch_type filter"},
                status_code=400,
            )

        batch_ids_to_delete = [int(getattr(r, "id", 0) or 0) for r in q.all()]
        batch_ids_to_delete = [bid for bid in batch_ids_to_delete if bid > 0]

    else:
        return JSONResponse(
            {"ok": False, "error": "Must provide either 'ids' array or scope='all_filtered'"},
            status_code=400,
        )

    if not batch_ids_to_delete:
        return JSONResponse({
            "ok": True,
            "deleted_count": 0,
            "item_deleted_count": 0,
            "deleted_ids": [],
            "not_found_ids": [],
            "message": "Không có đợt công văn nào phù hợp để xóa.",
        })

    # --- Perform hard delete ---
    # Get batches that actually exist
    batches = (
        db.query(BgCongVanBatchRow)
        .filter(BgCongVanBatchRow.id.in_(batch_ids_to_delete))
        .all()
    )
    confirmed_ids = [int(getattr(b, "id", 0) or 0) for b in batches]
    not_found = [bid for bid in batch_ids_to_delete if bid not in confirmed_ids]

    deleted_item_count = 0
    for batch in batches:
        # Hard delete process logs for items in this batch
        item_ids = [
            int(r.id)
            for r in db.query(BgCongVanRow.id).filter(BgCongVanRow.batch_id == int(getattr(batch, "id", 0))).all()
        ]
        if item_ids:
            db.query(BgCongVanProcessRow).filter(
                BgCongVanProcessRow.cong_van_id.in_(item_ids)
            ).delete(synchronize_session=False)

        # Hard delete all items in this batch
        item_count = db.query(BgCongVanRow).filter(
            BgCongVanRow.batch_id == int(getattr(batch, "id", 0))
        ).delete(synchronize_session=False)
        deleted_item_count += int(item_count or 0)

        # Hard delete the batch itself
        db.delete(batch)

    db.commit()

    return JSONResponse({
        "ok": True,
        "deleted_count": len(confirmed_ids),
        "item_deleted_count": deleted_item_count,
        "deleted_ids": confirmed_ids,
        "not_found_ids": not_found,
        "message": f"Đã xóa vĩnh viễn {len(confirmed_ids)} đợt công văn và {deleted_item_count} công văn.",
    })


@router.post("/items/bulk-delete")
def bulk_delete_items(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Bulk soft delete dispatch items.
    Two modes:
      A) explicit IDs:  { "ids": [211, 212] }
      B) all_filtered:  { "scope": "all_filtered", "filters": {...}, "confirm_text": "XOA TOAN BO" }

    Soft deletes matching non-deleted items only.
    Does NOT auto-delete parent batch.
    Does NOT delete any files.
    """
    user = _get_current_user(credentials=credentials, db=db)
    actor = str(getattr(user, "username", "") or "")
    now = datetime.utcnow()
    reason = str(payload.get("delete_reason") or "").strip() or None
    scope = str(payload.get("scope") or "").strip().lower()

    item_ids_to_delete: list[int] = []

    # --- Mode A: explicit IDs ---
    if scope == "ids" or ("ids" in payload and not scope):
        ids_raw = payload.get("ids", [])
        if not isinstance(ids_raw, list) or len(ids_raw) == 0:
            return JSONResponse(
                {"ok": False, "error": "ids must be a non-empty array"},
                status_code=400,
            )
        try:
            item_ids_to_delete = [int(x) for x in ids_raw]
        except (ValueError, TypeError):
            return JSONResponse({"ok": False, "error": "ids must contain integers"}, status_code=400)

    # --- Mode B: all_filtered ---
    elif scope == "all_filtered":
        confirm_text = str(payload.get("confirm_text") or "").strip().upper()
        if confirm_text != "XOA TOAN BO":
            return JSONResponse(
                {"ok": False, "error": "confirm_text must be exactly 'XOA TOAN BO'"},
                status_code=400,
            )
        filters = payload.get("filters", {})
        if not isinstance(filters, dict):
            filters = {}

        q = (
            db.query(BgCongVanRow)
            .filter(BgCongVanRow.batch_id.isnot(None))
            .filter(BgCongVanRow.deleted_at.is_(None))
        )

        f_year = filters.get("year")
        if f_year:
            try:
                yr = int(f_year)
                q = q.filter(BgCongVanRow.issue_date >= date(yr, 1, 1))
                q = q.filter(BgCongVanRow.issue_date <= date(yr, 12, 31))
            except (ValueError, TypeError):
                pass

        f_type = filters.get("dispatch_type")
        if f_type and str(f_type).strip():
            q = q.filter(BgCongVanRow.dispatch_type == str(f_type).strip())

        f_contact = filters.get("trang_thai_lien_he")
        if f_contact and str(f_contact).strip():
            q = q.filter(BgCongVanRow.trang_thai_lien_he == str(f_contact).strip())

        f_contract = filters.get("trang_thai_hop_dong")
        if f_contract and str(f_contract).strip():
            q = q.filter(BgCongVanRow.trang_thai_hop_dong == str(f_contract).strip())

        # Never allow all_filtered with no constraints
        if not (f_year or (f_type and str(f_type).strip())):
            return JSONResponse(
                {"ok": False, "error": "all_filtered requires at least year or dispatch_type filter"},
                status_code=400,
            )

        item_ids_to_delete = [int(getattr(r, "id", 0) or 0) for r in q.all()]
        item_ids_to_delete = [iid for iid in item_ids_to_delete if iid > 0]

    else:
        return JSONResponse(
            {"ok": False, "error": "Must provide either 'ids' array or scope='all_filtered'"},
            status_code=400,
        )

    if not item_ids_to_delete:
        return JSONResponse({
            "ok": True,
            "deleted_count": 0,
            "deleted_ids": [],
            "not_found_ids": [],
            "message": "Không có công văn nào phù hợp để xóa.",
        })

    # --- Perform hard delete ---
    # Delete process logs for these items first
    if item_ids_to_delete:
        db.query(BgCongVanProcessRow).filter(
            BgCongVanProcessRow.cong_van_id.in_(item_ids_to_delete)
        ).delete(synchronize_session=False)

    # Hard delete the items
    confirmed_count = (
        db.query(BgCongVanRow)
        .filter(BgCongVanRow.id.in_(item_ids_to_delete))
        .delete(synchronize_session=False)
    )

    confirmed_ids = [int(x) for x in item_ids_to_delete]
    not_found: list[int] = []

    db.commit()

    return JSONResponse({
        "ok": True,
        "deleted_count": int(confirmed_count or 0),
        "deleted_ids": confirmed_ids,
        "not_found_ids": not_found,
        "message": f"Đã xóa vĩnh viễn {int(confirmed_count or 0)} công văn.",
    })


@router.get("/tracking")
def list_tracking_items(
    year: int | None = Query(default=None),
    dispatch_type: str | None = Query(default=None),
    trang_thai_lien_he: str | None = Query(default=None),
    trang_thai_hop_dong: str | None = Query(default=None),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=50, ge=1, le=200),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    List all dispatch items for the tracking tab.
    Shows every unit across all batches with current status.
    """
    _get_current_user(credentials=credentials, db=db)
    selected_year = int(year) if year else None
    q = db.query(BgCongVanRow).filter(BgCongVanRow.batch_id.isnot(None))
    q = q.filter(BgCongVanRow.deleted_at.is_(None))
    if dispatch_type:
        q = q.filter(BgCongVanRow.dispatch_type == dispatch_type)
    if trang_thai_lien_he:
        q = q.filter(BgCongVanRow.trang_thai_lien_he == trang_thai_lien_he)
    if trang_thai_hop_dong:
        q = q.filter(BgCongVanRow.trang_thai_hop_dong == trang_thai_hop_dong)
    if selected_year:
        q = q.filter(BgCongVanRow.issue_date >= date(selected_year, 1, 1))
        q = q.filter(BgCongVanRow.issue_date <= date(selected_year, 12, 31))
    total = q.count()
    offset = (page - 1) * page_size
    items = q.order_by(BgCongVanRow.issue_date.desc(), BgCongVanRow.id.desc()).offset(offset).limit(page_size).all()
    rows_out = []
    for r in items:
        batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == int(getattr(r, "batch_id") or 0)).first()
        rows_out.append({
            "id": int(getattr(r, "id", 0) or 0),
            "batch_id": int(getattr(r, "batch_id", 0) or 0),
            "cong_van_no": str(getattr(r, "cong_van_no", "") or ""),
            "issue_date": _format_dd_mm_yyyy(getattr(r, "issue_date", None)),
            "recipient_unit": str(getattr(r, "recipient_unit", "") or ""),
            "recipient_address": str(getattr(r, "recipient_address", "") or ""),
            "so_dien_thoai": str(getattr(r, "recipient_phone", "") or ""),
            "dong_nguoi_nhan_bia_thu": str(getattr(r, "dong_nguoi_nhan_bia_thu", "") or ""),
            "lan_gui": int(getattr(r, "lan_gui", 1) or 1),
            "trang_thai_lien_he": str(getattr(r, "trang_thai_lien_he", "CHUA_LIEN_HE") or "CHUA_LIEN_HE"),
            "trang_thai_hop_dong": str(getattr(r, "trang_thai_hop_dong", "CHUA_KY_HOP_DONG") or "CHUA_KY_HOP_DONG"),
            "ngay_lien_he_gan_nhat": (
                getattr(r, "ngay_lien_he_gan_nhat", None).isoformat()
                if getattr(r, "ngay_lien_he_gan_nhat", None) else None
            ),
            "ghi_chu_lien_he": str(getattr(r, "ghi_chu_lien_he", "") or ""),
            "ngay_ky_hop_dong": (
                getattr(r, "ngay_ky_hop_dong", None).isoformat()
                if getattr(r, "ngay_ky_hop_dong", None) else None
            ),
            "contract_id": int(getattr(r, "contract_id") or 0) or None,
            "contract_no": str(getattr(r, "contract_no", "") or ""),
            "dispatch_type": str(getattr(r, "dispatch_type", "") or ""),
            "download_url": _download_url_for_path(getattr(r, "docx_path", None), year=None) if getattr(r, "docx_path", None) else "",
            "envelope_download_url": (
                _download_url_for_path(getattr(batch, "envelope_docx_path", None), year=None)
                if batch and getattr(batch, "envelope_docx_path", None) else ""
            ),
            "batch_cong_van_no": str(getattr(batch, "cong_van_no", "") or "") if batch else "",
            "created_at": getattr(r, "created_at", None).isoformat() if getattr(r, "created_at", None) else "",
        })
    total_pages = (total + page_size - 1) // page_size if total > 0 else 0
    return JSONResponse({
        "ok": True,
        "rows": rows_out,
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": total_pages,
    })


@router.get("/{dispatch_id}")
def get_dispatch_detail(
    dispatch_id: int,
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    row = db.query(BgCongVanRow).filter(BgCongVanRow.id == int(dispatch_id)).first()
    if row is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy công văn"}, status_code=404)
    issue_date = getattr(row, "issue_date", None)
    batch_id_val = int(getattr(row, "batch_id", 0) or 0)
    batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == batch_id_val).first() if batch_id_val > 0 else None
    return JSONResponse({
        "ok": True,
        "row": {
            "id": int(getattr(row, "id", 0) or 0),
            "batch_id": batch_id_val,
            "cong_van_no": str(getattr(row, "cong_van_no", "") or ""),
            "issue_date": _format_dd_mm_yyyy(issue_date),
            "contract_no": str(getattr(row, "contract_no", "") or ""),
            "recipient_unit": str(getattr(row, "recipient_unit", "") or ""),
            "recipient_address": str(getattr(row, "recipient_address", "") or ""),
            "recipient_contact": str(getattr(row, "recipient_contact", "") or ""),
            "recipient_phone": str(getattr(row, "recipient_phone", "") or ""),
            "expiry_date": _format_dd_mm_yyyy(getattr(row, "expiry_date", None)),
            "status": str(getattr(row, "status", "") or "draft"),
            "docx_path": str(getattr(row, "docx_path", "") or ""),
            "download_url": _download_url_for_path(getattr(row, "docx_path", None), year=(issue_date.year if issue_date else None)) if issue_date and getattr(row, "docx_path", None) else "",
            "batch_merged_download_url": _download_url_for_path(getattr(batch, "merged_docx_path", None), year=(issue_date.year if issue_date else None)) if batch else "",
            "batch_envelope_download_url": _download_url_for_path(getattr(batch, "envelope_docx_path", None), year=(issue_date.year if issue_date else None)) if batch else "",
            "batch_envelope_calibration_download_url": _download_url_for_path(getattr(batch, "envelope_calibration_docx_path", None), year=(issue_date.year if issue_date else None)) if batch else "",
            "note": str(getattr(row, "note", "") or ""),
        }
    })


@router.delete("/{dispatch_id}")
def delete_dispatch(
    dispatch_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    delete_file = bool(payload.get("delete_file", True)) if isinstance(payload, dict) else True
    batch_doc_paths: list[str] = []
    row = db.query(BgCongVanRow).filter(BgCongVanRow.id == int(dispatch_id)).first()
    if row is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy công văn"}, status_code=404)
    docx_path = str(getattr(row, "docx_path", "") or "").strip()
    row_batch_id = int(getattr(row, "batch_id", 0) or 0)
    db.query(BgCongVanProcessRow).filter(BgCongVanProcessRow.cong_van_id == int(dispatch_id)).delete(synchronize_session=False)
    db.delete(row)
    db.flush()
    if row_batch_id > 0:
        remains = db.query(BgCongVanRow.id).filter(BgCongVanRow.batch_id == row_batch_id).first()
        if remains is None:
            batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == row_batch_id).first()
            if batch is not None:
                batch_doc_paths = [
                    str(getattr(batch, "merged_docx_path", "") or "").strip(),
                    str(getattr(batch, "envelope_docx_path", "") or "").strip(),
                ]
                db.delete(batch)
    if delete_file and docx_path:
        try:
            p = Path(docx_path)
            if p.exists() and p.is_file():
                p.unlink(missing_ok=True)
        except Exception:
            pass
    if delete_file and batch_doc_paths:
        for raw in batch_doc_paths:
            if not raw:
                continue
            try:
                p = Path(raw)
                if p.exists() and p.is_file():
                    p.unlink(missing_ok=True)
            except Exception:
                pass
    return JSONResponse({"ok": True})


@router.post("/create-renewal")
def create_renewal_batch(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    user = _get_current_user(credentials=credentials, db=db)
    contract_ids_raw = payload.get("contract_ids") if isinstance(payload, dict) else None
    if not isinstance(contract_ids_raw, list):
        return JSONResponse({"ok": False, "error": "Thiếu danh sách contract_ids"}, status_code=400)
    today = date.today()
    issue_date_raw = str(payload.get("issue_date") or "").strip()
    issue_date = today
    if issue_date_raw:
        try:
            issue_date = date.fromisoformat(issue_date_raw)
        except Exception:
            issue_date = today
    start_cong_van_no = str(payload.get("start_cong_van_no") or "").strip()
    if not start_cong_van_no:
        return JSONResponse(
            {"ok": False, "error": "Số công văn là bắt buộc. Không tự sinh số công văn."},
            status_code=400,
        )
    should_merge = bool(payload.get("merge_output", True))
    should_create_envelope = bool(payload.get("create_envelope", True))
    contract_ids = [int(item or 0) for item in contract_ids_raw if item]
    contract_ids = [cid for cid in contract_ids if cid > 0]
    if not contract_ids:
        return JSONResponse({"ok": False, "error": "Danh sách hợp đồng rỗng"}, status_code=400)
    created_rows = []
    generated_paths: list[Path] = []
    merged_download_url = ""
    envelope_download_url = ""
    batch_cong_van_no = start_cong_van_no
    try:
        _ensure_cong_van_no_available(db=db, issue_date=issue_date, cong_van_no=batch_cong_van_no)
    except ValueError as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=409)
    batch_row = BgCongVanBatchRow(
        domain_group="background", field_code="karaoke",
        cong_van_no=batch_cong_van_no, issue_date=issue_date,
        dispatch_type=str(payload.get("dispatch_type") or "renewal_reminder").strip() or "renewal_reminder",
        template_name=str(payload.get("template_name") or "cong van_tai ky_karaoke.docx").strip() or None,
        note=str(payload.get("note") or "").strip() or None,
        total_items=0, envelope_total_items=0,
        created_by=str(getattr(user, "username", "") or ""),
        updated_by=str(getattr(user, "username", "") or ""),
    )
    db.add(batch_row)
    db.flush()
    batch_id = int(getattr(batch_row, "id", 0) or 0)
    for contract_id in contract_ids:
        contract = db.query(ContractRecordRow).filter(ContractRecordRow.id == contract_id).first()
        if contract is None or not _is_karaoke_contract_row(contract):
            continue
        out_dir = _CONGVAN_DOCX_ROOT / str(issue_date.year)
        out_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
        file_name = f"CONGVAN_TAIKY_{_slug_name(getattr(contract, 'don_vi_ten', '') or 'don-vi')}_{ts}.docx"
        out_path = (out_dir / file_name).resolve()
        recipient_unit = _normalize_unit_name(str(getattr(contract, "don_vi_ten", "") or ""))
        recipient_address = _normalize_address_spacing(str(getattr(contract, "don_vi_dia_chi", "") or ""))
        recipient_contact, recipient_phone = _resolve_recipient_contact_phone(contract)
        contract_no = str(getattr(contract, "contract_no", "") or "")

        context = {
            "so_cong_van": batch_cong_van_no,
            "TEN_DON_VI": recipient_unit,
            "ngay_ky_cong_van": f"{issue_date.day:02d}",
            "thang_ky_cong_van": f"{issue_date.month:02d}",
            "nam_ky_cong_van": str(issue_date.year),
            "ngay_ky_hop_dong": _format_dd_mm_yyyy(getattr(contract, "ngay_lap_hop_dong", None)),
            "ngay_het_hieu_luc_HD": _format_dd_mm_yyyy(getattr(contract, "ngay_ket_thuc", None)),
            "so_hop_dong": contract_no,
        }
        _render_cong_van_from_template(
            template_path=_CONG_VAN_TEMPLATE_PATH,
            output_path=out_path,
            context=context,
        )
        existing_count = (
            db.query(BgCongVanRow.id)
            .filter(BgCongVanRow.contract_id == int(getattr(contract, "id", 0) or 0))
            .filter(BgCongVanRow.dispatch_type == "renewal_reminder")
            .count()
        )
        attempt_no = existing_count + 1
        row = BgCongVanRow(
            domain_group="background", field_code="karaoke",
            batch_id=batch_id, cong_van_no=batch_cong_van_no, issue_date=issue_date,
            contract_id=int(getattr(contract, "id", 0) or 0), contract_no=contract_no,
            recipient_unit=recipient_unit, recipient_address=recipient_address,
            recipient_contact=recipient_contact, recipient_phone=recipient_phone,
            expiry_date=getattr(contract, "ngay_ket_thuc", None),
            dispatch_type="renewal_reminder",
            attempt_no=attempt_no,
            status="draft", note=str(payload.get("note") or "").strip() or None,
            docx_path=str(out_path),
            created_by=str(getattr(user, "username", "") or ""),
            updated_by=str(getattr(user, "username", "") or ""),
        )
        db.add(row)
        db.flush()
        db.add(BgCongVanProcessRow(
            cong_van_id=int(getattr(row, "id", 0) or 0),
            action_type="created", status_after="draft",
            note="Tạo công văn nhắc tái ký từ danh sách hợp đồng hết hạn",
            actor=str(getattr(user, "username", "") or ""),
        ))
        generated_paths.append(out_path)
        created_rows.append({
            "id": int(getattr(row, "id", 0) or 0),
            "batch_id": batch_id,
            "cong_van_no": batch_cong_van_no,
            "contract_no": contract_no,
            "attempt_no": attempt_no,
            "recipient_unit": recipient_unit,
            "download_url": _download_url_for_path(out_path, year=issue_date.year),
        })
    if not created_rows:
        db.delete(batch_row)
        return JSONResponse({"ok": False, "error": "Không có hợp đồng hợp lệ để tạo công văn trong batch."}, status_code=400)
    batch_row.total_items = len(created_rows)
    batch_row.updated_by = str(getattr(user, "username", "") or "")
    if should_merge and generated_paths:
        merged_name = f"CONGVAN_TAIKY_GOP_{issue_date.year}_{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
        merged_path = (_CONGVAN_DOCX_ROOT / str(issue_date.year) / merged_name).resolve()
        _merge_docx_files(inputs=generated_paths, output=merged_path)
        batch_row.merged_docx_path = str(merged_path)
        merged_download_url = _download_url_for_path(merged_path, year=issue_date.year)
    if should_create_envelope:
        envelope_path, envelope_count = _generate_envelope_for_batch(db=db, batch=batch_row, force_regenerate=True)
        batch_row.envelope_docx_path = str(envelope_path)
        batch_row.envelope_generated_at = datetime.utcnow()
        batch_row.envelope_total_items = int(envelope_count)
        envelope_download_url = _download_url_for_path(envelope_path, year=issue_date.year)

    # Extract envelope filename and timestamp from path
    envelope_filename = str(Path(envelope_download_url).name) if envelope_download_url else ""
    envelope_generated_at = batch_row.envelope_generated_at.isoformat() if batch_row.envelope_generated_at else ""

    return JSONResponse({
        "ok": True,
        "rows": created_rows,
        "merged_download_url": merged_download_url,
        "envelope_download_url": envelope_download_url,
        "envelope_filename": envelope_filename,
        "envelope_generated_at": envelope_generated_at,
        "total_created": len(created_rows),
        "batch_cong_van_no": batch_cong_van_no,
        "batch_id": batch_id,
    })


@router.post("/create-new-karaoke")
def create_new_karaoke_batch(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    user = _get_current_user(credentials=credentials, db=db)
    today = date.today()
    issue_date_raw = str(payload.get("issue_date") or "").strip()
    issue_date = today
    if issue_date_raw:
        try:
            issue_date = date.fromisoformat(issue_date_raw)
        except Exception:
            issue_date = today

    start_cong_van_no = str(payload.get("start_cong_van_no") or "").strip()
    if not start_cong_van_no:
        return JSONResponse(
            {"ok": False, "error": "Số công văn là bắt buộc. Không tự sinh số công văn."},
            status_code=400,
        )

    rows, issues = _coerce_new_karaoke_rows(payload.get("rows"))
    if not rows:
        return JSONResponse({"ok": False, "error": "Danh sách ký mới rỗng", "issues": issues}, status_code=400)
    if issues and not bool(payload.get("skip_invalid", False)):
        return JSONResponse({"ok": False, "error": "Có dòng thiếu dữ liệu bắt buộc", "issues": issues}, status_code=400)

    valid_row_indexes = {issue["row_index"] for issue in issues}
    rows_to_export = [row for index, row in enumerate(rows, start=1) if index not in valid_row_indexes] if issues else rows
    if not rows_to_export:
        return JSONResponse({"ok": False, "error": "Không còn dòng hợp lệ để xuất", "issues": issues}, status_code=400)

    should_merge = bool(payload.get("merge_output", True))
    should_create_envelope = bool(payload.get("create_envelope", False))
    try:
        _ensure_cong_van_no_available(db=db, issue_date=issue_date, cong_van_no=start_cong_van_no)
    except ValueError as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=409)

    base_note = str(payload.get("note") or "").strip()
    envelope_recipient_mode = str(payload.get("envelope_recipient_mode", "keep")).strip()
    envelope_custom_prefix = str(payload.get("envelope_custom_prefix", "")).strip()
    env_note_parts = []
    if envelope_recipient_mode and envelope_recipient_mode != "keep":
        env_note_parts.append(f"env_mode={envelope_recipient_mode}")
    if envelope_custom_prefix:
        env_note_parts.append(f"env_prefix={envelope_custom_prefix}")
    env_note_suffix = (" | " + "; ".join(env_note_parts)) if env_note_parts else ""
    batch_note = (base_note + env_note_suffix) or None

    batch_row = BgCongVanBatchRow(
        domain_group="background",
        field_code="karaoke",
        cong_van_no=start_cong_van_no,
        issue_date=issue_date,
        dispatch_type="new_karaoke",
        template_name="cong van_ky moi_karaoke.docx",
        note=batch_note,
        create_envelope=should_create_envelope,
        merge_output=should_merge,
        envelope_recipient_mode=envelope_recipient_mode if should_create_envelope else None,
        envelope_custom_prefix=envelope_custom_prefix if should_create_envelope else None,
        total_items=0,
        ready_items=len(rows_to_export),
        missing_items=len(issues),
        envelope_total_items=0,
        created_by=str(getattr(user, "username", "") or ""),
        updated_by=str(getattr(user, "username", "") or ""),
    )
    db.add(batch_row)
    db.flush()
    batch_id = int(getattr(batch_row, "id", 0) or 0)

    created_rows: list[dict[str, Any]] = []
    generated_paths: list[Path] = []
    out_dir = _CONGVAN_DOCX_ROOT / str(issue_date.year)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Envelope options for DONG_NGUOI_NHAN_BIA_THU computation — already loaded above

    def _compute_dong_nguoi_nhan_bia_thu(
        *, ten_don_vi: str, nguoi_nhan_bia_thu_row: str, mode: str, custom_prefix: str
    ) -> str:
        """Priority 1: NGUOI_NHAN_BIA_THU row value (if non-empty).
        Priority 2: apply UI mode to TEN_DON_VI.
        No duplication guard: skip adding prefix if TEN_DON_VI already starts with the prefix.
        """
        # Priority 1
        if nguoi_nhan_bia_thu_row.strip():
            return nguoi_nhan_bia_thu_row.strip()
        # Priority 2
        unit = ten_don_vi.strip()
        prefix_map = {
            "co_so": "Cơ sở kinh doanh ",
            "chu_co_so": "Chủ cơ sở kinh doanh ",
            "cong_ty": "Công ty ",
            "ho_kinh_doanh": "Hộ kinh doanh ",
        }
        if mode == "keep":
            return unit
        if mode == "custom":
            prefix = custom_prefix.strip() + " " if custom_prefix.strip() else ""
        else:
            prefix = prefix_map.get(mode, "")
        # Duplication guard
        if prefix and unit.startswith(prefix):
            return unit
        return prefix + unit

    # Pre-compute lan_gui (send round) for each row BEFORE creating items
    def _compute_lan_gui(ten_don_vi: str, dia_chi: str) -> int:
        """Look up previous dispatch items with same unit+address, return max lan_gui + 1."""
        norm_unit = _normalize_unit_name(ten_don_vi)
        norm_addr = _normalize_address_spacing(dia_chi)
        prev = (
            db.query(BgCongVanRow.lan_gui)
            .filter(BgCongVanRow.domain_group == "background")
            .filter(BgCongVanRow.dispatch_type == "new_karaoke")
            .filter(BgCongVanRow.recipient_unit == norm_unit)
            .filter(BgCongVanRow.recipient_address == norm_addr)
            .all()
        )
        if not prev:
            return 1
        return max(int(r.lan_gui or 1) for r in prev) + 1

    for index, row_data in enumerate(rows_to_export, start=1):
        ts = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
        file_name = f"CONGVAN_KYMOI_KARAOKE_{_slug_name(row_data.get('ten_don_vi') or 'don-vi')}_{ts}.docx"
        out_path = (out_dir / file_name).resolve()
        recipient_unit = row_data.get("ten_don_vi", "")
        recipient_address = row_data.get("dia_chi", "")
        recipient_phone = row_data.get("so_dien_thoai", "")
        # DONG_NGUOI_NHAN_BIA_THU: Priority 1 = row NGUOI_NHAN_BIA_THU; Priority 2 = apply UI mode
        dong_nguoi_nhan = _compute_dong_nguoi_nhan_bia_thu(
            ten_don_vi=recipient_unit,
            nguoi_nhan_bia_thu_row=str(row_data.get("nguoi_nhan_bia_thu") or ""),
            mode=envelope_recipient_mode,
            custom_prefix=envelope_custom_prefix,
        )
        # Compute lan_gui for this unit+address
        lan_gui = _compute_lan_gui(recipient_unit, recipient_address)
        context = _build_new_karaoke_context(row=row_data, cong_van_no=start_cong_van_no, issue_date=issue_date)
        _render_cong_van_from_template(
            template_path=_NEW_KARAOKE_TEMPLATE_PATH,
            output_path=out_path,
            context=context,
        )
        new_row = BgCongVanRow(
            domain_group="background",
            field_code="karaoke",
            batch_id=batch_id,
            cong_van_no=start_cong_van_no,
            issue_date=issue_date,
            contract_id=None,
            contract_no=None,
            recipient_unit=recipient_unit,
            recipient_address=recipient_address,
            recipient_contact=dong_nguoi_nhan,   # stored in recipient_contact (maps to DONG_NGUOI_NHAN_BIA_THU)
            recipient_phone=recipient_phone,
            expiry_date=None,
            dispatch_type="new_karaoke",
            attempt_no=index,
            lan_gui=lan_gui,
            dong_nguoi_nhan_bia_thu=dong_nguoi_nhan,
            trang_thai_lien_he="DA_GUI_CONG_VAN",
            trang_thai_hop_dong="CHUA_KY_HOP_DONG",
            status="draft",
            note=row_data.get("ghi_chu") or str(payload.get("note") or "").strip() or None,
            docx_path=str(out_path),
            created_by=str(getattr(user, "username", "") or ""),
            updated_by=str(getattr(user, "username", "") or ""),
        )
        db.add(new_row)
        db.flush()
        db.add(BgCongVanProcessRow(
            cong_van_id=int(getattr(new_row, "id", 0) or 0),
            action_type="created",
            status_after="draft",
            note="Tạo công văn ký mới Karaoke từ danh sách ngoài hệ thống",
            actor=str(getattr(user, "username", "") or ""),
        ))
        generated_paths.append(out_path)
        created_rows.append({
            "id": int(getattr(new_row, "id", 0) or 0),
            "batch_id": batch_id,
            "cong_van_no": start_cong_van_no,
            "contract_no": "",
            "attempt_no": index,
            "recipient_unit": recipient_unit,
            "recipient_address": recipient_address,
            "so_dien_thoai": recipient_phone,
            "nguoi_nhan_bia_thu": str(row_data.get("nguoi_nhan_bia_thu") or ""),
            "dong_nguoi_nhan_bia_thu": dong_nguoi_nhan,
            "lan_gui": lan_gui,
            "trang_thai_lien_he": "DA_GUI_CONG_VAN",
            "trang_thai_hop_dong": "CHUA_KY_HOP_DONG",
            "download_url": _download_url_for_path(out_path, year=issue_date.year),
        })

    if not created_rows:
        db.delete(batch_row)
        return JSONResponse({"ok": False, "error": "Không có đơn vị hợp lệ để tạo công văn", "issues": issues}, status_code=400)

    batch_row.total_items = len(created_rows)
    batch_row.updated_by = str(getattr(user, "username", "") or "")
    merged_download_url = ""
    envelope_download_url = ""
    if should_merge and generated_paths:
        merged_name = f"CONGVAN_KYMOI_KARAOKE_GOP_{issue_date.year}_{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
        merged_path = (_CONGVAN_DOCX_ROOT / str(issue_date.year) / merged_name).resolve()
        _merge_docx_files(inputs=generated_paths, output=merged_path)
        batch_row.merged_docx_path = str(merged_path)
        merged_download_url = _download_url_for_path(merged_path, year=issue_date.year)
    if should_create_envelope:
        try:
            # Build envelope using the new physical layout generator
            # Normalize raw row_data to recipients for the physical generator
            recipients_for_env: list[dict] = []
            for row_data in rows_to_export:
                normalized = _normalize_envelope_recipient(row_data)
                recipients_for_env.append({
                    "recipient_unit": normalized["name"],
                    "recipient_address": normalized["address"],
                    "recipient_contact": str(row_data.get("recipient_contact") or row_data.get("contact") or row_data.get("nguoi_nhan_bia_thu") or ""),
                    "recipient_phone": normalized["phone"],
                })
            ts = datetime.now().strftime("%Y%m%d-%H%M%S")
            out_name = f"BITHU_TAIKY_GOP_{issue_date.year}_{ts}.docx"
            out_path = (_CONGVAN_DOCX_ROOT / str(issue_date.year) / out_name).resolve()
            layout = _load_envelope_layout_config(db=db)
            _generate_vcpmc_envelope_docx(
                recipients=recipients_for_env,
                output=out_path,
                layout=layout,
                calibration=False,
            )
            batch_row.envelope_docx_path = str(out_path)
            batch_row.envelope_generated_at = datetime.utcnow()
            batch_row.envelope_total_items = len(rows_to_export)
            envelope_download_url = _download_url_for_path(out_path, year=issue_date.year)
            envelope_generated_at = batch_row.envelope_generated_at
        except Exception:
            envelope_download_url = ""
            envelope_generated_at = None

    db.commit()

    # Compute file count from actual URLs (not backend total_created alone)
    has_dispatch = bool(merged_download_url or (generated_paths and len(generated_paths) > 0))
    has_envelope = bool(envelope_download_url)
    total_files = (1 if has_dispatch else 0) + (1 if has_envelope else 0)

    # Extract envelope filename from path
    envelope_filename = str(Path(envelope_download_url).name) if envelope_download_url else ""

    return JSONResponse({
        "ok": True,
        "rows": created_rows,
        "merged_download_url": merged_download_url,
        "envelope_download_url": envelope_download_url,
        "envelope_filename": envelope_filename,
        "envelope_generated_at": envelope_generated_at.isoformat() if envelope_generated_at else "",
        "total_created": len(created_rows),
        "total_files": total_files,
        "batch_cong_van_no": start_cong_van_no,
        "batch_id": batch_id,
        "issues": issues,
        "template_name": "cong van_ky moi_karaoke.docx",
        "placeholders": NEW_KARAOKE_PLACEHOLDERS,
        "ready_count": len(rows_to_export),
        "total_input": len(rows),
    })


@router.post("/batches/{batch_id}/envelope")
def generate_batch_envelope(
    batch_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    force_regenerate = bool(payload.get("force_regenerate", True)) if isinstance(payload, dict) else True
    batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == int(batch_id)).first()
    if batch is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy batch công văn"}, status_code=404)
    try:
        envelope_path, envelope_count = _generate_envelope_for_batch(db=db, batch=batch, force_regenerate=force_regenerate)
    except ValueError as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=400)
    batch.envelope_docx_path = str(envelope_path)
    batch.envelope_generated_at = datetime.utcnow()
    batch.envelope_total_items = int(envelope_count)
    issue_date = getattr(batch, "issue_date", None) or date.today()
    envelope_url = _download_url_for_path(envelope_path, year=issue_date.year)
    layout = _load_envelope_layout_config(db=db)
    return JSONResponse({
        "ok": True,
        "batch_id": int(batch_id),
        "envelope_download_url": envelope_url,
        "envelope_total_items": int(envelope_count),
        "envelope_generated_at": datetime.utcnow().isoformat(),
        "layout": layout,
    })


@router.post("/batches/{batch_id}/envelope-calibration")
def generate_batch_envelope_calibration(
    batch_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    force_regenerate = bool(payload.get("force_regenerate", True)) if isinstance(payload, dict) else True
    batch = db.query(BgCongVanBatchRow).filter(BgCongVanBatchRow.id == int(batch_id)).first()
    if batch is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy batch công văn"}, status_code=404)
    try:
        calibration_path, envelope_count = _generate_envelope_for_batch(db=db, batch=batch, force_regenerate=force_regenerate, calibration=True)
    except ValueError as e:
        return JSONResponse({"ok": False, "error": str(e)}, status_code=400)
    batch.envelope_calibration_docx_path = str(calibration_path)
    batch.envelope_calibration_generated_at = datetime.utcnow()
    issue_date = getattr(batch, "issue_date", None) or date.today()
    calibration_url = _download_url_for_path(calibration_path, year=issue_date.year)
    layout = _load_envelope_layout_config(db=db)
    return JSONResponse({
        "ok": True,
        "batch_id": int(batch_id),
        "calibration_download_url": calibration_url,
        "envelope_total_items": int(envelope_count),
        "envelope_calibration_generated_at": datetime.utcnow().isoformat(),
        "layout": layout,
    })


@router.post("/envelope-test-230x170")
def create_envelope_test_230x170(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Create a test envelope file using main tray feed profile (flap left, face up)."""
    _get_current_user(credentials=credentials, db=db)

    # Use PORTRAIT_NORMAL profile (170x230, text upright)
    layout = _resolve_envelope_layout(VCPMC_PORTRAIT_NORMAL)

    recipients = [
        {
            "recipient_unit": "TEST BIA THU MAIN TRAY (1)",
            "recipient_address": "Vi tri Kinh gui",
            "recipient_contact": "",
            "recipient_phone": "0900 000 000",
        },
        {
            "recipient_unit": "TEST BIA THU MAIN TRAY (2)",
            "recipient_address": "Vi tri Kinh gui - trang 2",
            "recipient_contact": "",
            "recipient_phone": "0900 000 001",
        },
    ]

    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_name = f"TEST_BIA_THU_MAIN_TRAY_{ts}.docx"
    out_path = (_CONGVAN_DOCX_ROOT / str(date.today().year) / out_name).resolve()

    _generate_vcpmc_envelope_docx(recipients=recipients, output=out_path, layout=layout, calibration=False)

    download_url = _download_url_for_path(out_path, year=date.today().year)

    return JSONResponse({
        "ok": True,
        "download_url": download_url,
        "filename": out_name,
        "layout": layout,
        "message": "Test bia thu main tray (flap left, face up) X=130 Y=134",
    })


@router.post("/envelope-calibration-feed-orientation")
def envelope_calibration_feed_orientation(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Generate all 4 envelope orientation variants for calibration.

    User prints each file with the SAME feed direction (flap left, face up),
    then reports which one has text reading correctly in the "Kinh gui" area.

    Files:
      A. TEST_BIA_A_PORTRAIT_NORMAL.docx     — 170x230, text upright
      B. TEST_BIA_B_PORTRAIT_ROTATE_180.docx — 170x230, text upside-down
      C. TEST_BIA_C_LANDSCAPE_NORMAL.docx     — 230x170, text upright
      D. TEST_BIA_D_LANDSCAPE_ROTATE_180.docx — 230x170, text upside-down
    """
    _get_current_user(credentials=credentials, db=db)

    import datetime
    import os

    ts = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    year = datetime.date.today().year
    out_dir = os.path.join(
        os.environ.get("APPS_STORAGE_ROOT", "F:\\APPs\\storage"),
        "docx", str(year)
    )
    os.makedirs(out_dir, exist_ok=True)

    results: list[dict] = []

    for profile_key, profile in ALL_ENVELOPE_CALIBRATION_PROFILES:
        layout = _resolve_envelope_layout(profile)

        recipients = [
            {
                "recipient_unit":  f"PROFILE: {profile_key}",
                "recipient_address": "TEST DIA CHI",
                "recipient_contact": "",
                "recipient_phone":  "DT: 0900 000 000",
            },
        ]

        filename = f"TEST_BIA_{profile_key}_{ts}.docx"
        out_path = Path(out_dir) / filename

        _generate_vcpmc_envelope_docx(
            recipients=recipients,
            output=out_path,
            layout=layout,
            calibration=False,
        )

        download_url = _download_url_for_path(out_path, year=year)

        results.append({
            "profile_key": profile_key,
            "filename": filename,
            "download_url": download_url,
            "page_width_mm":  float(layout.get("page_width_mm", 0)),
            "page_height_mm": float(layout.get("page_height_mm", 0)),
            "rotate_180":     bool(layout.get("rotate_180", False)),
            "safe_margin_mm": float(layout.get("safe_margin_mm", 0)),
            "recipient_x_mm": float(layout.get("recipient_x_mm", 0)),
        })

    return JSONResponse({
        "ok": True,
        "count": len(results),
        "files": results,
        "message": "4 orientation variants generated. Print each file with the SAME feed direction to find the correct orientation.",
        "feed_instructions": {
            "flap_left": True,
            "face_up": True,
            "paper_direction": "head_first_into_tray",
            "rotate_between_tests": False,
        },
    })


@router.post("/envelope-test-canon")
def create_envelope_test_canon(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Create a test envelope file for Canon LBP325x printer calibration using physical layout."""
    _get_current_user(credentials=credentials, db=db)

    # Use PRINTER_PROFILES directly, allow payload overrides
    layout = PRINTER_PROFILES["canon_lbp325x"].copy()
    if isinstance(payload, dict):
        for k in payload:
            if k in layout:
                layout[k] = payload[k]

    recipients = [{
        "recipient_unit": "TEST CANON LBP325X",
        "recipient_address": "Vi tri can chinh bia thu",
        "recipient_contact": "",
        "recipient_phone": "0900 000 000",
    }]

    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_name = f"BITHU_TEST_CANON_{ts}.docx"
    out_path = (_CONGVAN_DOCX_ROOT / str(date.today().year) / out_name).resolve()

    _generate_vcpmc_envelope_docx(recipients=recipients, output=out_path, layout=layout, calibration=False)

    download_url = _download_url_for_path(out_path, year=date.today().year)

    return JSONResponse({
        "ok": True,
        "download_url": download_url,
        "filename": out_name,
        "layout": layout,
        "message": "Test envelope created with Canon LBP325x Portrait 17x23",
    })


@router.post("/envelope-test-brother")
def create_envelope_test_brother(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """Create a test envelope file for Brother HL-L2360D printer calibration using physical layout."""
    _get_current_user(credentials=credentials, db=db)

    # Merge Brother profile with base preset, allowing payload overrides
    brother_layout = {**VCPMC_BIA_THU_PRINTED, **BROTHER_HL_L2360D_PRINTER_PROFILE}
    if isinstance(payload, dict):
        for k in ("printer_offset_x_mm", "printer_offset_y_mm", "safe_margin_mm",
                  "recipient_x_mm", "recipient_top_mm", "recipient_block_width_mm",
                  "recipient_block_height_mm", "rotate_180"):
            if k in payload:
                brother_layout[k] = payload[k]

    layout = _resolve_envelope_layout(brother_layout)

    recipients = [{
        "recipient_unit": "TEST BROTHER HL-L2360D",
        "recipient_address": "Vi tri can chinh bia thu",
        "recipient_contact": "",
        "recipient_phone": "0900 000 000",
    }]

    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_name = f"BITHU_TEST_BROTHER_{ts}.docx"
    out_path = (_CONGVAN_DOCX_ROOT / str(date.today().year) / out_name).resolve()

    _generate_vcpmc_envelope_docx(recipients=recipients, output=out_path, layout=layout, calibration=False)

    download_url = _download_url_for_path(out_path, year=date.today().year)

    # Log printer profile debug info
    _debug_log("dispatches.py:envelope-test-brother", "Brother HL-L2360D test created", {
        "filename": out_name,
        "printer_profile": layout.get("printer_profile"),
        "rotate_180": layout.get("rotate_180"),
        "safe_margin_mm": layout.get("safe_margin_mm"),
        "printer_offset_x_mm": layout.get("printer_offset_x_mm"),
        "printer_offset_y_mm": layout.get("printer_offset_y_mm"),
    }, hypothesis_id="C,D")

    return JSONResponse({
        "ok": True,
        "download_url": download_url,
        "filename": out_name,
        "layout": layout,
        "message": f"Test envelope created with Brother HL-L2360D preset (rotate_180={layout.get('rotate_180')})",
    })


@router.post("/envelope-alignment-test")
def create_envelope_alignment_test(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Create a test file for the VCPMC physical layout.

    Body (optional overrides from DB config):
      - recipient_x_mm, recipient_width_mm, first_line_baseline_from_bottom_mm
      - font_baseline_offset_mm, line_gap_mm, font_size_pt
      - printer_offset_x_mm, printer_offset_y_mm

    File test: TEST TEN DON VI / Dia chi test (NO +, NO phone)
    Position: left_indent=recipient_x_mm, space_before=computed from baseline.
    """
    _get_current_user(credentials=credentials, db=db)

    base = _load_envelope_layout_config(db=db)
    if isinstance(payload, dict):
        for k in (
            "recipient_x_mm", "recipient_width_mm",
            "first_line_baseline_from_bottom_mm",
            "font_baseline_offset_mm", "line_gap_mm", "font_size_pt",
            "font_name",
            "printer_offset_x_mm", "printer_offset_y_mm",
            "page_width_mm", "page_height_mm",
            "phone_on_envelope", "phone_render_mode",
            # Printer profile / rotate support
            "printer_profile", "rotate_180", "safe_margin_mm",
            "recipient_block_width_mm", "recipient_block_height_mm",
        ):
            if k in payload:
                base[k] = payload[k]

    layout = _resolve_envelope_layout(base)

    page_h = float(layout.get("page_height_mm", 170.0))
    baseline_bot = float(layout.get("first_line_baseline_from_bottom_mm", 32.0))
    font_offset = float(layout.get("font_baseline_offset_mm", 4.0))
    baseline_y = page_h - baseline_bot
    para_top = baseline_y - font_offset
    rec_x = float(layout.get("recipient_x_mm", 130.0))
    rec_w = float(layout.get("recipient_width_mm", 95.0))
    line_gap = float(layout.get("line_gap_mm", 8.0))

    recipients = [{
        "recipient_unit": "TEST TEN DON VI",
        "recipient_address": "Dia chi test",
        "recipient_contact": "",
        "recipient_phone": "",
    }]

    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_name = f"BITHU_CHUAN_{ts}.docx"
    out_path = (_CONGVAN_DOCX_ROOT / str(date.today().year) / out_name).resolve()

    _generate_vcpmc_envelope_docx(recipients=recipients, output=out_path, layout=layout, calibration=False)

    download_url = _download_url_for_path(out_path, year=date.today().year)

    return JSONResponse({
        "ok": True,
        "download_url": download_url,
        "filename": out_name,
        "layout": layout,
        "anchor": {"x_mm": rec_x, "y_mm": para_top},
        "message": f"Test chuan: X={rec_x}mm, paragraph_top={para_top:.1f}mm (baseline {baseline_bot}mm from bottom). In bia thu that de kiem tra.",
    })


@router.post("/envelope-alignment-test-32mm")
def create_envelope_alignment_test_32mm(
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    """
    Create a test file for VCPMC physical layout using BASELINE-32mm anchor.

    File test:
      Line 1: TEST TEN DON VI
      Line 2: Dia chi test
    (NO +, NO border, NO phone)

    Body (optional overrides from DB config):
      - recipient_x_mm, recipient_width_mm, first_line_baseline_from_bottom_mm
      - font_baseline_offset_mm, line_gap_mm, font_size_pt
      - printer_offset_x_mm, printer_offset_y_mm
    """
    _get_current_user(credentials=credentials, db=db)

    base = _load_envelope_layout_config(db=db)
    if isinstance(payload, dict):
        for k in (
            "recipient_x_mm", "recipient_width_mm",
            "first_line_baseline_from_bottom_mm",
            "font_baseline_offset_mm", "line_gap_mm", "font_size_pt",
            "font_name",
            "printer_offset_x_mm", "printer_offset_y_mm",
            "page_width_mm", "page_height_mm",
            "phone_on_envelope", "phone_render_mode",
        ):
            if k in payload:
                base[k] = payload[k]

    layout = _resolve_envelope_layout(base)
    # Force no phone in test
    layout["phone_on_envelope"] = False

    recipients = [{
        "recipient_unit": "TEST TEN DON VI",
        "recipient_address": "Dia chi test",
        "recipient_contact": "",
        "recipient_phone": "",
    }]

    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    out_name = f"BITHU_CHUAN_{ts}.docx"
    out_path = (_CONGVAN_DOCX_ROOT / str(date.today().year) / out_name).resolve()

    _generate_vcpmc_envelope_docx(recipients=recipients, output=out_path, layout=layout, calibration=False)

    page_h = float(layout.get("page_height_mm", 170.0))
    baseline_bot = float(layout.get("first_line_baseline_from_bottom_mm", 32.0))
    font_offset = float(layout.get("font_baseline_offset_mm", 4.0))
    baseline_y = page_h - baseline_bot
    para_top = baseline_y - font_offset
    rec_x = float(layout.get("recipient_x_mm", 130.0))

    download_url = _download_url_for_path(out_path, year=date.today().year)

    return JSONResponse({
        "ok": True,
        "download_url": download_url,
        "filename": out_name,
        "layout": layout,
        "anchor": {"x_mm": rec_x, "y_mm": para_top},
        "baseline": {
            "from_bottom_mm": baseline_bot,
            "y_from_top_mm": baseline_y,
            "font_offset_mm": font_offset,
            "start_y_mm": para_top,
        },
        "message": f"Test chuan: X={rec_x}mm, paragraph_top={para_top:.1f}mm (baseline {baseline_bot}mm from bottom). Khong co + hay DT.",
    })


@router.get("/{dispatch_id}/logs")
def get_dispatch_logs(
    dispatch_id: int,
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    _get_current_user(credentials=credentials, db=db)
    rows = (
        db.query(BgCongVanProcessRow)
        .filter(BgCongVanProcessRow.cong_van_id == int(dispatch_id))
        .order_by(BgCongVanProcessRow.created_at.desc(), BgCongVanProcessRow.id.desc())
        .all()
    )
    out = [{
        "id": int(getattr(r, "id", 0) or 0),
        "action_type": str(getattr(r, "action_type", "") or ""),
        "status_after": str(getattr(r, "status_after", "") or ""),
        "note": str(getattr(r, "note", "") or ""),
        "actor": str(getattr(r, "actor", "") or ""),
        "created_at": getattr(r, "created_at", None).isoformat() if getattr(r, "created_at", None) else "",
    } for r in rows]
    return JSONResponse({"ok": True, "rows": out})


@router.post("/{dispatch_id}/logs")
def add_dispatch_log(
    dispatch_id: int,
    payload: dict = Body(default={}),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> JSONResponse:
    user = _get_current_user(credentials=credentials, db=db)
    action_type = str(payload.get("action_type") or "updated").strip().lower()
    status_after = str(payload.get("status_after") or "").strip().lower() or None
    note = str(payload.get("note") or "").strip() or None
    row = db.query(BgCongVanRow).filter(BgCongVanRow.id == int(dispatch_id)).first()
    if row is None:
        return JSONResponse({"ok": False, "error": "Không tìm thấy công văn"}, status_code=404)
    if status_after:
        row.status = status_after
    if note:
        row.note = note
    row.updated_by = str(getattr(user, "username", "") or "")
    db.add(BgCongVanProcessRow(
        cong_van_id=int(dispatch_id),
        action_type=action_type or "updated",
        status_after=status_after or str(getattr(row, "status", "") or ""),
        note=note,
        actor=str(getattr(user, "username", "") or ""),
    ))
    return JSONResponse({"ok": True})


# =============================================================================
# File download endpoint
# =============================================================================

@download_router.get("/{year}/{filename}", response_model=None)
def download_dispatch_file(
    year: str,
    filename: str,
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
) -> FileResponse | JSONResponse:
    """Download a generated DOCX file for Công văn."""
    # Auth required - verify user is logged in
    _get_current_user(credentials=credentials, db=db)
    safe_year = str(year or "").strip()
    safe_filename = str(filename or "").strip()
    if not safe_filename:
        return JSONResponse({"detail": "Missing filename"}, status_code=400)
    # Look in F:\APPs\storage\docx\{year}\{filename}
    file_path = _CONGVAN_DOCX_ROOT / safe_year / safe_filename
    if not file_path.exists():
        return JSONResponse({"detail": f"File not found: {safe_filename}"}, status_code=404)
    return FileResponse(
        path=str(file_path.resolve()),
        filename=safe_filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
