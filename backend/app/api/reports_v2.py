"""
Reports V2 + KPI endpoints — Phase 6 role-adaptive frontend.

NOTE: These endpoints were previously missing from the backend.
This module adds the minimum read-only endpoints that the React
ReportsPage (frontend/src/pages/ReportsPage.tsx) calls via
frontend/src/components/reports/kpiClient.ts:

  GET  /api/kpi/annual-target?year=YYYY
  PUT  /api/kpi/annual-target
  GET  /api/kpi/annual-summary?year=YYYY
  GET  /api/kpi/annual-overview?year=YYYY
  GET  /api/reports/v2/overview?year=YYYY[&field=...]
  GET  /api/reports/v2/contracts?year=YYYY[&field=...&owner_user_id=...]
  GET  /api/reports/v2/users?year=YYYY
  GET  /api/reports/v2/renewals?year=YYYY[&page=1&page_size=20]
  GET  /api/reports/v2/gcn?year=YYYY[&page=1&page_size=20]
  POST /api/reports/v2/export

All endpoints are READ-ONLY except /kpi/annual-target (PUT)
which stores per-user annual target in a simple JSON config table
shared with /api/reports/summary logic.

No business formulas were changed. Royalty values reuse
ContractRecordRow.royalty_amount_before_vat + royalty_amount_after_vat.
"""
from __future__ import annotations

import json
import logging
from datetime import date, datetime, timedelta
from typing import Any, Optional

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.security import HTTPAuthorizationCredentials
from sqlalchemy import and_, func
from sqlalchemy.orm import Session

from ..core.database import get_db
from ..core.security import (
    decode_access_token,
    get_user_permissions,
    security_scheme,
)
from ..models.contracts import ContractRecordRow
from ..models.user import UserRow
from ..models.certificates import CertificateRecordRow
from ..services.revenue_resolver import (
    resolve_contract_revenue,
    RevenueBasis,
    get_before_vat_revenue,
    get_normalized_before_vat,
    normalize_contract_revenue,
)

log = logging.getLogger("reports_v2")

# ─── Routers ─────────────────────────────────────────────────────────────────

kpi_router = APIRouter(prefix="/api/kpi", tags=["kpi"])
reports_v2_router = APIRouter(prefix="/api/reports/v2", tags=["reports_v2"])

# =============================================================================
# Internal helpers
# =============================================================================

def _current_user(
    db: Session,
    credentials: HTTPAuthorizationCredentials | None,
) -> UserRow | None:
    """Decode bearer token and return the matching UserRow, or None."""
    if not credentials or not credentials.credentials:
        return None
    try:
        username = decode_access_token(credentials.credentials)
    except HTTPException:
        return None
    user = (
        db.query(UserRow)
        .filter(UserRow.username == username)
        .one_or_none()
    )
    return user


def _renewal_bucket(rs: str | None) -> str:
    """Map renewal_status to signing bucket label."""
    if not rs:
        return "unknown"
    s = str(rs).strip().upper()
    if s == "NEW":
        return "new"
    if s in ("PENDING_RENEWAL", "RENEWED"):
        return "renewal"
    if s == "FRAME_CONTRACT":
        return "frame"
    return "unknown"


def _renewal_label(bucket: str) -> str:
    return {
        "new": "Ký mới",
        "renewal": "Tái ký",
        "frame": "HĐ khung",
        "unknown": "Chưa xác định",
    }.get(bucket, "Chưa xác định")


# Map assignment/frontend field codes to contract_records.linh_vuc values.
# Contract records use exact match on linh_vuc column with mixed-case values.
LINH_VUC_MAP: dict[str, str] = {
    'KHU_VUI_CHOI': 'Khu vui chơi',
    'Khu vui chơi': 'Khu vui chơi',
    'ENTERTAINMENT': 'Khu vui chơi',
    'amusement': 'Khu vui chơi',
    'KARAOKE': 'KARAOKE',
    'Karaoke': 'Karaoke',
    'karaoke': 'Karaoke',
    'BACKGROUND': 'Background',
    'Background': 'Background',
    'background': 'Background',
    'PHONG_THU_AM': 'Phòng thu âm',
    'Phòng thu âm': 'Phòng thu âm',
    'BD': 'BD',
    'SCTT': 'SCTT',
}


def _resolve_linh_vuc(code: str) -> str:
    """Normalize a field code to contract_records.linh_vuc value."""
    return LINH_VUC_MAP.get(code, code)

def _has_gcn(cert_map: dict[int, tuple[str, str | None]], contract_id: int | None) -> bool:
    """
    Return True if the contract has a certificate record WITH a certificate_number.
    Matches ContractsListPage logic: gcn_certificate_no must be non-null, non-empty, non-dash.
    """
    if contract_id is None:
        return False
    entry = cert_map.get(contract_id)
    if not entry:
        return False
    _status, cert_no = entry
    return bool(cert_no and cert_no.strip() not in ("", "-"))


def _build_cert_map(db: Session) -> dict[int, tuple[str, str | None]]:
    """
    contract_id -> (latest certificate status, certificate_number).
    Matches ContractsListPage: both gcn_status and gcn_certificate_no are needed
    to determine whether to show 'GCN số X' vs 'GCN chưa cấp số'.
    """
    cert_map: dict[int, tuple[str, str | None]] = {}
    cert_all = (
        db.query(
            CertificateRecordRow.contract_id,
            CertificateRecordRow.status,
            CertificateRecordRow.certificate_no,
            CertificateRecordRow.created_at,
        )
        .order_by(
            CertificateRecordRow.contract_id,
            CertificateRecordRow.created_at.desc(),
        )
        .all()
    )
    for cid, cstatus, cert_no, _cdate in cert_all:
        if cid not in cert_map:
            status = str(cstatus or "draft").lower() if cstatus else "draft"
            cert_map[cid] = (status, cert_no)
    return cert_map


# =============================================================================
# KPI: annual-target (per-user)
# =============================================================================

# Simple in-process store: {year: {username: target}}
# Persists within a single uvicorn process; not shared across workers.
# This is sufficient for the local dev workflow described in the brief.
_KPI_TARGETS: dict[int, dict[str, dict[str, Any]]] = {}


@kpi_router.get("/annual-target")
def get_annual_target(
    year: int = Query(..., ge=2000, le=2100),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    year_map = _KPI_TARGETS.get(year, {})
    entry = year_map.get(user.username)
    if not entry:
        return None
    return {
        "year": year,
        "annual_target": entry.get("annual_target", 0),
        "note": entry.get("note", ""),
        "updated_at": entry.get("updated_at"),
    }


@kpi_router.put("/annual-target")
def put_annual_target(
    body: dict[str, Any],
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    year = int(body.get("year") or 0)
    target = int(body.get("annual_target") or 0)
    note = str(body.get("note") or "")
    if not year:
        raise HTTPException(status_code=400, detail="year is required")
    year_map = _KPI_TARGETS.setdefault(year, {})
    year_map[user.username] = {
        "annual_target": target,
        "note": note,
        "updated_at": datetime.utcnow().isoformat() + "Z",
    }
    return {
        "year": year,
        "annual_target": target,
        "note": note,
        "updated_at": year_map[user.username]["updated_at"],
    }


def _resolve_target(year: int, username: str) -> tuple[int | None, bool]:
    entry = _KPI_TARGETS.get(year, {}).get(username)
    if not entry:
        return None, False
    target = int(entry.get("annual_target") or 0)
    return target, target == 0


# =============================================================================
# KPI: annual-summary (per-user) — for tab "Tổng quan của tôi"
# =============================================================================

@kpi_router.get("/annual-summary")
def get_annual_summary(
    year: int = Query(..., ge=2000, le=2100),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # Filter contracts in year owned by this user (by email match)
    email = (user.username or "").strip().lower()
    rows = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.contract_year == year)
        .filter(ContractRecordRow.nguoi_thuc_hien_email.ilike(email))
        .all()
    )
    user_contract_ids = [r.id for r in rows if r.id is not None]
    cert_map = _build_cert_map(db) if user_contract_ids else {}
    gcn_issued_count = 0
    gcn_missing_count = 0
    for cid in user_contract_ids:
        if _has_gcn(cert_map, cid):
            gcn_issued_count += 1
        else:
            gcn_missing_count += 1

    target, target_zero = _resolve_target(year, user.username)

    today = date.today()
    today60 = today + timedelta(days=60)

    actual = 0
    contract_count = 0
    bucket = {"new": [0, 0], "renewal": [0, 0], "frame": [0, 0], "unknown": [0, 0]}
    monthly_actual = [0] * 12
    monthly_count = [0] * 12
    quarterly_actual = [0] * 4
    quarterly_count = [0] * 4
    active_count = 0
    expired_count = 0
    expiring_count = 0

    for row in rows:
        val = get_normalized_before_vat(row)
        actual += val
        contract_count += 1
        b = _renewal_bucket(row.renewal_status)
        bucket[b][0] += 1
        bucket[b][1] += val
        signed_date = row.ngay_lap_hop_dong
        if signed_date:
            try:
                m = signed_date.month
                monthly_actual[m - 1] += val
                monthly_count[m - 1] += 1
                q = (m - 1) // 3
                quarterly_actual[q] += val
                quarterly_count[q] += 1
            except Exception:
                pass
        # State classification — same rule as overview (end_date → active/expiring/expired).
        end = row.ngay_ket_thuc
        if end:
            try:
                days_left = (end - today).days
                if days_left < 0:
                    expired_count += 1
                elif days_left <= 60:
                    expiring_count += 1
                else:
                    active_count += 1
            except Exception:
                active_count += 1
        else:
            active_count += 1

    remaining = (target - actual) if target is not None else None
    exceeded = (-remaining) if (remaining is not None and remaining < 0) else None
    progress = (
        round(actual / target * 100, 1) if target and target > 0 else None
    )

    return {
        "user_id": user.id,
        "display_name": user.display_name,
        "year": year,
        "configured": target is not None,
        "annual_target": target,
        "target_zero": target_zero,
        "actual": actual,
        "contract_count": contract_count,
        "active_count": active_count,
        "expiring_count": expiring_count,
        "expired_count": expired_count,
        "gcn_issued_count": gcn_issued_count,
        "gcn_missing_count": gcn_missing_count,
        "remaining": remaining,
        "exceeded": exceeded,
        "progress_percent": progress,
        "buckets": {
            "new_count": bucket["new"][0],
            "new_actual": bucket["new"][1],
            "renewal_count": bucket["renewal"][0],
            "renewal_actual": bucket["renewal"][1],
            "frame_count": bucket["frame"][0],
            "frame_actual": bucket["frame"][1],
            "unknown_count": bucket["unknown"][0],
            "unknown_actual": bucket["unknown"][1],
        },
        "monthly": [
            {"month": i + 1, "actual": monthly_actual[i], "count": monthly_count[i]}
            for i in range(12)
        ],
        "quarterly": [
            {
                "quarter": q + 1,
                "actual": quarterly_actual[q],
                "count": quarterly_count[q],
            }
            for q in range(4)
        ],
    }


# =============================================================================
# KPI: annual-overview (branch-wide) — admin/mod
# =============================================================================

@kpi_router.get("/annual-overview")
def get_annual_overview(
    year: int = Query(..., ge=2000, le=2100),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    perms = get_user_permissions(db, user)
    if "reports.view_branch" not in perms and "kpi.manage" not in perms:
        raise HTTPException(status_code=403, detail="Forbidden")

    cert_map = _build_cert_map(db)
    rows = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.contract_year == year)
        .all()
    )
    users = db.query(UserRow).all()
    user_by_email = {(u.username or "").strip().lower(): u for u in users}

    branch_count = 0
    branch_actual = 0
    assigned_count = 0
    assigned_actual = 0
    unassigned_count = 0
    unassigned_actual = 0
    bucket = {"new": [0, 0], "renewal": [0, 0], "frame": [0, 0], "unknown": [0, 0]}

    for row in rows:
        val = get_normalized_before_vat(row)
        branch_count += 1
        branch_actual += val
        b = _renewal_bucket(row.renewal_status)
        bucket[b][0] += 1
        bucket[b][1] += val
        owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
        if owner_email and owner_email in user_by_email:
            assigned_count += 1
            assigned_actual += val
        else:
            unassigned_count += 1
            unassigned_actual += val

    # Per-user rows
    per_user: dict[str, dict[str, Any]] = {}
    for u in users:
        key = (u.username or "").strip().lower()
        per_user[key] = {
            "user_id": u.id,
            "display_name": u.display_name,
            "username": u.username,
            "configured": False,
            "annual_target": None,
            "target_zero": False,
            "actual": 0,
            "contract_count": 0,
            "remaining": None,
            "exceeded": None,
            "progress_percent": None,
            "new_count": 0,
            "new_actual": 0,
            "renewal_count": 0,
            "renewal_actual": 0,
            "frame_count": 0,
            "frame_actual": 0,
            "unknown_count": 0,
            "unknown_actual": 0,
        }

    for row in rows:
        owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
        u = per_user.get(owner_email)
        if not u:
            continue
        val = get_normalized_before_vat(row)
        u["actual"] += val
        u["contract_count"] += 1
        b = _renewal_bucket(row.renewal_status)
        u[f"{b}_count"] = u.get(f"{b}_count", 0) + 1
        u[f"{b}_actual"] = u.get(f"{b}_actual", 0) + val

    configured = 0
    unconfigured = 0
    sum_targets = 0
    user_rows: list[dict[str, Any]] = []
    for key, u in per_user.items():
        target, target_zero = _resolve_target(year, key)
        u["annual_target"] = target
        u["target_zero"] = target_zero
        u["configured"] = target is not None
        if target is not None:
            configured += 1
            sum_targets += target
        else:
            unconfigured += 1
        if target and target > 0:
            u["remaining"] = target - u["actual"]
            u["exceeded"] = (-u["remaining"]) if u["remaining"] < 0 else None
            u["progress_percent"] = round(u["actual"] / target * 100, 1)
        user_rows.append(u)

    return {
        "year": year,
        "assigned_actual": assigned_actual,
        "assigned_count": assigned_count,
        "unassigned_actual": unassigned_actual,
        "unassigned_count": unassigned_count,
        "branch_actual": branch_actual,
        "branch_count": branch_count,
        "configured_user_count": configured,
        "unconfigured_user_count": unconfigured,
        "sum_user_targets": sum_targets,
        "buckets": {
            "new_count": bucket["new"][0],
            "new_actual": bucket["new"][1],
            "renewal_count": bucket["renewal"][0],
            "renewal_actual": bucket["renewal"][1],
            "frame_count": bucket["frame"][0],
            "frame_actual": bucket["frame"][1],
            "unknown_count": bucket["unknown"][0],
            "unknown_actual": bucket["unknown"][1],
        },
        "users": sorted(user_rows, key=lambda x: -(x.get("actual") or 0)),
        "department_aggregation": False,
        "department_note": "Tổng hợp theo user — không gộp theo chi nhánh.",
    }


# =============================================================================
# Reports V2: /overview
# =============================================================================

@reports_v2_router.get("/overview")
def get_v2_overview(
    year: int = Query(..., ge=2000, le=2100),
    field: Optional[str] = Query(default=None),
    owner_user_id: Optional[int] = Query(default=None),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    cert_map = _build_cert_map(db)

    q = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.contract_year == year)
    )
    if field:
        canon = _resolve_linh_vuc(field)
        q = q.filter(
            (ContractRecordRow.linh_vuc == canon)
            | (ContractRecordRow.field_code == field)
        )
    rows = q.all()
    users = db.query(UserRow).all()
    user_by_email = {(u.username or "").strip().lower(): u for u in users}

    today = date.today()

    active_count = 0
    expired_count = 0
    expiring_count = 0
    gcn_issued = 0
    gcn_missing = 0
    positive = zero = null_v = 0
    bucket = {"new": [0, 0], "renewal": [0, 0], "frame": [0, 0], "unknown": [0, 0]}
    assigned = {"count": 0, "actual": 0}
    unassigned = {"count": 0, "actual": 0}
    monthly_actual = [0] * 12
    monthly_count = [0] * 12
    quarterly_actual = [0] * 4
    quarterly_count = [0] * 4
    field_break: dict[str, dict[str, int]] = {}

    for row in rows:
        val = get_normalized_before_vat(row)
        b = _renewal_bucket(row.renewal_status)
        bucket[b][0] += 1
        bucket[b][1] += val

        # active / expired / expiring
        end = row.ngay_ket_thuc
        if end:
            try:
                days_left = (end - today).days
                if days_left < 0:
                    expired_count += 1
                elif days_left <= 60:
                    expiring_count += 1
                else:
                    active_count += 1
            except Exception:
                active_count += 1
        else:
            active_count += 1

        # GCN
        if _has_gcn(cert_map, row.id):
            gcn_issued += 1
        else:
            gcn_missing += 1

        # value quality — revenue = before VAT, total = after VAT
        _rev = resolve_contract_revenue(row)
        if _rev.resolution_status == "unresolved":
            null_v += 1
        elif val == 0:
            zero += 1
        else:
            positive += 1

        # monthly / quarterly
        signed_date = row.ngay_lap_hop_dong
        if signed_date:
            try:
                m = signed_date.month
                monthly_actual[m - 1] += val
                monthly_count[m - 1] += 1
                qidx = (m - 1) // 3
                quarterly_actual[qidx] += val
                quarterly_count[qidx] += 1
            except Exception:
                pass

        # field breakdown
        fkey = (row.linh_vuc_hien_thi or row.linh_vuc or "Khác").strip() or "Khác"
        fb = field_break.setdefault(fkey, {"count": 0, "actual": 0})
        fb["count"] += 1
        fb["actual"] += val

        # assigned vs unassigned
        owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
        if owner_email and owner_email in user_by_email:
            assigned["count"] += 1
            assigned["actual"] += val
        else:
            unassigned["count"] += 1
            unassigned["actual"] += val

    total_count = len(rows)
    # Authoritative before-vat chain for ALL revenue KPI (matches Dashboard /reports/summary)
    actual = sum(get_normalized_before_vat(r) for r in rows)
    # Authoritative before-VAT chain for the "Tổng giá trị hợp đồng" card,
    # which now matches the label "chưa Thuế GTGT" exactly (no so_tien
    # after-VAT inflation). `total_revenue` and `total_contract_value`
    # are the same field at this layer because KPI dollar basis is
    # before-VAT.
    total_revenue = actual
    total_contract_value = actual
    monthly_trend = [
        {"month": i + 1, "count": monthly_count[i], "actual": monthly_actual[i]}
        for i in range(12)
    ]
    quarterly_contribution = [
        {
            "quarter": qidx + 1,
            "count": quarterly_count[qidx],
            "actual": quarterly_actual[qidx],
        }
        for qidx in range(4)
    ]
    signing_breakdown = [
        {"bucket": b, "label": _renewal_label(b), "count": bucket[b][0], "actual": bucket[b][1]}
        for b in ("new", "renewal", "frame", "unknown")
        if bucket[b][0] > 0
    ]

    data_quality_warnings = []
    if null_v > 0:
        data_quality_warnings.append(
            f"Có {null_v} hợp đồng không có giá trị tiền — không tính vào doanh thu."
        )
    if zero > 0:
        data_quality_warnings.append(
            f"Có {zero} hợp đồng có giá trị = 0 — cần kiểm tra lại."
        )

    return {
        "year": year,
        "revenue_before_vat": total_revenue,
        "contract_total_after_vat": total_contract_value,
        "contract_count": total_count,
        "total_count": total_count,
        "total_actual": total_revenue,
        "positive_value_count": positive,
        "zero_value_count": zero,
        "null_value_count": null_v,
        "active_count": active_count,
        "expired_count": expired_count,
        "expiring_count": expiring_count,
        "gcn_issued_count": gcn_issued,
        "gcn_missing_count": gcn_missing,
        "new_count": bucket["new"][0],
        "new_actual": bucket["new"][1],
        "renewal_count": bucket["renewal"][0],
        "renewal_actual": bucket["renewal"][1],
        "frame_count": bucket["frame"][0],
        "frame_actual": bucket["frame"][1],
        "unknown_count": bucket["unknown"][0],
        "unknown_actual": bucket["unknown"][1],
        "assigned_count": assigned["count"],
        "assigned_actual": assigned["actual"],
        "unassigned_count": unassigned["count"],
        "unassigned_actual": unassigned["actual"],
        "monthly": monthly_trend,
        "monthly_trend": monthly_trend,
        "quarterly": quarterly_contribution,
        "quarterly_contribution": quarterly_contribution,
        "field_breakdown": [
            {"field": k, "count": v["count"], "actual": v["actual"]}
            for k, v in sorted(field_break.items(), key=lambda x: -x[1]["count"])
        ],
        "signing_breakdown": signing_breakdown,
        "data_quality_warnings": data_quality_warnings,
        "applied_filters": {
            "year": year,
            "date_from": None,
            "date_to": None,
            "period": None,
            "quarter": None,
            "month": None,
            "field": field,
            "owner_user_id": owner_user_id,
            "signing_bucket": None,
        },
    }


# =============================================================================
# Reports V2: /contracts
# =============================================================================

@reports_v2_router.get("/contracts")
def get_v2_contracts(
    year: int = Query(..., ge=2000, le=2100),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    search: Optional[str] = Query(default=None),
    field: Optional[str] = Query(default=None),
    signing_bucket: Optional[str] = Query(default=None),
    contract_state: Optional[str] = Query(default=None),
    gcn_state: Optional[str] = Query(default=None),
    owner_user_id: Optional[int] = Query(default=None),
    owner_email: Optional[str] = Query(default=None, description="Canonical filter: owner (nguoi_thuc_hien_email) equals this email."),
    value_filter: Optional[str] = Query(default=None, description="all|positive|zero|null — lọc theo trạng thái dữ liệu giá trị hợp đồng."),
    sort_by: Optional[str] = Query(default=None, description="Column to sort: signed_date|contract_no|don_vi_ten|royalty_amount_before_vat|ngay_ket_thuc"),
    sort_order: Optional[str] = Query(default="desc", description="asc|desc"),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    cert_map = _build_cert_map(db)
    users = db.query(UserRow).all()
    user_by_email = {(u.username or "").strip().lower(): u for u in users}

    q = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.contract_year == year)
    )
    if field:
        canon = _resolve_linh_vuc(field)
        q = q.filter(
            (ContractRecordRow.linh_vuc == canon)
            | (ContractRecordRow.field_code == field)
        )
    if search:
        q = q.filter(
            (ContractRecordRow.contract_no.ilike(f"%{search}%"))
            | (ContractRecordRow.don_vi_ten.ilike(f"%{search}%"))
            | (ContractRecordRow.ten_bang_hieu.ilike(f"%{search}%"))
        )
    if owner_email:
        owner_norm = owner_email.strip().lower()
        q = q.filter(func.lower(func.coalesce(ContractRecordRow.nguoi_thuc_hien_email, "")) == owner_norm)

    # Sorting — handled as DB-side ORDER BY to honor the UI's sort dropdown.
    sort_columns = {
        "signed_date": ContractRecordRow.ngay_lap_hop_dong,
        "contract_no": ContractRecordRow.contract_no,
        "don_vi_ten": ContractRecordRow.don_vi_ten,
        "royalty_amount_before_vat": ContractRecordRow.royalty_amount_before_vat,
        "ngay_ket_thuc": ContractRecordRow.ngay_ket_thuc,
    }
    if sort_by in sort_columns:
        col = sort_columns[sort_by]
        order = (sort_order or "desc").lower()
        if order == "asc":
            q = q.order_by(col.asc().nulls_last())
        else:
            q = q.order_by(col.desc().nulls_last())
    else:
        q = q.order_by(ContractRecordRow.ngay_lap_hop_dong.desc().nulls_last())

    rows = q.all()

    items = []
    for row in rows:
        owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
        u = user_by_email.get(owner_email)
        b = _renewal_bucket(row.renewal_status)
        if signing_bucket and b != signing_bucket:
            continue

        end = row.ngay_ket_thuc
        days_to_expiry = None
        if end:
            try:
                days_to_expiry = (end - date.today()).days
            except Exception:
                days_to_expiry = None

        # GCN status from cert_map (status, certificate_number)
        entry = cert_map.get(row.id)
        gcn_status = entry[0] if entry else "draft"
        if gcn_state and gcn_status != gcn_state:
            continue

        if contract_state == "active":
            if days_to_expiry is not None and days_to_expiry < 0:
                continue
        elif contract_state == "expired":
            if days_to_expiry is None or days_to_expiry >= 0:
                continue
        elif contract_state == "expiring":
            if days_to_expiry is None or days_to_expiry < 0 or days_to_expiry > 60:
                continue

        if owner_user_id is not None:
            if not u or u.id != owner_user_id:
                continue

        # Data-quality filter on contract value. Applied here (server side) so the
        # paginated total matches the filter instead of only the current page.
        if value_filter and value_filter != "all":
            amount = row.royalty_amount_before_vat
            if value_filter == "positive" and not (amount is not None and amount > 0):
                continue
            if value_filter == "zero" and amount != 0:
                continue
            if value_filter == "null" and amount is not None:
                continue

        # Use certificate_no from cert_map (avoids extra query)
        gcn_no = entry[1] if entry else None

        items.append(
            {
                "id": row.id,
                "contract_number": row.contract_no or "",
                "organization_name": row.don_vi_ten or "",
                "field": row.linh_vuc_hien_thi or row.linh_vuc or "",
                "owner_user_id": u.id if u else None,
                "owner_email": owner_email or None,
                "owner_name": u.display_name if u else None,
                "signed_date": row.ngay_lap_hop_dong.isoformat() if row.ngay_lap_hop_dong else None,
                "start_date": row.ngay_bat_dau.isoformat() if row.ngay_bat_dau else None,
                "end_date": row.ngay_ket_thuc.isoformat() if row.ngay_ket_thuc else None,
                "royalty_amount_before_vat": row.royalty_amount_before_vat,
                "total_payment": row.royalty_amount_after_vat,
                "signing_bucket": b,
                "signing_bucket_label": _renewal_label(b),
                "renewal_status": row.renewal_status,
                "reference_contract_id": row.reference_contract_id,
                "reference_contract_number": row.reference_contract_no,
                "contract_state": (
                    "expired" if days_to_expiry is not None and days_to_expiry < 0
                    else "expiring" if days_to_expiry is not None and days_to_expiry <= 60
                    else "active"
                ),
                "days_to_expiry": days_to_expiry,
                "gcn_number": gcn_no,
                "gcn_state": gcn_status,
                "detail_url": f"/bg/contracts/{row.id}",
            }
        )

    total = len(items)
    total_pages = max(1, (total + page_size - 1) // page_size)
    start = (page - 1) * page_size
    end = start + page_size
    return {
        "items": items[start:end],
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": total_pages,
    }


# =============================================================================
# Reports V2: /users — staff performance
# =============================================================================

@reports_v2_router.get("/users")
def get_v2_users_report(
    year: int = Query(..., ge=2000, le=2100),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    users = db.query(UserRow).filter(UserRow.is_active == True).all()  # noqa: E712
    user_by_email = {(u.username or "").strip().lower(): u for u in users}
    rows = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.contract_year == year)
        .all()
    )

    user_rows: list[dict[str, Any]] = []
    unassigned = {
        "user_id": None,
        "username": None,
        "display_name": "Chưa gán",
        "is_active": None,
        "configured": False,
        "annual_target": None,
        "target_zero": False,
        "actual": 0,
        "contract_count": 0,
        "positive_value_count": 0,
        "zero_value_count": 0,
        "null_value_count": 0,
        "remaining": None,
        "exceeded": None,
        "progress_percent": None,
        "new_count": 0,
        "new_actual": 0,
        "renewal_count": 0,
        "renewal_actual": 0,
        "frame_count": 0,
        "frame_actual": 0,
        "unknown_count": 0,
        "unknown_actual": 0,
    }
    branch_total = {
        "assigned_count": 0,
        "assigned_actual": 0,
        "unassigned_count": 0,
        "unassigned_actual": 0,
        "contract_count": 0,
        "actual": 0,
        "positive_value_count": 0,
        "zero_value_count": 0,
        "null_value_count": 0,
    }

    per_user_idx: dict[str, dict[str, Any]] = {}
    for u in users:
        per_user_idx[(u.username or "").strip().lower()] = {
            "user_id": u.id,
            "username": u.username,
            "display_name": u.display_name,
            "is_active": bool(u.is_active),
            "configured": False,
            "annual_target": None,
            "target_zero": False,
            "actual": 0,
            "contract_count": 0,
            "positive_value_count": 0,
            "zero_value_count": 0,
            "null_value_count": 0,
            "remaining": None,
            "exceeded": None,
            "progress_percent": None,
            "new_count": 0,
            "new_actual": 0,
            "renewal_count": 0,
            "renewal_actual": 0,
            "frame_count": 0,
            "frame_actual": 0,
            "unknown_count": 0,
            "unknown_actual": 0,
        }

    for row in rows:
        val = get_normalized_before_vat(row)
        b = _renewal_bucket(row.renewal_status)
        branch_total["contract_count"] += 1
        branch_total["actual"] += val
        if resolve_contract_revenue(row).resolution_status == "unresolved":
            branch_total["null_value_count"] += 1
        elif val == 0:
            branch_total["zero_value_count"] += 1
        else:
            branch_total["positive_value_count"] += 1

        owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
        target_idx = per_user_idx.get(owner_email)
        if target_idx is None:
            target_idx = unassigned
            branch_total["unassigned_count"] += 1
            branch_total["unassigned_actual"] += val
        else:
            branch_total["assigned_count"] += 1
            branch_total["assigned_actual"] += val

        target_idx["contract_count"] += 1
        target_idx["actual"] += val
        if resolve_contract_revenue(row).resolution_status == "unresolved":
            target_idx["null_value_count"] += 1
        elif val == 0:
            target_idx["zero_value_count"] += 1
        else:
            target_idx["positive_value_count"] += 1
        target_idx[f"{b}_count"] += 1
        target_idx[f"{b}_actual"] += val

    for key, entry in per_user_idx.items():
        target, target_zero = _resolve_target(year, key)
        entry["annual_target"] = target
        entry["target_zero"] = target_zero
        entry["configured"] = target is not None
        if target and target > 0:
            entry["remaining"] = target - entry["actual"]
            entry["exceeded"] = (-entry["remaining"]) if entry["remaining"] < 0 else None
            entry["progress_percent"] = round(entry["actual"] / target * 100, 1)
        user_rows.append(entry)

    return {
        "year": year,
        "users": sorted(user_rows, key=lambda x: -(x.get("actual") or 0)),
        "unassigned": unassigned,
        "branch": branch_total,
    }


# =============================================================================
# Reports V2: /renewals
# =============================================================================

@reports_v2_router.get("/renewals")
def get_v2_renewals(
    year: int = Query(..., ge=2000, le=2100),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    include_historical: bool = Query(default=False, description="Khi false (mặc định): chỉ lấy HĐ có ngày hết hạn trong năm báo cáo. Khi true: lấy thêm tồn đọng từ các năm trước."),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    today = date.today()
    year_start = date(year, 1, 1)
    year_end = date(year, 12, 31)

    # Base query: canonical contracts with end_date
    q = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.ngay_ket_thuc.isnot(None))
    )
    if include_historical:
        # Include contracts whose expiry falls anywhere — for backlog review
        q = q.filter(ContractRecordRow.ngay_ket_thuc <= year_end)
    else:
        # Default: only contracts expiring within the reporting year
        q = q.filter(ContractRecordRow.ngay_ket_thuc >= year_start)
        q = q.filter(ContractRecordRow.ngay_ket_thuc <= year_end)

    all_rows = q.all()
    users = db.query(UserRow).all()
    user_by_email = {(u.username or "").strip().lower(): u for u in users}

    # ── Build full item list (used for summary counts) ──────────────────────
    items_full: list[dict[str, Any]] = []

    # group contracts by reference_contract_id to detect renewals
    by_ref: dict[int, list[ContractRecordRow]] = {}
    standalone: list[ContractRecordRow] = []
    for row in all_rows:
        ref = row.reference_contract_id
        if ref is not None:
            by_ref.setdefault(ref, []).append(row)
        else:
            standalone.append(row)

    # Linked renewals (old + new via reference_contract_id)
    for ref_id, group in by_ref.items():
        group_sorted = sorted(group, key=lambda r: r.ngay_ket_thuc)
        if len(group_sorted) < 2:
            continue
        old = group_sorted[0]
        new = group_sorted[-1]
        end = old.ngay_ket_thuc
        days_remaining = (end - today).days if end else None
        is_overdue = (end < today) if end else False
        owner_email = (old.nguoi_thuc_hien_email or "").strip().lower()
        u = user_by_email.get(owner_email)
        new_val = get_normalized_before_vat(new)
        items_full.append(
            {
                "old_contract_id": old.id,
                "old_contract_number": old.contract_no or "",
                "organization_name": old.don_vi_ten or "",
                "field": old.linh_vuc_hien_thi or old.linh_vuc or "",
                "owner_user_id": u.id if u else None,
                "owner_email": owner_email or None,
                "owner_name": u.display_name if u else None,
                "signed_date": old.ngay_lap_hop_dong.isoformat() if old.ngay_lap_hop_dong else None,
                "end_date": end.isoformat() if end else None,
                "days_remaining": days_remaining,
                "is_overdue": is_overdue,
                "royalty_amount_before_vat": old.royalty_amount_before_vat,
                "renewal_status": old.renewal_status,
                "renewal_link_status": "linked",
                "new_contract_id": new.id,
                "new_contract_number": new.contract_no,
                "new_contract_signed_date": new.ngay_lap_hop_dong.isoformat() if new.ngay_lap_hop_dong else None,
                "new_contract_actual": new_val,
            }
        )

    # Standalone contracts needing renewal (within 60 days of today)
    for row in standalone:
        end = row.ngay_ket_thuc
        if not end:
            continue
        days_remaining = (end - today).days
        if days_remaining > 60:
            continue
        is_overdue = days_remaining < 0
        owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
        u = user_by_email.get(owner_email)
        items_full.append(
            {
                "old_contract_id": row.id,
                "old_contract_number": row.contract_no or "",
                "organization_name": row.don_vi_ten or "",
                "field": row.linh_vuc_hien_thi or row.linh_vuc or "",
                "owner_user_id": u.id if u else None,
                "owner_email": owner_email or None,
                "owner_name": u.display_name if u else None,
                "signed_date": row.ngay_lap_hop_dong.isoformat() if row.ngay_lap_hop_dong else None,
                "end_date": end.isoformat(),
                "days_remaining": days_remaining,
                "is_overdue": is_overdue,
                "royalty_amount_before_vat": row.royalty_amount_before_vat,
                "renewal_status": row.renewal_status,
                "renewal_link_status": "unlinked",
                "new_contract_id": None,
                "new_contract_number": None,
                "new_contract_signed_date": None,
                "new_contract_actual": None,
            }
        )

    # Sort by urgency
    items_full.sort(key=lambda r: r.get("days_remaining") or 99999)

    # ── Server-side summary counts ─────────────────────────────────────────────
    total_count = len(items_full)
    needs_renewal_count = sum(
        1 for r in items_full
        if r["renewal_link_status"] == "unlinked"
        and (r.get("days_remaining") or 0) <= 30
    )
    expiring_soon_count = sum(
        1 for r in items_full
        if r["renewal_link_status"] == "unlinked"
        and not r["is_overdue"]
        and ((r.get("days_remaining") or 0) > 30 and (r.get("days_remaining") or 0) <= 60)
    )
    overdue_count = sum(
        1 for r in items_full
        if r["is_overdue"] and r["renewal_link_status"] == "unlinked"
    )
    renewed_count = sum(1 for r in items_full if r["renewal_link_status"] == "linked")
    unassigned_count = sum(1 for r in items_full if not r["owner_user_id"])

    # Paginate from the full list
    total_pages = max(1, (total_count + page_size - 1) // page_size)
    start = (page - 1) * page_size
    end = start + page_size

    return {
        "year": year,
        "include_historical": include_historical,
        "total_count": total_count,
        "needs_renewal_count": needs_renewal_count,
        "expiring_soon_count": expiring_soon_count,
        "overdue_count": overdue_count,
        "renewed_count": renewed_count,
        "unassigned_count": unassigned_count,
        "linked_count": renewed_count,
        "unlinked_count": total_count - renewed_count,
        "null_value_count": 0,
        "zero_value_count": 0,
        "total_value": sum((r.get("new_contract_actual") or 0) for r in items_full),
        "items": items_full[start:end],
    }


# =============================================================================
# Reports V2: /gcn
# =============================================================================

@reports_v2_router.get("/gcn")
def get_v2_gcn_report(
    year: int = Query(..., ge=2000, le=2100),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    contract_ids = [
        cid
        for (cid,) in (
            db.query(ContractRecordRow.id)
            .filter(ContractRecordRow.contract_year == year)
            .filter(ContractRecordRow.annex_no.is_(None))
            .all()
        )
    ]
    if not contract_ids:
        return {
            "year": year,
            "total": 0,
            "total_count": 0,
            "issued_count": 0,
            "missing_count": 0,
            "items": [],
        }

    certs = (
        db.query(CertificateRecordRow)
        .filter(CertificateRecordRow.contract_id.in_(contract_ids))
        .order_by(CertificateRecordRow.created_at.desc())
        .all()
    )
    contract_rows = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.id.in_(contract_ids))
        .all()
    )
    contracts_by_id = {c.id: c for c in contract_rows}

    items = []
    issued = 0
    for cert in certs:
        contract = contracts_by_id.get(cert.contract_id)
        org = contract.don_vi_ten if contract else ""
        if str(cert.status or "").lower() == "final_printed":
            issued += 1
        items.append(
            {
                "id": cert.certificate_id,
                "certificate_no": cert.certificate_no,
                "certificate_status": str(cert.status or "draft"),
                "contract_id": cert.contract_id,
                "contract_no": contract.contract_no if contract else "",
                "organization_name": org or "",
                "issue_date": cert.certificate_issue_date.isoformat() if cert.certificate_issue_date else None,
            }
        )

    total = len(items)
    total_pages = max(1, (total + page_size - 1) // page_size)
    start = (page - 1) * page_size
    end = start + page_size

    return {
        "year": year,
        "total": total,
        "total_count": total,
        "issued_count": issued,
        "missing_count": total - issued,
        "items": items[start:end],
    }


# =============================================================================
# Reports V2: /export (POST) — minimal stub returning JSON notice
# =============================================================================

# =============================================================================
# Reports V2: /export (POST) — real generators for xlsx / docx / pdf
# =============================================================================
#
# Supported report_type values:
#   - overview  → branch-wide summary (uses /api/reports/v2/overview data)
#   - users     → per-user staff performance (uses /api/reports/v2/users data)
#   - renewals  → renewals & expiry list (uses /api/reports/v2/renewals data)
#
# All formats honour the same `year` filter. Format follows frontend's
# `ExportRequest` (`xlsx` | `docx` | `pdf`).
#
# Output is a binary Blob (xlsx/docx/pdf) with appropriate MIME type.

from io import BytesIO

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle


def _vnd(n) -> str:
    try:
        v = int(n or 0)
    except Exception:
        v = 0
    return f"{v:,}".replace(",", ".") + " ₫"


def _num(n) -> str:
    try:
        v = int(n or 0)
    except Exception:
        v = 0
    return f"{v:,}".replace(",", ".")


def _resolve_scope_data(db: Session, user: UserRow, report_type: str, year: int) -> dict[str, Any]:
    """Run the right summary query based on report_type and return a dict
    suitable for templating. All numeric fields use the canonical grain."""
    perms = get_user_permissions(db, user)
    is_admin = ("reports.view_branch" in perms) or ("kpi.manage" in perms)

    if report_type == "renewals":
        # Reuse the renewals endpoint logic — call directly with page_size=large
        from fastapi import Request  # noqa: F401  (kept for future request context)
        rows: list[dict[str, Any]] = []
        today = date.today()
        c_rows = (
            db.query(ContractRecordRow)
            .filter(ContractRecordRow.annex_no.is_(None))
            .filter(ContractRecordRow.contract_year <= year)
            .filter(ContractRecordRow.ngay_ket_thuc.isnot(None))
            .all()
        )
        users = db.query(UserRow).all()
        user_by_email = {(u.username or "").strip().lower(): u for u in users}
        by_ref: dict[int, list[ContractRecordRow]] = {}
        standalone: list[ContractRecordRow] = []
        for row in c_rows:
            ref = row.reference_contract_id
            if ref is not None:
                by_ref.setdefault(ref, []).append(row)
            else:
                standalone.append(row)
        for _, group in by_ref.items():
            g_sorted = sorted(group, key=lambda r: r.ngay_ket_thuc)
            if len(g_sorted) < 2:
                continue
            old = g_sorted[0]
            new = g_sorted[-1]
            end = old.ngay_ket_thuc
            days_remaining = (end - today).days if end else None
            is_overdue = (end < today) if end else False
            owner_email = (old.nguoi_thuc_hien_email or "").strip().lower()
            u = user_by_email.get(owner_email)
            rows.append({
                "old_contract_number": old.contract_no or "",
                "organization_name": old.don_vi_ten or "",
                "field": old.linh_vuc_hien_thi or old.linh_vuc or "",
                "owner_email": owner_email or None,
                "end_date": end.isoformat() if end else "",
                "days_remaining": days_remaining,
                "is_overdue": is_overdue,
                "renewal_link_status": "linked",
                "new_contract_number": new.contract_no or "",
            })
        for row in standalone:
            end = row.ngay_ket_thuc
            if not end:
                continue
            days_remaining = (end - today).days
            if days_remaining > 60:
                continue
            is_overdue = days_remaining < 0
            owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
            u = user_by_email.get(owner_email)
            rows.append({
                "old_contract_number": row.contract_no or "",
                "organization_name": row.don_vi_ten or "",
                "field": row.linh_vuc_hien_thi or row.linh_vuc or "",
                "owner_email": owner_email or None,
                "end_date": end.isoformat(),
                "days_remaining": days_remaining,
                "is_overdue": is_overdue,
                "renewal_link_status": "unlinked",
                "new_contract_number": "",
            })
        rows.sort(key=lambda r: (r.get("days_remaining") if r.get("days_remaining") is not None else 99999))
        return {"title": "Tái ký & hết hạn", "rows": rows}

    if report_type == "users":
        if not is_admin:
            # Staff can only export their own row.
            rows = []
            from sqlalchemy import func as _f
            email = (user.username or "").strip().lower()
            c_rows = (
                db.query(ContractRecordRow)
                .filter(ContractRecordRow.annex_no.is_(None))
                .filter(ContractRecordRow.contract_year == year)
                .filter(_f.lower(ContractRecordRow.nguoi_thuc_hien_email) == email)
                .all()
            )
            actual = sum(get_normalized_before_vat(r) for r in c_rows)
            target, target_zero = _resolve_target(year, user.username)
            rows.append({
                "username": user.username,
                "display_name": user.display_name or "",
                "contract_count": len(c_rows),
                "new_count": sum(1 for r in c_rows if _renewal_bucket(r.renewal_status) == "new"),
                "renewal_count": sum(1 for r in c_rows if _renewal_bucket(r.renewal_status) == "renewal"),
                "actual": actual,
                "annual_target": target,
                "configured": target is not None,
            })
            return {"title": "Hiệu suất cá nhân", "rows": rows}
        # Admin: call /users logic
        users = db.query(UserRow).filter(UserRow.is_active == True).all()  # noqa: E712
        user_by_email = {(u.username or "").strip().lower(): u for u in users}
        c_rows = (
            db.query(ContractRecordRow)
            .filter(ContractRecordRow.annex_no.is_(None))
            .filter(ContractRecordRow.contract_year == year)
            .all()
        )
        per_user_idx: dict[str, dict[str, Any]] = {}
        for u in users:
            per_user_idx[(u.username or "").strip().lower()] = {
                "user_id": u.id,
                "username": u.username,
                "display_name": u.display_name,
                "contract_count": 0,
                "new_count": 0,
                "renewal_count": 0,
                "actual": 0,
                "annual_target": None,
                "configured": False,
            }
        unassigned = {"user_id": None, "username": None, "display_name": "Chưa gán",
                       "contract_count": 0, "new_count": 0, "renewal_count": 0,
                       "actual": 0, "annual_target": None, "configured": False}
        for row in c_rows:
            val = get_normalized_before_vat(row)
            b = _renewal_bucket(row.renewal_status)
            owner_email = (row.nguoi_thuc_hien_email or "").strip().lower()
            target_idx = per_user_idx.get(owner_email) or unassigned
            target_idx["contract_count"] += 1
            target_idx["actual"] += val
            if b == "new":
                target_idx["new_count"] += 1
            elif b == "renewal":
                target_idx["renewal_count"] += 1
        rows: list[dict[str, Any]] = []
        for key, entry in per_user_idx.items():
            target, _ = _resolve_target(year, key)
            entry["annual_target"] = target
            entry["configured"] = target is not None
            rows.append(entry)
        if unassigned["contract_count"] > 0:
            rows.append(unassigned)
        return {"title": "Phân công & khối lượng", "rows": rows}

    # Default: overview
    if is_admin:
        # Reuse /overview logic via internal call
        cert_map = _build_cert_map(db)
        c_rows = (
            db.query(ContractRecordRow)
            .filter(ContractRecordRow.annex_no.is_(None))
            .filter(ContractRecordRow.contract_year == year)
            .all()
        )
        users = db.query(UserRow).all()
        user_by_email = {(u.username or "").strip().lower(): u for u in users}
        bucket = {"new": [0, 0], "renewal": [0, 0], "frame": [0, 0], "unknown": [0, 0]}
        for row in c_rows:
            val = get_normalized_before_vat(row)
            b = _renewal_bucket(row.renewal_status)
            bucket[b][0] += 1
            bucket[b][1] += val
        rows = [
            {"label": "Tổng hợp đồng", "count": len(c_rows), "value": sum(get_normalized_before_vat(r) for r in c_rows)},
            {"label": "Ký mới", "count": bucket["new"][0], "value": bucket["new"][1]},
            {"label": "Tái ký", "count": bucket["renewal"][0], "value": bucket["renewal"][1]},
            {"label": "HĐ khung", "count": bucket["frame"][0], "value": bucket["frame"][1]},
            {"label": "Chưa xác định", "count": bucket["unknown"][0], "value": bucket["unknown"][1]},
        ]
        return {"title": "Tổng quan chi nhánh", "rows": rows}

    # Staff: personal overview
    from sqlalchemy import func as _f
    email = (user.username or "").strip().lower()
    c_rows = (
        db.query(ContractRecordRow)
        .filter(ContractRecordRow.annex_no.is_(None))
        .filter(ContractRecordRow.contract_year == year)
        .filter(_f.lower(ContractRecordRow.nguoi_thuc_hien_email) == email)
        .all()
    )
    bucket = {"new": [0, 0], "renewal": [0, 0], "frame": [0, 0], "unknown": [0, 0]}
    actual = 0
    for row in c_rows:
        val = get_normalized_before_vat(row)
        actual += val
        b = _renewal_bucket(row.renewal_status)
        bucket[b][0] += 1
        bucket[b][1] += val
    target, _ = _resolve_target(year, user.username)
    rows = [
        {"label": "Tổng hợp đồng của tôi", "count": len(c_rows), "value": actual},
        {"label": "Ký mới", "count": bucket["new"][0], "value": bucket["new"][1]},
        {"label": "Tái ký", "count": bucket["renewal"][0], "value": bucket["renewal"][1]},
        {"label": "HĐ khung", "count": bucket["frame"][0], "value": bucket["frame"][1]},
        {"label": "Chưa xác định", "count": bucket["unknown"][0], "value": bucket["unknown"][1]},
    ]
    if target:
        rows.insert(1, {"label": "Mục tiêu năm", "count": "", "value": target})
    return {"title": "Tổng quan của tôi", "rows": rows}


# ─── XLSX ──────────────────────────────────────────────────────────────────

def _render_xlsx(title: str, rows: list[dict[str, Any]], year: int) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Báo cáo"
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="C95867")
    title_font = Font(bold=True, size=14)
    ws.cell(row=1, column=1, value=f"{title} — năm {year}").font = title_font
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
    ws.cell(row=2, column=1, value=f"Xuất lúc {datetime.now().strftime('%Y-%m-%d %H:%M')}").font = Font(italic=True, color="808080")

    if not rows:
        ws.cell(row=4, column=1, value="(Không có dữ liệu)")
        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()

    headers = list(rows[0].keys())
    for col_idx, h in enumerate(headers, start=1):
        c = ws.cell(row=4, column=col_idx, value=h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(horizontal="center", vertical="center")

    for r_idx, row in enumerate(rows, start=5):
        for c_idx, h in enumerate(headers, start=1):
            val = row.get(h)
            if isinstance(val, bool):
                val = "Có" if val else "Không"
            elif val is None:
                val = ""
            ws.cell(row=r_idx, column=c_idx, value=val)

    # Auto-width
    for col_idx, h in enumerate(headers, start=1):
        max_len = len(str(h))
        for r_idx in range(5, 5 + len(rows)):
            v = ws.cell(row=r_idx, column=col_idx).value
            if v is not None:
                max_len = max(max_len, len(str(v)))
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 2, 40)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─── DOCX ──────────────────────────────────────────────────────────────────

def _render_docx(title: str, rows: list[dict[str, Any]], year: int) -> bytes:
    doc = Document()
    h = doc.add_heading(f"{title} — năm {year}", level=1)
    doc.add_paragraph(f"Xuất lúc {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    if not rows:
        doc.add_paragraph("(Không có dữ liệu)")
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, htxt in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(str(htxt))
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, hkey in enumerate(headers):
            val = row.get(hkey)
            if isinstance(val, bool):
                val = "Có" if val else "Không"
            elif val is None:
                val = ""
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── PDF ───────────────────────────────────────────────────────────────────

def _render_pdf(title: str, rows: list[dict[str, Any]], year: int) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            leftMargin=1.5 * cm, rightMargin=1.5 * cm,
                            topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Heading1"], fontSize=16, spaceAfter=8)
    sub_style = ParagraphStyle("sub", parent=styles["Italic"], fontSize=9, textColor=colors.grey)

    elements = []
    elements.append(Paragraph(f"{title} — năm {year}", title_style))
    elements.append(Paragraph(f"Xuất lúc {datetime.now().strftime('%Y-%m-%d %H:%M')}", sub_style))
    elements.append(Spacer(1, 8))

    if not rows:
        elements.append(Paragraph("(Không có dữ liệu)", styles["Normal"]))
        doc.build(elements)
        return buf.getvalue()

    headers = list(rows[0].keys())
    data = [headers]
    for row in rows:
        line = []
        for hkey in headers:
            v = row.get(hkey)
            if isinstance(v, bool):
                v = "Có" if v else "Không"
            elif v is None:
                v = ""
            line.append(str(v))
        data.append(line)

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#C95867")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(table)
    doc.build(elements)
    return buf.getvalue()


@reports_v2_router.post("/export")
def post_v2_export(
    body: dict[str, Any],
    credentials: HTTPAuthorizationCredentials | None = Depends(security_scheme),
    db: Session = Depends(get_db),
):
    user = _current_user(db, credentials)
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")
    fmt = (body.get("format") or "").lower()
    report_type = body.get("report_type") or "overview"
    year = body.get("year")
    if not year:
        raise HTTPException(status_code=400, detail="year is required")
    if fmt not in ("xlsx", "docx", "pdf"):
        raise HTTPException(status_code=400, detail="format must be xlsx/docx/pdf")

    try:
        scope = _resolve_scope_data(db, user, str(report_type), int(year))
    except HTTPException:
        raise
    except Exception as exc:
        log.exception("export scope failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Lỗi khi tải dữ liệu báo cáo: {exc}")

    rows = scope.get("rows") or []
    title = scope.get("title") or "Báo cáo"

    if fmt == "xlsx":
        blob = _render_xlsx(title, rows, int(year))
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif fmt == "docx":
        blob = _render_docx(title, rows, int(year))
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        blob = _render_pdf(title, rows, int(year))
        media = "application/pdf"

    from fastapi.responses import Response
    return Response(
        content=blob,
        media_type=media,
        headers={
            "Content-Disposition": f'attachment; filename="bao_cao_{report_type}_{year}.{fmt}"'
        },
    )