"""DOCX validation helper for verifying Word file integrity.

This module provides utilities to validate DOCX files after rendering
to ensure they can be opened by Microsoft Word.
"""
from __future__ import annotations

import logging
import shutil
import zipfile
from pathlib import Path

from docx import Document

logger = logging.getLogger("uvicorn.error")

# Sentinel used in karaoke templates for block insertion
KARAOKE_ROOM_BLOCK_SENTINEL = "__KARAOKE_ROOM_BLOCK__"
KVC_USAGE_BLOCK_SENTINEL = "__KVC_USAGE_BLOCK__"


def validate_docx_can_open(path: Path | str, *, strict: bool = False) -> tuple[bool, str | None]:
    """Validate that a DOCX file can be opened by python-docx.

    This catches XML corruption, malformed structures, and other issues
    that would cause Microsoft Word to refuse opening the file.

    Args:
        path: Path to the DOCX file to validate.
        strict: If True, raises exceptions on validation failure.

    Returns:
        Tuple of (is_valid, error_message).
        is_valid is True if the file can be opened.
        error_message is None if valid, or a description of the error.
    """
    path = Path(path)
    if not path.exists():
        msg = f"File not found: {path}"
        logger.error(f"[DOCX_VALIDATE] FAIL: {msg}")
        if strict:
            raise FileNotFoundError(msg)
        return False, msg

    try:
        doc = Document(str(path))
        # Basic sanity check: document should have a body
        if doc.element.body is None:
            msg = "Document body is None"
            logger.error(f"[DOCX_VALIDATE] FAIL: {msg}")
            return False, msg
        logger.info(f"[DOCX_VALIDATE] OK: {path}")
        return True, None

    except zipfile.BadZipFile as e:
        msg = f"Bad ZIP file (corrupted archive): {e}"
        logger.error(f"[DOCX_VALIDATE] FAIL: {msg}")
        if strict:
            raise
        return False, msg

    except Exception as e:
        msg = f"Failed to open DOCX: {type(e).__name__}: {e}"
        logger.error(f"[DOCX_VALIDATE] FAIL: {msg}")
        if strict:
            raise
        return False, msg


def validate_and_save_debug_copy(
    output_path: Path | str,
    *,
    debug_dir: Path | str | None = None,
) -> Path | None:
    """Validate DOCX and save a debug copy if validation fails.

    Args:
        output_path: Path to the rendered DOCX.
        debug_dir: Directory to save debug copy. Defaults to same dir as output.

    Returns:
        Path to debug copy if validation failed, None otherwise.
    """
    output_path = Path(output_path)
    is_valid, error_msg = validate_docx_can_open(output_path)

    if is_valid:
        return None

    # Save debug copy
    if debug_dir is None:
        debug_dir = output_path.parent
    debug_dir = Path(debug_dir)
    debug_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    debug_name = f"DEBUG_CORRUPT_{output_path.stem}_{timestamp}{output_path.suffix}"
    debug_path = debug_dir / debug_name

    try:
        shutil.copy2(output_path, debug_path)
        logger.warning(
            f"[DOCX_VALIDATE] Saved debug copy: {debug_path} "
            f"(reason: {error_msg})"
        )
        return debug_path
    except Exception as e:
        logger.error(f"[DOCX_VALIDATE] Failed to save debug copy: {e}")
        return None


def check_template_placeholders(
    template_path: Path | str,
    required_placeholders: list[str],
) -> tuple[bool, list[str]]:
    """Check that a template contains required placeholders.

    Args:
        template_path: Path to the DOCX template.
        required_placeholders: List of placeholder names (without {{}}).

    Returns:
        Tuple of (all_found, missing_placeholders).
    """
    import re

    template_path = Path(template_path)
    if not template_path.exists():
        return False, required_placeholders

    try:
        with zipfile.ZipFile(template_path, "r") as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")

        found = []
        missing = []
        for ph in required_placeholders:
            # Build pattern without backslash in f-string expression
            pattern = r"\{\{" + "{" + ph + "}" + r"\}\}"
            if re.search(pattern, xml):
                found.append(ph)
            else:
                missing.append(ph)

        return len(missing) == 0, missing

    except Exception as e:
        logger.error(f"[TEMPLATE_CHECK] Failed to check placeholders: {e}")
        return False, required_placeholders


def check_output_has_table(output_path: Path | str, *, min_rows: int = 1) -> tuple[bool, int]:
    """Check that DOCX contains at least one table with minimum rows.

    Args:
        output_path: Path to the DOCX file.
        min_rows: Minimum number of rows expected in the table.

    Returns:
        Tuple of (has_table, actual_row_count).
    """
    try:
        doc = Document(str(output_path))
        for table in doc.tables:
            if len(table.rows) >= min_rows:
                return True, len(table.rows)
        return False, 0
    except Exception as e:
        logger.error(f"[TABLE_CHECK] Failed: {e}")
        return False, 0


def check_placeholder_preserved(
    output_path: Path | str,
    placeholder_name: str,
) -> bool:
    """Check that a placeholder is still present in the output DOCX.

    This is used for PRESERVED placeholders like {{tien_ban_quyen}}
    that should NOT be modified by the renderer.

    Args:
        output_path: Path to the DOCX file.
        placeholder_name: Placeholder name (without {{}}).

    Returns:
        True if the placeholder text is found in the output.
    """
    import re

    output_path = Path(output_path)
    try:
        with zipfile.ZipFile(output_path, "r") as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")

        pattern = r"\{\{" + "{" + placeholder_name + "}" + r"\}\}"
        return bool(re.search(pattern, xml))

    except Exception as e:
        logger.error(f"[PLACEHOLDER_CHECK] Failed: {e}")
        return False


# Import datetime for debug copy naming
from datetime import datetime
